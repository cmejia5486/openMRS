# scripts/ai_correlate.py
import os, sys, json, re, argparse, time
from pathlib import Path
from typing import List, Dict, Any, Optional
import orjson
from tqdm import tqdm

# OpenAI 1.x (opcional)
USE_OPENAI = False
try:
    from openai import OpenAI
    USE_OPENAI = True
except Exception:
    USE_OPENAI = False

# DOCX
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

CODE_EXT = {".kt",".java",".xml",".gradle",".kts",".properties",".json",".yml",".yaml"}

PATTERNS = {
    "ssl_allow_all": r"ALLOW_ALL_HOSTNAME_VERIFIER|TrustAll|InsecureTrustManager|verify\(\)\s*{\s*return\s+true",
    "webview_js": r"setJavaScriptEnabled\s*\(\s*true\s*\)",
    "add_js_interface": r"addJavascriptInterface\s*\(",
    "hardcoded_key": r"(api[_-]?key|secret|token)\s*[:=]\s*['\"][A-Za-z0-9_\-]{12,}",
    "exported_activity": r'android:exported="true"',
    "plaintext_http": r'http://[A-Za-z0-9\.\-]',
}

SYSTEM_PROMPT = """Eres un auditor de seguridad móvil senior.
Emites veredictos por requisito (Yes/No/N_a/Insufficient_Evidence) con justificación breve (<=6 líneas), severidad y referencias a evidencias.
Si la evidencia es ambigua o insuficiente, usa 'Insufficient_Evidence' y explica qué faltó.
Responde SIEMPRE en JSON con campos: puid, status, severity, rationale, references (lista de strings), tags (lista)."""

def load_json_any(path: Path) -> Optional[Any]:
    if not path or not path.exists():
        return None
    try:
        return orjson.loads(path.read_bytes())
    except Exception:
        try:
            return json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            return None

def find_files(root: Path, patterns: List[str]) -> List[Path]:
    out: List[Path] = []
    for pat in patterns:
        out.extend(root.rglob(pat))
    # quitar duplicados conservando orden
    seen = set(); res=[]
    for p in out:
        if p.resolve() not in seen:
            seen.add(p.resolve()); res.append(p)
    return res

def load_requirements(req_path: Path) -> List[Dict[str, Any]]:
    raw = load_json_any(req_path)
    if isinstance(raw, dict) and "requirements" in raw:
        data = raw["requirements"]
    else:
        data = raw if isinstance(raw, list) else []
    out = []
    for i, d in enumerate(data, 1):
        if not isinstance(d, dict): 
            continue
        puid = d.get("PUID") or d.get("puid") or d.get("id") or f"REQ-{i:04d}"
        # prioridad a "Requirement description"
        text = (d.get("Requirement description") or d.get("Requirement") or
                d.get("requirement") or d.get("controles") or
                d.get("description") or d.get("Description") or
                d.get("text") or d.get("Text") or "")
        if not str(text).strip():
            # fallback concatenado corto
            text = " ".join([str(v) for v in d.values() if isinstance(v, (str,int,float))])[:2000]
        out.append({"id": str(puid).strip(), "text": str(text).strip(), "raw": d})
    return out

def summarize_trivy(reports_dir: Path) -> Dict[str, Any]:
    # prefer payload enriquecido
    cand = find_files(reports_dir, ["agent_payload.json","trivy.json"])
    js = load_json_any(cand[0]) if cand else {}
    if isinstance(js, dict) and "summary" in js:
        return js
    if isinstance(js, dict) and "Results" in js:
        return {"findings": js.get("Results"), "summary": {"source": "trivy.json"}}
    return {}

def summarize_mobsf(reports_dir: Path) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    static = None; dynamic = None
    cand = find_files(reports_dir, ["*mobsf*static*.json","*mobsf*results*.json","*mobsf*report*.json","*report_json*.json"])
    if cand:
        static = load_json_any(cand[0])
    cand2 = find_files(reports_dir, ["*mobsf*dynamic*.json"])
    if cand2:
        dynamic = load_json_any(cand2[0])
    sev = {"CRITICAL":0,"HIGH":0,"MEDIUM":0,"LOW":0,"INFO":0}
    if isinstance(static, dict):
        try:
            # intenta recorrer listas con 'severity' o 'Severity'
            for v in static.values():
                if isinstance(v, list):
                    for it in v:
                        s = (it.get("severity") or it.get("Severity") or "").upper()
                        if s in sev: sev[s]+=1
        except Exception:
            pass
    if static: out["static"] = {"severity_counts": sev, "has_static": True}
    if dynamic: out["dynamic"] = {"has_dynamic": True}
    return out

def summarize_sast(reports_dir: Path) -> Dict[str, Any]:
    sarifs = find_files(reports_dir, ["*.sarif"])
    findings = []
    for s in sarifs:
        js = load_json_any(s) or {}
        for run in (js.get("runs") or []):
            tool = (((run.get("tool") or {}).get("driver") or {}).get("name") or "sarif-tool")
            for r in (run.get("results") or []):
                level = r.get("level") or ((r.get("properties") or {}).get("severity")) or "warning"
                rule = r.get("ruleId") or ((r.get("rule") or {}).get("id")) or ""
                msg  = ((r.get("message") or {}).get("text")) or ""
                findings.append({"tool": tool, "level": str(level), "rule": str(rule), "message": msg})
    # también JSON SAST genéricos
    for g in find_files(reports_dir, ["sast*.json","*codeql*results*.json"]):
        js = load_json_any(g) or {}
        if isinstance(js, list):
            for it in js[:1000]:
                findings.append({"tool":"sast-json","level":str(it.get("severity","")), "rule":str(it.get("ruleId","")), "message":str(it.get("message",""))})
    return {"findings": findings[:5000]}

def scan_codebase(root: Path, max_bytes: int = 2_000_000) -> Dict[str,int]:
    totals = {k:0 for k in PATTERNS}
    size = 0
    for p in root.rglob("*"):
        if not p.is_file(): 
            continue
        if p.suffix.lower() not in CODE_EXT: 
            continue
        try:
            b = p.read_bytes()
            size += len(b)
            if size > max_bytes: 
                break
            txt = b.decode("utf-8", errors="ignore")
            for name, rx in PATTERNS.items():
                if re.search(rx, txt, flags=re.IGNORECASE|re.MULTILINE):
                    totals[name]+=1
        except Exception:
            continue
    return totals

def build_openai() -> Optional["OpenAI"]:
    if not USE_OPENAI:
        return None
    api_key = os.environ.get("OPENAI_API_KEY","").strip()
    if not api_key:
        return None
    try:
        # Se puede configurar timeout si lo deseas
        return OpenAI()
    except Exception:
        return None

def ask_llm(client: Optional["OpenAI"], model: str, req: Dict[str,Any], evidence: Dict[str,Any]) -> Dict[str,Any]:
    base = {"puid": req["id"], "status":"Insufficient_Evidence", "severity":"unknown",
            "rationale":"No AI available or insufficient inputs.", "references":[], "tags":[]}
    if not client:
        return base
    user = {
        "role": "user",
        "content": f"""
REQUISITO:
ID: {req['id']}
Texto: {req['text']}

CONTEXTO:
- Trivy summary: {json.dumps(evidence.get('trivy',{}))[:3500]}
- MobSF summary: {json.dumps(evidence.get('mobsf',{}))[:3500]}
- SAST sample: {json.dumps((evidence.get('sast') or [])[:20])[:3500]}
- Code patterns: {json.dumps(evidence.get('code',{}))}

INSTRUCCIONES:
1) Decide status (Yes/No/N_a/Insufficient_Evidence).
2) 1-2 frases de rationale.
3) Asigna severidad aproximada (critical/high/medium/low/unknown).
4) Lista 1..3 referencias (texto libre: indicios o archivo).
5) tags (p.ej., ['ssl','webview','hardcoded'])
6) Solo JSON, sin comentarios.
""".strip()
    }
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role":"system","content":SYSTEM_PROMPT}, user],
            temperature=0.2,
        )
        txt = (resp.choices[0].message.content or "").strip().strip("` ")
        if txt.lower().startswith("json"):
            txt = txt[4:].lstrip()
        js = json.loads(txt)
        js["puid"] = js.get("puid") or req["id"]
        return js
    except Exception as e:
        base["rationale"] = f"AI error: {e.__class__.__name__}"
        return base

def write_docx_seccat(path: Path, app_name: str, findings: List[Dict[str,Any]]):
    doc = Document()
    doc.add_heading("SRS for Security on mHealth applications (SEC-CAT*) – Adapted", level=0)
    p = doc.add_paragraph(f"Audited application: {app_name}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "PUID"
    hdr[1].text = "Requirement"
    hdr[2].text = "Status"
    hdr[3].text = "Severity"
    for f in findings:
        r = table.add_row().cells
        r[0].text = f.get("puid","")
        r[1].text = f.get("text","")
        r[2].text = f.get("status","")
        r[3].text = f.get("severity","")
    doc.save(path)

def write_docx_checklist(path: Path, findings: List[Dict[str,Any]]):
    doc = Document()
    doc.add_heading("Compliance Checklist – SEC-CAT*", level=0)
    t = doc.add_table(rows=1, cols=5)
    h = t.rows[0].cells
    h[0].text="PUID"; h[1].text="Requirement"; h[2].text="Yes"; h[3].text="No"; h[4].text="N/a"
    for f in findings:
        s = (f.get("status","") or "").lower()
        yes = "X" if s == "yes" else ""
        no  = "X" if s == "no" else ""
        na  = "X" if s in ("n_a","n/a","na") else ""
        row = t.add_row().cells
        row[0].text = f.get("puid","")
        row[1].text = f.get("text","")
        row[2].text = yes
        row[3].text = no
        row[4].text = na
    doc.save(path)

def write_docx_summary(path: Path, app_name: str, stats: Dict[str,Any], notes: List[str]):
    doc = Document()
    doc.add_heading("Audit Summary", level=0)
    doc.add_paragraph(f"Application: {app_name}")
    doc.add_paragraph(f"Total requirements evaluated: {stats.get('total',0)}")
    doc.add_paragraph(f"Results: Yes={stats.get('yes',0)}, No={stats.get('no',0)}, N/a={stats.get('na',0)}, Insufficient={stats.get('ins',0)}")
    doc.add_paragraph("Key observations:")
    for n in notes[:12]:
        doc.add_paragraph(f"• {n}")
    doc.save(path)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--checklist", required=True)
    ap.add_argument("--reports", required=True)
    ap.add_argument("--source-root", default=".")
    ap.add_argument("--output-dir", required=True)
    ap.add_argument("--openai-model", default=os.environ.get("OPENAI_MODEL","gpt-4o-mini"))
    ap.add_argument("--max-requirements", type=int, default=-1)
    args = ap.parse_args()

    outdir = Path(args.output_dir); outdir.mkdir(parents=True, exist_ok=True)
    reports_dir = Path(args.reports)
    app_name = Path(os.getenv("GITHUB_REPOSITORY","openMRS")).name

    # Requisitos (469, con PUID y "Requirement description")
    reqs = load_requirements(Path(args.checklist))
    if args.max_requirements and args.max_requirements > 0:
        reqs = reqs[: int(args.max_requirements)]

    # Artefactos
    trivy = summarize_trivy(reports_dir)
    mobsf = summarize_mobsf(reports_dir)
    sast  = summarize_sast(reports_dir)
    codep = scan_codebase(Path(args.source_root))

    client = build_openai()
    model  = args.openai_model

    findings = []
    for req in tqdm(reqs, desc="Auditing requirements"):
        evidence = {
            "trivy": trivy.get("summary") or trivy,
            "mobsf": {
                "static": mobsf.get("static",{}).get("severity_counts",{}),
                "dynamic": bool(mobsf.get("dynamic"))
            },
            "sast": (sast.get("findings") or [])[:50],
            "code": codep
        }
        verdict = ask_llm(client, model, req, evidence)
        findings.append({
            "puid": verdict.get("puid", req["id"]),
            "id": verdict.get("puid", req["id"]),
            "text": req["text"],
            "status": (verdict.get("status") or "Insufficient_Evidence"),
            "severity": (verdict.get("severity") or "unknown"),
            "rationale": verdict.get("rationale",""),
            "references": verdict.get("references",[]),
            "tags": verdict.get("tags",[]),
        })
        time.sleep(0.03)

    # Stats
    stats = {"total": len(findings), "yes":0,"no":0,"na":0,"ins":0}
    for f in findings:
        s = (f["status"] or "").lower()
        if s=="yes": stats["yes"]+=1
        elif s=="no": stats["no"]+=1
        elif s in ("n_a","n/a","na"): stats["na"]+=1
        else: stats["ins"]+=1

    # Persistencia
    (outdir/"audit-findings.json").write_text(
        json.dumps({"app": app_name, "findings": findings, "stats": stats}, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )
    (outdir/"audit-summary.md").write_text(
        f"# Audit Summary – {app_name}\n\n"
        f"- Total: {stats['total']} | Yes: {stats['yes']} | No: {stats['no']} | N/a: {stats['na']} | Insufficient: {stats['ins']}\n\n"
        f"## Notes\n\n"
        f"- MobSF dynamic present: {'yes' if mobsf.get('dynamic') else 'no'}\n"
        f"- Code patterns: {json.dumps(codep)}\n",
        encoding="utf-8"
    )

    # DOCX
    write_docx_seccat(outdir/"secm-cat_adapted.docx", app_name, findings)
    write_docx_checklist(outdir/"checklist.docx", findings)
    top_notes = [
        f"Trivy summary keys: {list((trivy.get('summary') or {}).keys())[:5]}",
        f"MobSF severities: {mobsf.get('static',{}).get('severity_counts',{})}",
        f"SAST findings (sample): {min(50, len(sast.get('findings',[])))}",
        f"Code flags: {codep}"
    ]
    write_docx_summary(outdir/"audit-summary.docx", app_name, stats, top_notes)

    print("Done. Reports in:", str(outdir))

if __name__ == "__main__":
    main()
