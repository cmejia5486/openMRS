#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import os
import re
from datetime import datetime, date
from pathlib import Path
from typing import Dict, Any, List

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

try:
    from lib.ai_runtime import AIRuntime
except Exception:
    AIRuntime = None  # type: ignore

def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _runtime_data_dir() -> Path:
    # Prefer explicit, cross-platform data-directory variables.
    # If none are provided, use GitHub Actions RUNNER_TEMP when available.
    # As a final local fallback, use a repository-local hidden directory.
    for name in ("VISION360_DATA_DIR", "AUDIT_DATA_DIR", "SECURITY_AUDIT_DATA_DIR"):
        raw = os.getenv(name, "").strip()
        if raw:
            base = Path(raw)
            base.mkdir(parents=True, exist_ok=True)
            return base

    runner_temp = os.getenv("RUNNER_TEMP", "").strip()
    if runner_temp:
        base = Path(runner_temp) / "vision360-data"
        base.mkdir(parents=True, exist_ok=True)
        return base

    base = _repo_root() / ".vision360-data"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _env_path(env_name: str, default_filename: str) -> Path:
    raw = os.getenv(env_name, "").strip()
    if raw:
        return Path(raw)
    return _runtime_data_dir() / default_filename


DEFAULT_IN = str(_env_path("AUDIT_ANALYSIS_JSON_PATH", "audit_summary_analysis_pack.json"))
DEFAULT_OUT = str(_env_path("AUDIT_SUMMARY_DOCX_PATH", "Audit Summary.docx"))
CHART_DIR = str(_env_path("AUDIT_SUMMARY_CHART_DIR", "_audit_summary_charts"))

HEADER_TEXT = "mSEC-AM Audit Summary - OpenMRS Android Client v3.1.1"


def _wrap_label(s: str, width: int = 30) -> str:
    s = str(s)
    if len(s) <= width:
        return s
    out = []
    while len(s) > width:
        cut = s.rfind(" ", 0, width)
        if cut == -1:
            cut = width
        out.append(s[:cut].strip())
        s = s[cut:].strip()
    if s:
        out.append(s)
    return "\n".join(out)


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _add_field_run(paragraph, field_instr: str) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def _add_header_footer(section, audit_date_str: str, report_title: str = "Mobile Application") -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = f"mSEC-AM Audit Summary - {report_title}"
    if p.runs:
        p.runs[0].font.size = Pt(9)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = f"{audit_date_str} | "
    if fp.runs:
        fp.runs[0].font.size = Pt(9)
    fp.add_run("Page ")
    _add_field_run(fp, " PAGE ")
    fp.add_run(" of ")
    _add_field_run(fp, " NUMPAGES ")
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_cover(doc: Document, audit_date_str: str, auditor: str, report_title: str = "Mobile Application") -> None:
    doc.add_paragraph()
    t = doc.add_paragraph("mSEC-AM Audit Summary")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(28)
    t.runs[0].bold = True
    s = doc.add_paragraph(report_title)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(16)
    s.runs[0].bold = True
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(f"Audit Date: {audit_date_str}\nAuditor: {auditor}\nClassification: Confidential / Internal Use")
    r.font.size = Pt(11)
    doc.add_page_break()


def _enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.xpath('./w:updateFields')
    if not existing:
        node = OxmlElement("w:updateFields")
        node.set(qn("w:val"), "true")
        settings.append(node)


def _make_bookmark_name(text: str, used: set[str]) -> str:
    base = re.sub(r"[^A-Za-z0-9_]+", "_", str(text)).strip("_")
    if not base:
        base = "section"
    if base[0].isdigit():
        base = f"s_{base}"
    name = base[:32]
    seed = name
    i = 2
    while name in used:
        suffix = f"_{i}"
        name = f"{seed[:32-len(suffix)]}{suffix}"
        i += 1
    used.add(name)
    return name


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def _add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _render_clickable_toc(after_paragraph, toc_entries: List[tuple[int, str, str]]) -> None:
    current = after_paragraph
    for level, title, anchor in toc_entries:
        p = _insert_paragraph_after(current)
        p.paragraph_format.left_indent = Inches(0.25 * max(level - 1, 0))
        p.paragraph_format.space_after = Pt(2)
        _add_internal_hyperlink(p, title, anchor)
        current = p


def _add_toc(doc: Document):
    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = doc.add_paragraph()
    doc.add_page_break()
    return p


def _add_two_col_table(doc: Document, rows: List[List[str]]) -> None:
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for k, v in rows:
        cells = tbl.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def _add_callout(doc: Document, title: str, bullets: List[str]) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, "EDEDED")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    for b in bullets:
        cell.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()


def _add_figure(doc: Document, img_path: str, caption: str) -> None:
    doc.add_picture(img_path, width=Inches(6.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def _donut(values: List[int], labels: List[str], title: str, center_text: str, out_path: str) -> None:
    fig, ax = plt.subplots(figsize=(7.4, 4.8))
    wedges, _, _ = ax.pie(values, labels=None, autopct="%1.1f%%", startangle=90, pctdistance=0.82)
    centre = plt.Circle((0, 0), 0.55, fc="white")
    fig.gca().add_artist(centre)
    ax.axis("equal")
    ax.set_title(title)
    ax.legend(wedges, labels, loc="center left", bbox_to_anchor=(1.02, 0.5), frameon=False)
    ax.text(0, 0, center_text, ha="center", va="center", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_share_noncompliances(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = sorted([(d["category_name"], int(d.get("non_compliant", 0))) for _, d in cat_stats.items()], key=lambda x: x[1])
    names = [x[0] for x in items]
    counts = [float(x[1]) for x in items]
    total = float(sum(counts))
    pct = [0.0 for _ in counts] if total <= 0.0 else [(c / total) * 100.0 for c in counts]
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    bars = ax.barh(y, pct)
    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(n, 30) for n in names])
    ax.set_xlabel("Share of total non-compliances (%)")
    ax.set_title("Share of non-compliances by category (high-level)")
    ax.set_xlim(0, max(5, float(max(pct) if pct else 0.0) + 5))
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    for b, c, p in zip(bars, counts, pct):
        ax.text(b.get_width() + 0.6, b.get_y() + b.get_height() / 2, f"{p:.1f}%  (n={int(c)})", va="center", fontsize=9)
    fig.text(0.01, 0.01, f"Source: audit workbook. Total non-compliances: {int(total)}.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_compliance_rate(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = []
    for _, d in cat_stats.items():
        app = int(d.get("applicable", 0))
        comp = int(d.get("compliant", 0))
        pct = float(d.get("compliance_pct", 0.0))
        items.append((d["category_name"], pct, comp, app))
    items = sorted(items, key=lambda x: x[1])
    names = [x[0] for x in items]
    pct_vals = [x[1] for x in items]
    comp_vals = [x[2] for x in items]
    app_vals = [x[3] for x in items]
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    bars = ax.barh(y, pct_vals)
    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(n, 30) for n in names])
    ax.set_xlabel("Compliance rate (%) - applicable controls only")
    ax.set_title("Compliance rate by category (applicable controls only)")
    ax.set_xlim(0, 100)
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    for b, p, c, a in zip(bars, pct_vals, comp_vals, app_vals):
        ax.text(p + 1, b.get_y() + b.get_height() / 2, f"{p:.1f}%  ({c}/{a})", va="center", fontsize=9)
    fig.text(0.01, 0.01, "Note: (c/a) indicates Compliant / Applicable controls per category. Source: audit workbook.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _stacked_counts(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = sorted([(d["category_name"], int(d.get("compliant", 0)), int(d.get("non_compliant", 0)), int(d.get("not_applicable", 0))) for _, d in cat_stats.items()], key=lambda x: (x[1] + x[2] + x[3]), reverse=True)
    names = [x[0] for x in items]
    c = [x[1] for x in items]
    n = [x[2] for x in items]
    na = [x[3] for x in items]
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    ax.barh(y, c, label="Compliant")
    ax.barh(y, n, left=c, label="Non-compliant")
    left_cn = [cc + nn for cc, nn in zip(c, n)]
    ax.barh(y, na, left=left_cn, label="Not applicable")
    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(nm, 30) for nm in names])
    ax.invert_yaxis()
    ax.set_xlabel("Count")
    ax.set_title("Counts by category and status")
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    ax.legend(loc="lower right", frameon=False)
    fig.tight_layout()
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _likelihood_from_count(cnt: int) -> str:
    if cnt >= 50:
        return "High"
    if cnt >= 20:
        return "Medium-High"
    if cnt >= 10:
        return "Medium"
    return "Low-Medium"


def _target_timeline(sev: str) -> str:
    return "0-90 days" if sev == "High" else ("0-180 days" if sev == "Medium" else "0-365 days")


def _target_date_str(audit_dt: date, sev: str) -> str:
    days = 90 if sev == "High" else (180 if sev == "Medium" else 365)
    due = audit_dt.toordinal() + days
    return date.fromordinal(due).strftime("%d %b %Y")



def _as_dict(value: Any) -> Dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if value is None:
        return []
    return [value]


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        if value is None:
            return default
        if isinstance(value, bool):
            return int(value)
        return int(float(str(value).strip()))
    except Exception:
        return default


def _safe_str(value: Any, default: str = "") -> str:
    if value is None:
        return default
    return str(value)


def _short(value: Any, limit: int = 170) -> str:
    text = re.sub(r"\s+", " ", _safe_str(value)).strip()
    if len(text) > limit:
        return text[:limit] + "..."
    return text


def _first_present(obj: Dict[str, Any], keys: List[str], default: Any = None) -> Any:
    for key in keys:
        if isinstance(obj, dict) and key in obj and obj[key] not in (None, ""):
            return obj[key]
    return default


def _deep_find_first(obj: Any, keys: List[str]) -> Any:
    if isinstance(obj, dict):
        for key in keys:
            if key in obj and obj[key] not in (None, ""):
                return obj[key]
        for value in obj.values():
            found = _deep_find_first(value, keys)
            if found not in (None, ""):
                return found
    elif isinstance(obj, list):
        for item in obj:
            found = _deep_find_first(item, keys)
            if found not in (None, ""):
                return found
    return None


def _deep_int(obj: Any, keys: List[str], default: int = 0) -> int:
    value = _deep_find_first(obj, keys)
    return _safe_int(value, default)


def _deep_dict(obj: Any, keys: List[str]) -> Dict[str, Any]:
    value = _deep_find_first(obj, keys)
    return value if isinstance(value, dict) else {}


def _deep_list(obj: Any, keys: List[str]) -> List[Any]:
    value = _deep_find_first(obj, keys)
    if isinstance(value, list):
        return value
    return []


def _block_available(block: Any) -> bool:
    if not isinstance(block, dict) or not block:
        return False
    if block.get("available") is True:
        return True
    for key, value in block.items():
        if key in {"available", "paths", "input_paths"}:
            continue
        if isinstance(value, (list, dict)) and len(value) > 0:
            return True
        if isinstance(value, (int, float)) and value > 0:
            return True
        if isinstance(value, str) and value.strip():
            return True
    return False


def _format_bool(value: Any) -> str:
    if value is True:
        return "Yes"
    if value is False:
        return "No"
    if value in (None, ""):
        return "Not reported"
    return _safe_str(value)


def _add_table(doc: Document, headers: List[str], rows: List[List[Any]], max_rows: int | None = None) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    for idx, txt in enumerate(headers):
        h[idx].text = str(txt)
        _set_cell_shading(h[idx], "D9E1F2")
        for run in h[idx].paragraphs[0].runs:
            run.bold = True
    actual_rows = rows[:max_rows] if max_rows is not None else rows
    if not actual_rows:
        actual_rows = [["No data available"] + [""] * (len(headers) - 1)]
    for row_values in actual_rows:
        r = tbl.add_row().cells
        for idx in range(len(headers)):
            r[idx].text = _short(row_values[idx] if idx < len(row_values) else "", 260)
    doc.add_paragraph()


def _app_report_title(app: Dict[str, Any]) -> str:
    app = _as_dict(app)
    name = _first_present(app, ["app_name", "application_name", "name", "title", "app", "project", "project_name"])
    version = _first_present(app, ["version", "app_version", "version_name", "release", "app_release"])

    if not name:
        for key, value in app.items():
            if "name" in str(key).lower() and value:
                name = value
                break
    if not version:
        for key, value in app.items():
            if "version" in str(key).lower() and value:
                version = value
                break

    name_s = _safe_str(name, "Mobile Application").strip() or "Mobile Application"
    version_s = _safe_str(version).strip()
    return f"{name_s} - {version_s}" if version_s and version_s not in name_s else name_s


def _technical_evidence(pack: Dict[str, Any]) -> Dict[str, Any]:
    return _as_dict(pack.get("technical_evidence"))


def _source_status_rows(technical: Dict[str, Any]) -> List[List[Any]]:
    labels = [
        ("Vision360 trace", "vision360"),
        ("Trivy SCA", "trivy_sca"),
        ("MobSF static", "mobsf_static"),
        ("MobSF dynamic", "mobsf_dynamic"),
        ("SAST app-code", "sast_app_code"),
    ]
    rows = []
    for label, key in labels:
        block = _as_dict(technical.get(key))
        available = _block_available(block)
        summary = "Available" if available else "Not available in analysis pack"
        if key == "trivy_sca" and available:
            summary = f"{_deep_int(block, ['total_vulnerabilities', 'vulnerabilities_total', 'total'])} vulnerability finding(s); {_deep_int(block, ['packages_detected', 'package_count', 'packages_total'])} package(s) detected."
        elif key == "sast_app_code" and available:
            summary = f"{_deep_int(block, ['total_findings', 'total', 'results_count'])} app-code finding(s) retained after scope filtering."
        elif key == "mobsf_static" and available:
            summary = "Static APK evidence parsed for manifest, certificate, permissions, signing, trackers, and hardening indicators."
        elif key == "mobsf_dynamic" and available:
            summary = "Runtime evidence parsed for local storage, databases, shared preferences, trackers, and observed artifacts."
        rows.append([label, "Yes" if available else "No", summary])
    return rows


def _technical_takeaways(technical: Dict[str, Any]) -> List[str]:
    takeaways: List[str] = []
    trivy = _as_dict(technical.get("trivy_sca"))
    if _block_available(trivy):
        total = _deep_int(trivy, ["total_vulnerabilities", "vulnerabilities_total", "total"])
        by_sev = _deep_dict(trivy, ["by_severity", "severity_counts"])
        critical = _safe_int(by_sev.get("CRITICAL") or by_sev.get("critical"))
        high = _safe_int(by_sev.get("HIGH") or by_sev.get("high"))
        if total > 0:
            takeaways.append(f"Trivy SCA evidence identified {total} known dependency vulnerability finding(s), including {critical} Critical and {high} High finding(s) where reported.")
    mobsf = _as_dict(technical.get("mobsf_static"))
    if _block_available(mobsf):
        indicators = []
        for label, keys in [
            ("debuggable build", ["debuggable", "has_debuggable", "debuggable_true", "android_debuggable"]),
            ("debug certificate", ["debug_certificate", "has_debug_cert", "signed_with_debug_certificate"]),
            ("backup enabled", ["allow_backup", "backup_enabled", "android_allow_backup"]),
            ("exported components", ["exported_components_count", "exported_component_count"]),
        ]:
            value = _deep_find_first(mobsf, keys)
            if value is True or _safe_int(value) > 0:
                indicators.append(label)
        if indicators:
            takeaways.append("MobSF static evidence reported Android hardening indicators requiring review, including " + ", ".join(indicators[:4]) + ".")
    sast = _as_dict(technical.get("sast_app_code"))
    if _block_available(sast):
        total = _deep_int(sast, ["total_findings", "total", "results_count"])
        if total > 0:
            takeaways.append(f"SAST evidence retained {total} application-code finding(s) after excluding CI/CD workflows and audit-generation scripts from scope.")
    return takeaways


def _trivy_findings(trivy: Dict[str, Any]) -> List[Dict[str, Any]]:
    for keys in (["cve_table"], ["findings"], ["vulnerabilities"], ["top_findings"]):
        items = _deep_list(trivy, keys)
        if items:
            return [x for x in items if isinstance(x, dict)]
    return []


def _add_trivy_section(doc: Document, technical: Dict[str, Any]) -> None:
    trivy = _as_dict(technical.get("trivy_sca"))
    if not _block_available(trivy):
        doc.add_paragraph("Trivy SCA evidence was not available in the analysis pack.")
        return
    total = _deep_int(trivy, ["total_vulnerabilities", "vulnerabilities_total", "total"])
    packages = _deep_int(trivy, ["packages_detected", "package_count", "packages_total"])
    fixable = _deep_int(trivy, ["fixable_total", "fixable_vulnerabilities", "fixable"])
    licenses = _deep_int(trivy, ["license_entries_detected", "licenses_detected", "license_count"])
    by_sev = _deep_dict(trivy, ["by_severity", "severity_counts"])
    severity_text = ", ".join(f"{k}: {_safe_int(v)}" for k, v in by_sev.items()) if by_sev else "not reported"

    doc.add_paragraph(
        f"Trivy Software Composition Analysis reported {packages} detected package(s), {total} known dependency vulnerability finding(s), "
        f"{fixable} finding(s) with a fixed version available, and {licenses} license entr(y/ies). Severity distribution: {severity_text}."
    )

    rows = []
    for finding in _trivy_findings(trivy)[:12]:
        rows.append([
            _first_present(finding, ["severity", "Severity"], ""),
            _first_present(finding, ["id", "VulnerabilityID", "cve", "rule_id"], ""),
            _first_present(finding, ["pkg", "PkgName", "package", "package_name"], ""),
            _first_present(finding, ["installed", "InstalledVersion", "installed_version"], ""),
            _first_present(finding, ["fixed", "FixedVersion", "fixed_version"], ""),
            _first_present(finding, ["title", "Title", "description", "Description"], ""),
        ])
    _add_table(doc, ["Severity", "CVE / ID", "Package", "Installed", "Fixed", "Title / description"], rows, max_rows=12)


def _mobsf_signal_rows(mobsf: Dict[str, Any]) -> List[List[Any]]:
    checks = [
        ("Debuggable", ["debuggable", "has_debuggable", "debuggable_true", "android_debuggable"]),
        ("Debug certificate", ["debug_certificate", "has_debug_cert", "signed_with_debug_certificate"]),
        ("v1 signature / Janus exposure", ["v1_signature", "janus", "has_janus", "v1_signature_present"]),
        ("SHA1 certificate/signature evidence", ["sha1", "uses_sha1", "has_sha1"]),
        ("Backup enabled", ["allow_backup", "backup_enabled", "android_allow_backup"]),
        ("Minimum SDK", ["min_sdk", "minsdk", "minSdk"]),
        ("Dangerous permissions", ["dangerous_permissions_count", "dangerous_permission_count"]),
        ("Exported components", ["exported_components_count", "exported_component_count"]),
        ("Trackers detected", ["trackers_detected", "tracker_count", "detected_trackers"]),
    ]
    rows = []
    for label, keys in checks:
        value = _deep_find_first(mobsf, keys)
        rows.append([label, _format_bool(value)])
    return rows


def _add_mobsf_static_section(doc: Document, technical: Dict[str, Any]) -> None:
    mobsf = _as_dict(technical.get("mobsf_static"))
    if not _block_available(mobsf):
        doc.add_paragraph("MobSF static evidence was not available in the analysis pack.")
        return
    doc.add_paragraph("MobSF static evidence was used to summarize Android APK, manifest, certificate, signing, permissions, tracker, and binary hardening indicators. Findings are treated as technical evidence supporting the workbook-level conclusions, not as a replacement for requirement-level adjudication.")
    _add_table(doc, ["Indicator", "Reported value"], _mobsf_signal_rows(mobsf), max_rows=20)

    finding_lists = []
    for key in ["manifest_findings", "certificate_findings", "findings", "high_findings", "warnings"]:
        items = _deep_list(mobsf, [key])
        if items:
            finding_lists.extend(items)
    rows = []
    for item in finding_lists[:10]:
        if isinstance(item, dict):
            rows.append([
                _first_present(item, ["severity", "level", "risk"], ""),
                _first_present(item, ["title", "name", "rule", "id"], ""),
                _first_present(item, ["description", "message", "summary"], ""),
            ])
        else:
            rows.append(["", "", item])
    if rows:
        _add_table(doc, ["Severity", "Finding", "Description"], rows, max_rows=10)


def _add_mobsf_dynamic_section(doc: Document, technical: Dict[str, Any]) -> None:
    dynamic = _as_dict(technical.get("mobsf_dynamic"))
    if not _block_available(dynamic):
        doc.add_paragraph("MobSF dynamic evidence was not available in the analysis pack. If dynamic analysis is not executed for a given application, runtime storage and behavioral observations should be treated as not assessed rather than clean.")
        return
    doc.add_paragraph("MobSF dynamic evidence was used to summarize runtime storage and behavioral observations, including local files, SharedPreferences, SQLite databases, cache artifacts, and trackers where reported.")
    rows = []
    for label, keys in [
        ("SharedPreferences artifacts", ["shared_preferences", "shared_preferences_files", "preferences"]),
        ("SQLite/database artifacts", ["sqlite_databases", "databases", "db_files"]),
        ("Local storage artifacts", ["local_storage_artifacts", "files", "storage_artifacts"]),
        ("Trackers", ["trackers", "detected_trackers"]),
    ]:
        items = _deep_list(dynamic, keys)
        count = len(items) if items else _deep_int(dynamic, [keys[0] + "_count"], 0)
        sample = ", ".join(_short(x, 80) for x in items[:4]) if items else ""
        rows.append([label, count, sample])
    _add_table(doc, ["Runtime evidence type", "Count", "Examples"], rows, max_rows=20)


def _sast_findings(sast: Dict[str, Any]) -> List[Dict[str, Any]]:
    for key in ["findings", "results", "app_code_findings", "top_findings"]:
        items = _deep_list(sast, [key])
        if items:
            return [x for x in items if isinstance(x, dict)]
    return []


def _add_sast_section(doc: Document, technical: Dict[str, Any]) -> None:
    sast = _as_dict(technical.get("sast_app_code"))
    if not _block_available(sast):
        doc.add_paragraph("SAST app-code evidence was not available in the analysis pack.")
        return
    total = _deep_int(sast, ["total_findings", "total", "results_count"])
    tool_counts = _deep_dict(sast, ["tool_counts", "by_tool"])
    scope_note = _first_present(sast, ["scope_note", "filter_note"], "CI/CD workflows and audit-generation scripts are excluded; retained findings are limited to application code and Android app artifacts according to the generic app-code scope filter.")
    doc.add_paragraph(f"SAST retained {total} application-code finding(s) after scope filtering. {scope_note}")
    if tool_counts:
        _add_table(doc, ["Tool", "Findings"], [[k, v] for k, v in tool_counts.items()], max_rows=10)
    rows = []
    for item in _sast_findings(sast)[:12]:
        rows.append([
            _first_present(item, ["tool", "driver", "source"], ""),
            _first_present(item, ["rule_id", "ruleId", "rule", "id"], ""),
            _first_present(item, ["severity", "level", "kind"], ""),
            _first_present(item, ["file", "path", "uri"], ""),
            _first_present(item, ["line", "start_line", "startLine"], ""),
            _first_present(item, ["message", "title", "description"], ""),
        ])
    _add_table(doc, ["Tool", "Rule", "Severity", "File", "Line", "Message"], rows, max_rows=12)


def _add_coverage_limitations(doc: Document, technical: Dict[str, Any]) -> None:
    limitations = _as_dict(technical.get("coverage_limitations"))
    rows = []
    if isinstance(limitations, dict):
        for key, value in limitations.items():
            if isinstance(value, list):
                rows.append([key, "; ".join(_short(x, 120) for x in value[:6])])
            elif isinstance(value, dict):
                rows.append([key, json.dumps(value, ensure_ascii=False)[:500]])
            elif value not in (None, ""):
                rows.append([key, value])
    if rows:
        _add_table(doc, ["Limitation", "Details"], rows, max_rows=12)
    else:
        doc.add_paragraph("No additional technical coverage limitations were reported in the analysis pack beyond the workbook-defined scope and tool-specific execution constraints.")


def _technical_kpi_for_pattern(pattern: str, technical: Dict[str, Any]) -> str:
    pat = pattern.lower()
    if "supply chain" in pat or "outdated" in pat or "dependency" in pat:
        return "No Critical or High dependency vulnerabilities remain in Trivy; all fixable CVEs have an upgrade, mitigation, or formally accepted-risk decision; dependency inventory/SBOM evidence is retained for the assessed release."
    if "storage" in pat or "key management" in pat:
        return "Local SharedPreferences, SQLite databases, caches, and files are reviewed for PHI, credentials, tokens, and keys; sensitive local data is encrypted or eliminated; backup exposure is explicitly controlled."
    if "tamper" in pat or "reverse" in pat or "binary" in pat:
        return "Release APK is not debuggable, is not signed with a debug certificate, debug-only tooling is absent, and release hardening controls are verified by MobSF or equivalent evidence."
    if "transport" in pat or "certificate" in pat or "tls" in pat:
        return "TLS configuration, certificate validation, and pinning decisions are documented and verified; weak signing or certificate indicators are remediated or formally justified."
    if "input" in pat or "injection" in pat:
        return "Application-code SAST findings for input handling and injection are triaged; exploitable issues are remediated and covered by regression tests."
    if "privacy" in pat or "permission" in pat:
        return "Dangerous permissions and privacy-sensitive data flows are justified, minimized, and covered by user-facing notices and runtime access controls."
    if "misconfiguration" in pat or "default" in pat:
        return "Manifest and runtime configuration findings are triaged; insecure defaults are removed or documented with compensating controls."
    return "Evidence recorded in workbook and technical artifacts; mapped controls can be re-tested and re-scored as compliant after remediation."

def _extract_json_object(text: str) -> str:
    text = (text or "").strip()
    if not text:
        raise ValueError("Model returned empty text.")

    # Direct JSON object.
    try:
        obj = json.loads(text)
        if isinstance(obj, dict):
            return json.dumps(obj, ensure_ascii=False)
    except Exception:
        pass

    # JSON inside Markdown code fence.
    fenced = re.search(
        r"```(?:json)?\s*(\{.*?\})\s*```",
        text,
        flags=re.DOTALL | re.IGNORECASE,
    )
    if fenced:
        return fenced.group(1)

    # Remove common fence wrappers if the model used an unterminated fence.
    cleaned = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, dict):
            return json.dumps(obj, ensure_ascii=False)
    except Exception:
        pass

    # Find the first balanced JSON object anywhere in the response.
    start = text.find("{")
    if start < 0:
        preview = text[:500].replace("\n", "\\n")
        raise ValueError(f"No JSON object found in model output. Preview: {preview}")

    depth = 0
    in_string = False
    escape = False

    for i in range(start, len(text)):
        ch = text[i]

        if escape:
            escape = False
            continue

        if ch == "\\":
            escape = True
            continue

        if ch == '"':
            in_string = not in_string
            continue

        if in_string:
            continue

        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                candidate = text[start:i + 1]
                obj = json.loads(candidate)
                if isinstance(obj, dict):
                    return candidate

    preview = text[:500].replace("\n", "\\n")
    raise ValueError(f"No complete JSON object found in model output. Preview: {preview}")


def _call_llm_for_style(patterns: List[Dict[str, Any]], likelihood_rubric: Dict[str, str], max_takeaways: int = 7) -> Dict[str, Any]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key or AIRuntime is None:
        return {}

    model = os.getenv("OPENAI_MODEL", "gpt-5.4").strip() or "gpt-5.4"
    max_tokens = int(os.getenv("OPENAI_MAX_OUTPUT_TOKENS", "6000"))
    effort = os.getenv("OPENAI_REASONING_EFFORT", "medium").strip() or "medium"
    runtime = AIRuntime(task=os.getenv("AI_TASK", "").strip() or "audit_summary_docx", api_key=api_key)

    inp = []
    for p in patterns[:10]:
        cnt = int(p.get("mapped_noncompliant_count", 0))
        inp.append({
            "pattern": p.get("pattern", ""),
            "count": cnt,
            "severity": p.get("severity", "Low"),
            "likelihood": _likelihood_from_count(cnt),
            "owner": p.get("recommended_owner", "Engineering"),
            "anchors": (p.get("description_anchors", []) or [])[:2],
        })

    system = (
        "You are a senior security audit reporting specialist for peer-review publications. "
        "You will improve wording and generate actionable recommendations ONLY at the weakness-pattern level. "
        "You must NOT invent specific implemented controls. You must NOT invent metrics. "
        "You must NOT claim facts beyond the provided anchors and counts. "
        "Return exactly one raw JSON object and nothing else. "
        "Do not use Markdown, code fences, prose outside JSON, comments, or tool calls."
    )
    user_payload = {
        "task": "Generate paper-quality prose components grounded in workbook-derived prevalence counts.",
        "constraints": {
            "no_invented_metrics": True,
            "no_category_level_bullet_dumps": True,
            "no_long_id_lists": True,
            "recommendations_no_time_headings": True,
            "max_key_takeaways": max_takeaways,
            "likelihood_rubric": likelihood_rubric,
        },
        "input_patterns": inp,
        "required_output_schema": {
            "key_takeaways": ["<5-7 bullets; each references prevalence count and pattern name>"],
            "pattern_writeups": [{"pattern": "<exact pattern name from input>", "expected": "<1-2 sentences>", "impact": "<1-2 sentences; CIA + regulatory for health data>", "recommendations": ["<6-10 bullets; practical; may include MFA/biometric step-up if appropriate>"]}]
        }
    }

    resp = runtime.create(
        model=model,
        input=[{"role": "system", "content": system}, {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}],
        max_output_tokens=max_tokens,
        reasoning={"effort": effort},
    )
    obj = json.loads(_extract_json_object((getattr(resp, "output_text", "") or "").strip()))
    return obj if isinstance(obj, dict) else {}


def main() -> None:
    in_path = os.getenv("AUDIT_ANALYSIS_JSON_PATH", DEFAULT_IN)
    out_path = os.getenv("AUDIT_SUMMARY_DOCX_PATH", DEFAULT_OUT)
    chart_dir = os.getenv("AUDIT_SUMMARY_CHART_DIR", CHART_DIR)

    if not os.path.isfile(in_path):
        raise SystemExit(f"[ERROR] analysis pack not found: {in_path}")

    Path(chart_dir).mkdir(parents=True, exist_ok=True)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    with open(in_path, "r", encoding="utf-8") as f:
        pack = json.load(f)

    metrics = pack["metrics"]
    cat_stats = pack["category_metrics"]
    app = pack["app_metadata"]
    actors = pack["actors"]
    patterns = pack["weakness_patterns"]
    pos_controls = pack.get("positive_controls_candidates", [])[:7]
    likelihood_rubric = pack.get("likelihood_rubric", {})
    technical = _technical_evidence(pack)
    report_title = _app_report_title(app)

    audit_dt = datetime.utcnow().date()
    audit_date_str = audit_dt.strftime("%d %b %Y")

    plt.rcParams.update({"font.size": 10, "figure.titlesize": 12, "axes.titlesize": 12, "axes.labelsize": 10})

    fig1 = os.path.join(chart_dir, "figure1_overall_donut.png")
    fig2 = os.path.join(chart_dir, "figure2_noncompliance_share_hbar.png")
    fig3 = os.path.join(chart_dir, "figure3_compliance_rate_hbar.png")
    fig4 = os.path.join(chart_dir, "figure4_counts_stacked_hbar.png")

    applicable = int(metrics["applicable"])
    compliant = int(metrics["compliant"])
    non_compliant = int(metrics["non_compliant"])
    not_applicable = int(metrics["not_applicable"])
    overall_pct = float(metrics["overall_compliance_pct"])

    _donut([compliant, non_compliant, not_applicable], ["Compliant", "Non-compliant", "Not applicable"], "Overall compliance distribution (workbook-derived)", f"{overall_pct:.2f}%\ncompliant\n(applicable)", fig1)
    _hbar_share_noncompliances(cat_stats, fig2)
    _hbar_compliance_rate(cat_stats, fig3)
    _stacked_counts(cat_stats, fig4)

    prose: Dict[str, Any] = {}
    if AIRuntime is not None and os.getenv("OPENAI_API_KEY", "").strip():
        try:
            prose = _call_llm_for_style(patterns, likelihood_rubric, max_takeaways=7)
        except Exception:
            prose = {}

    key_takeaways = prose.get("key_takeaways", [])
    writeups = {w["pattern"]: w for w in prose.get("pattern_writeups", []) if isinstance(w, dict) and "pattern" in w}
    if not key_takeaways:
        key_takeaways = [f"{p['pattern']} - {p['severity']} severity; {int(p['mapped_noncompliant_count'])} related non-compliant control(s) in the workbook." for p in patterns[:7]]
    technical_takeaways = _technical_takeaways(technical)
    if technical_takeaways:
        key_takeaways = (technical_takeaways + key_takeaways)[:7]

    doc = Document()
    _set_doc_defaults(doc)
    _add_header_footer(doc.sections[0], audit_date_str, report_title)
    _add_cover(doc, audit_date_str, actors["Auditor"], report_title)
    used_bookmarks: set[str] = set()
    toc_entries: List[tuple[int, str, str]] = []
    bookmark_id = 1000

    def add_nav_heading(text: str, level: int):
        nonlocal bookmark_id
        style = f"Heading {level}"
        p = doc.add_paragraph(text, style=style)
        anchor = _make_bookmark_name(text, used_bookmarks)
        _add_bookmark(p, anchor, bookmark_id)
        toc_entries.append((level, text, anchor))
        bookmark_id += 1
        return p

    toc_placeholder = _add_toc(doc)

    add_nav_heading("1. App information", 1)
    _add_two_col_table(doc, [[k, v] for k, v in app.items()])

    add_nav_heading("2. Actors", 1)
    _add_two_col_table(doc, [["Auditor", actors["Auditor"]], ["Requirement Engineering team", "\n".join(actors["Requirement Engineering team"])], ["Engineering Group (EN)", "\n".join(actors["Engineering Group (EN)"])]])

    add_nav_heading("3. Scope and limitations", 1)
    doc.add_paragraph("This Audit Summary consolidates the compliance determinations recorded in the audit workbook. The results reflect the assessed application version and the workbook-defined scope. Controls not evidenced as implemented in the workbook are reported as non-compliant for summary purposes. Where available, technical evidence from Vision360, Trivy, MobSF, and app-code SAST is used to support and qualify the workbook-derived findings. CI/CD workflows and audit-generation scripts are outside the application audit scope unless explicitly included by project configuration.")

    add_nav_heading("4. Evidence criteria", 1)
    doc.add_paragraph("- Compliant: the workbook provides sufficient evidence that the control is implemented and effective for the assessed scope.\n- Non-compliant: the workbook indicates the control is missing, insufficient, or not evidenced.\n- Not applicable: the control is recorded as out of scope or not relevant for the assessed context.\n- Technical evidence: automated scan artifacts are used as supporting evidence for application code, Android manifest, binary, dependency, runtime, and storage observations. They do not replace requirement-level judgment and are interpreted within the declared audit scope.")

    add_nav_heading("5. Audit summary", 1)
    doc.add_paragraph("The audit was carried out using the mSEC-AM (mobile SECurity Audit Method).")
    doc.add_paragraph(f"Overall, {int(metrics['total_assessed'])} requirements were assessed. {applicable} were applicable controls and {not_applicable} were recorded as not applicable. Of the applicable controls, {compliant} were compliant and {non_compliant} were non-compliant, resulting in an overall compliance rate of {overall_pct:.2f}% (applicable controls only).")
    doc.add_paragraph("This report summarizes the dominant weakness patterns evidenced by non-compliant requirements and proposes actionable remediations suitable for mHealth/EMR environments handling sensitive health information.")
    if technical:
        doc.add_paragraph("The report also incorporates technical scan evidence where available, including dependency vulnerability evidence from Trivy, Android APK and manifest evidence from MobSF, runtime observations from MobSF dynamic analysis, and SAST findings filtered to application code.")

    add_nav_heading("5.1 Key takeaways (Top findings)", 2)
    _add_callout(doc, "Key takeaways (Top findings)", key_takeaways[:7])

    add_nav_heading("5.2 Positive controls observed", 2)
    doc.add_paragraph("All statements below are derived exclusively from controls recorded as Compliant in the audit workbook and include supporting signals (flags and/or evidence). Verification traceability is provided in Appendix B.")
    if pos_controls:
        for pc in pos_controls:
            doc.add_paragraph(pc["declarative_statement"], style="List Bullet")
    else:
        doc.add_paragraph("No compliant controls with supporting evidence/flags were available for verification in the workbook.", style="List Bullet")

    add_nav_heading("5.3 Risk scoring approach", 2)
    doc.add_paragraph("Severity and likelihood ratings in this report follow a qualitative rubric grounded in the audit workbook:\n- Severity reflects potential impact on confidentiality, integrity, and availability of health information, including regulatory exposure.\n- Likelihood is derived from workbook prevalence: the count of non-compliant controls mapped to a weakness pattern as a proxy for exposure.\nLikelihood mapping: High (>=50), Medium-High (20-49), Medium (10-19), Low-Medium (<10).")

    add_nav_heading("5.4 Risk triage (prioritized)", 2)
    rt = doc.add_table(rows=1, cols=7)
    rt.style = "Table Grid"
    h = rt.rows[0].cells
    for idx, txt in enumerate(["Weakness pattern", "Severity (rubric)", "Impact", "Likelihood (workbook prevalence)", "Workbook basis", "Recommended owner", "Target timeline"]):
        h[idx].text = txt
        _set_cell_shading(h[idx], "D9E1F2")
        for run in h[idx].paragraphs[0].runs:
            run.bold = True
    for p in patterns[:10]:
        cnt = int(p["mapped_noncompliant_count"])
        lik = _likelihood_from_count(cnt)
        sev = p["severity"]
        owner = p["recommended_owner"]
        impact = writeups.get(p["pattern"], {}).get("impact", "The weakness pattern can compromise confidentiality/integrity/availability of health information and increase regulatory exposure.")
        row = rt.add_row().cells
        row[0].text = p["pattern"]
        row[1].text = sev
        row[2].text = impact
        row[3].text = lik
        row[4].text = f"{cnt} mapped non-compliant control(s) in the workbook."
        row[5].text = owner
        row[6].text = _target_timeline(sev)
    doc.add_paragraph()

    add_nav_heading("5.5 Technical scan coverage", 2)
    doc.add_paragraph("The following table summarizes which automated technical evidence sources were available to support the audit summary. Absence of a technical source means that the source was not available to the report generator, not necessarily that the corresponding risk is absent.")
    _add_table(doc, ["Evidence source", "Available", "Summary"], _source_status_rows(technical), max_rows=10)

    add_nav_heading("6. Technical evidence from automated analysis", 1)
    doc.add_paragraph("This section summarizes technical scan evidence relevant to the assessed application and its code. The evidence is used to reinforce and qualify workbook findings while preserving the workbook as the authoritative requirement-level adjudication source.")
    add_nav_heading("6.1 Software Composition Analysis from Trivy", 2)
    _add_trivy_section(doc, technical)
    add_nav_heading("6.2 Android static evidence from MobSF", 2)
    _add_mobsf_static_section(doc, technical)
    add_nav_heading("6.3 Runtime evidence from MobSF dynamic analysis", 2)
    _add_mobsf_dynamic_section(doc, technical)
    add_nav_heading("6.4 Static Application Security Testing evidence", 2)
    _add_sast_section(doc, technical)
    add_nav_heading("6.5 Technical coverage limitations", 2)
    _add_coverage_limitations(doc, technical)

    add_nav_heading("7. Main deficiencies", 1)
    doc.add_paragraph("The following deficiencies are synthesized as common weakness patterns based on non-compliant requirements and supported, where available, by technical scan evidence. They are not grouped by category; instead they represent cross-cutting gaps evidenced in the audit workbook and related artifacts.")
    for p in patterns[:10]:
        pat = p["pattern"]
        cnt = int(p["mapped_noncompliant_count"])
        sev = p["severity"]
        owner = p["recommended_owner"]
        ex_ids = p.get("example_puids", [])[:4]
        anchors = p.get("description_anchors", [])[:2]
        doc.add_paragraph(f"{pat} ({sev})", style="Heading 2")
        doc.add_paragraph(f"Workbook basis: {cnt} related non-compliant control(s) mapped to this pattern.")
        expected = writeups.get(pat, {}).get("expected", "Controls in this area should provide robust, consistently enforced safeguards appropriate to health data processing.")
        impact = writeups.get(pat, {}).get("impact", "Deficiencies can increase the likelihood and impact of security incidents affecting confidentiality, integrity, or availability.")
        doc.add_paragraph(f"Expected: {expected}")
        doc.add_paragraph("Observed: The audit workbook indicates the related controls are missing, insufficient, or not evidenced for the assessed scope.")
        doc.add_paragraph(f"Impact: {impact}")
        doc.add_paragraph(f"Recommended owner: {owner}")
        if ex_ids:
            doc.add_paragraph(f"Traceability (examples, non-exhaustive): {', '.join(ex_ids)}.")
        for a in anchors:
            doc.add_paragraph(f"Evidence anchor (from workbook description): {a}", style="List Bullet")

    doc.add_page_break()
    add_nav_heading("8. Recommendations", 1)
    doc.add_paragraph("Recommendations are organized by the same weakness patterns presented in the Main deficiencies section. They target remediation of workbook-evidenced gaps and may include strengthening controls to improve security posture.")
    for p in patterns[:10]:
        pat = p["pattern"]
        doc.add_paragraph(pat, style="Heading 2")
        recs = writeups.get(pat, {}).get("recommendations", [])
        if not recs:
            recs = [
                "Define and document secure-by-default requirements for this control area, and implement automated tests to prevent regressions.",
                "Apply least-privilege, defense-in-depth, and secure configuration baselines aligned with mHealth/EMR risk profiles.",
                "Introduce step-up authentication, such as MFA, TOTP, FIDO2, or Android BiometricPrompt, for sensitive actions where feasible.",
                "Validate effectiveness through security testing and re-assessment of the mapped non-compliant controls.",
            ]
        for r in recs[:12]:
            doc.add_paragraph(str(r).strip(), style="List Bullet")

    doc.add_page_break()
    add_nav_heading("9. Visual Analytics", 1)
    doc.add_paragraph("Figures below summarize workbook-derived outcomes and distributions. All figures: source: audit workbook.")
    _add_figure(doc, fig1, "Figure 1. Overall compliance distribution (donut chart; source: audit workbook).")
    _add_figure(doc, fig2, "Figure 2. Share of non-compliances by category (legible horizontal bars; source: audit workbook).")
    _add_figure(doc, fig3, "Figure 3. Compliance rate by category (applicable controls only; source: audit workbook).")
    _add_figure(doc, fig4, "Figure 4. Counts by category and status (horizontal stacked bars; source: audit workbook).")

    doc.add_page_break()
    add_nav_heading("10. Management Action Plan (MAP)", 1)
    doc.add_paragraph("Severity and likelihood nomenclature follow the rubric in Section 5.3. Likelihood is supported by workbook prevalence counts recorded in the Workbook basis column.")
    mp = doc.add_table(rows=1, cols=9)
    mp.style = "Table Grid"
    mh = mp.rows[0].cells
    headers = ["Finding / weakness pattern", "Severity", "Likelihood", "Workbook basis", "Owner", "Management action", "Target window", "Target date", "Acceptance criteria / KPI"]
    for idx, txt in enumerate(headers):
        mh[idx].text = txt
        _set_cell_shading(mh[idx], "D9E1F2")
        for run in mh[idx].paragraphs[0].runs:
            run.bold = True
    for p in patterns[:10]:
        pat = p["pattern"]
        cnt = int(p["mapped_noncompliant_count"])
        sev = p["severity"]
        lik = _likelihood_from_count(cnt)
        owner = p["recommended_owner"]
        target_window = _target_timeline(sev)
        target_date = _target_date_str(audit_dt, sev)
        recs = writeups.get(pat, {}).get("recommendations", [])
        action = " ".join([r.strip() for r in recs[:3]]) if recs else "Implement remediation actions aligned to the weakness pattern and validate effectiveness."
        row = mp.add_row().cells
        row[0].text = pat
        row[1].text = sev
        row[2].text = lik
        row[3].text = f"{cnt} mapped non-compliant control(s) in workbook."
        row[4].text = owner
        row[5].text = action
        row[6].text = target_window
        row[7].text = target_date
        row[8].text = _technical_kpi_for_pattern(pat, technical)

    doc.add_page_break()
    add_nav_heading("Appendix A - Traceability index (non-exhaustive)", 1)
    doc.add_paragraph("For complete traceability and evidence, refer to the audit workbook.")
    for p in patterns[:10]:
        ex = p.get("example_puids", [])[:5]
        if ex:
            doc.add_paragraph(f"{p['pattern']}: {', '.join(ex)}", style="List Bullet")

    doc.add_page_break()
    add_nav_heading("Appendix B - Positive controls verification (workbook traceability)", 1)
    doc.add_paragraph("This appendix verifies each Positive controls observed statement by providing the originating PUID, flags used, and an evidence excerpt when available.")
    vb = doc.add_table(rows=1, cols=4)
    vb.style = "Table Grid"
    vh = vb.rows[0].cells
    for idx, txt in enumerate(["Positive control statement (as reported)", "Workbook PUID", "Flags used", "Evidence / justification (excerpt)"]):
        vh[idx].text = txt
        _set_cell_shading(vh[idx], "D9E1F2")
        for run in vh[idx].paragraphs[0].runs:
            run.bold = True
    if pos_controls:
        for pc in pos_controls:
            r = vb.add_row().cells
            r[0].text = pc["declarative_statement"]
            r[1].text = pc["puid"]
            r[2].text = pc.get("flags_used", "") or ""
            r[3].text = pc.get("evidence_excerpt", "") or ""
    else:
        r = vb.add_row().cells
        r[0].text = "No verified positive controls available."
        r[1].text = ""
        r[2].text = ""
        r[3].text = ""

    doc.add_page_break()
    add_nav_heading("Appendix C - Technical evidence summary", 1)
    doc.add_paragraph("This appendix provides a compact index of the technical evidence parsed from scan artifacts. Detailed raw JSON, SARIF, and tool reports remain in their original pipeline artifacts and are not embedded in the report package.")
    _add_table(doc, ["Evidence source", "Available", "Summary"], _source_status_rows(technical), max_rows=10)

    _render_clickable_toc(toc_placeholder, toc_entries)
    _enable_update_fields_on_open(doc)
    doc.save(out_path)
    print(f"[OK] DOCX generated -> {out_path}")


if __name__ == "__main__":
    main()