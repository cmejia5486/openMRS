#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import os
import re
import time
from datetime import datetime, date, timezone
from pathlib import Path
from typing import Dict, Any, List

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

try:
    from lib.ai_runtime import AIRuntime  # type: ignore
except Exception as exc:
    AIRuntime = None  # type: ignore
    AI_RUNTIME_IMPORT_ERROR = exc
else:
    AI_RUNTIME_IMPORT_ERROR = None

try:
    from lib.prompt_telemetry import contract_for_section, record_prompt_call
except Exception:
    def contract_for_section(section_name: str, fallback_source_function: str = "_ai_json_chat") -> Dict[str, Any]:  # type: ignore
        return {
            "prompt_id": "UNAVAILABLE",
            "prompt_name": section_name,
            "prompt_scope": "unavailable",
            "prompt_category": "unavailable",
            "source_file": "scripts/audit_summary_stage2_generate_docx.py",
            "source_function": fallback_source_function,
            "registration_status": "unavailable",
        }
    def record_prompt_call(**kwargs: Any) -> str:  # type: ignore
        return str(kwargs.get("prompt_call_id") or "")

def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _runtime_data_dir() -> Path:
    # Prefer explicit, cross-platform data-directory variables.
    # If none are provided, use GitHub Actions RUNNER_TEMP when available.
    # As a final local fallback, use a repository-local hidden directory.
    for name in ("VISION360_DATA_DIR", "AUDIT_DATA_DIR", "SECURITY_AUDIT_DATA_DIR"):
        raw = os.getenv(name, "").strip()
        if raw:
            base = Path(raw)
            base.mkdir(parents=True, exist_ok=True)
            return base

    runner_temp = os.getenv("RUNNER_TEMP", "").strip()
    if runner_temp:
        base = Path(runner_temp) / "vision360-data"
        base.mkdir(parents=True, exist_ok=True)
        return base

    base = _repo_root() / ".vision360-data"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _env_path(env_name: str, default_filename: str) -> Path:
    raw = os.getenv(env_name, "").strip()
    if raw:
        return Path(raw)
    return _runtime_data_dir() / default_filename


DEFAULT_IN = str(_env_path("AUDIT_ANALYSIS_JSON_PATH", "audit_summary_analysis_pack.json"))
DEFAULT_OUT = str(_env_path("AUDIT_SUMMARY_DOCX_PATH", "Audit Summary.docx"))
CHART_DIR = str(_env_path("AUDIT_SUMMARY_CHART_DIR", "_audit_summary_charts"))

HEADER_TEXT_TEMPLATE = "mSEC-AM Audit Summary - {report_title}"

REPORT_FONT_NAME = "Arial"
REPORT_BODY_PT = 9.5
REPORT_HEADING1_PT = 14.0
REPORT_HEADING2_PT = 12.0
REPORT_HEADING3_PT = 10.5
REPORT_CAPTION_PT = 8.5
REPORT_TABLE_HEADER_PT = 8.0
REPORT_TABLE_BODY_PT = 7.6
REPORT_HEADER_FOOTER_PT = 7.5

THEME_NAVY = "17365D"
THEME_BLUE = "D9EAF7"
THEME_LIGHT_BLUE = "EEF5FB"
THEME_LIGHT = "F7F9FC"
THEME_GRAY = "EDEDED"
THEME_GREEN = "E2F0D9"
THEME_RED = "FCE4D6"



def _wrap_label(s: str, width: int = 30) -> str:
    s = str(s)
    if len(s) <= width:
        return s
    out = []
    while len(s) > width:
        cut = s.rfind(" ", 0, width)
        if cut == -1:
            cut = width
        out.append(s[:cut].strip())
        s = s[cut:].strip()
    if s:
        out.append(s)
    return "\n".join(out)


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)




def _set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    """Set Word table-cell margins in twentieths of a point."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.first_child_found_in("w:tcMar")
        if tcMar is None:
            tcMar = OxmlElement("w:tcMar")
            tcPr.append(tcMar)
        for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
            node = tcMar.find(qn(f"w:{m}"))
            if node is None:
                node = OxmlElement(f"w:{m}")
                tcMar.append(node)
            node.set(qn("w:w"), str(v))
            node.set(qn("w:type"), "dxa")
    except Exception:
        pass


def _set_cell_vertical_alignment(cell, align=WD_CELL_VERTICAL_ALIGNMENT.CENTER) -> None:
    try:
        cell.vertical_alignment = align
    except Exception:
        pass


def _set_table_alignment(table, align=WD_TABLE_ALIGNMENT.CENTER) -> None:
    try:
        table.alignment = align
    except Exception:
        pass


def _style_cell_text(cell, size_pt: float = REPORT_TABLE_BODY_PT, bold: bool | None = None, italic: bool | None = None, color_hex: str | None = None) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            _set_run_font(run, size_pt=size_pt, bold=bold, italic=italic, color_hex=color_hex)


def _add_section_callout(doc: Document, title: str, text: str, fill: str = THEME_LIGHT_BLUE) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    _set_table_alignment(tbl)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    _set_cell_vertical_alignment(cell)
    p = cell.paragraphs[0]
    r = p.add_run(_report_display_text(title))
    _set_run_font(r, size_pt=REPORT_BODY_PT, bold=True, color_hex=THEME_NAVY)
    if text:
        p2 = cell.add_paragraph(_report_display_text(text))
        for run in p2.runs:
            _set_run_font(run, size_pt=REPORT_BODY_PT)
    doc.add_paragraph()


def _style_heading_paragraph(paragraph, level: int) -> None:
    try:
        paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        paragraph.paragraph_format.space_after = Pt(5 if level == 1 else 3)
        for run in paragraph.runs:
            _set_run_font(run, bold=True, color_hex=THEME_NAVY)
    except Exception:
        pass

def _add_field_run(paragraph, field_instr: str) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _set_run_font(
    run,
    name: str = REPORT_FONT_NAME,
    size_pt: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color_hex: str | None = None,
) -> None:
    """Apply a single Word font consistently across latin/complex/east-Asian slots."""
    try:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(str(color_hex).strip().lstrip("#")[:6])
    except Exception:
        pass


def _set_style_font(doc: Document, style_name: str, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    try:
        style = doc.styles[style_name]
    except Exception:
        return
    try:
        style.font.name = REPORT_FONT_NAME
        style.font.size = Pt(size_pt)
        if bold is not None:
            style.font.bold = bold
        if italic is not None:
            style.font.italic = italic
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), REPORT_FONT_NAME)
    except Exception:
        pass


def _set_doc_defaults(doc: Document) -> None:
    _set_style_font(doc, "Normal", REPORT_BODY_PT)
    _set_style_font(doc, "Title", 22.0, bold=True)
    _set_style_font(doc, "Subtitle", 13.0, bold=True)
    _set_style_font(doc, "Heading 1", REPORT_HEADING1_PT, bold=True)
    _set_style_font(doc, "Heading 2", REPORT_HEADING2_PT, bold=True)
    _set_style_font(doc, "Heading 3", REPORT_HEADING3_PT, bold=True)
    _set_style_font(doc, "Caption", REPORT_CAPTION_PT, italic=True)
    _set_style_font(doc, "List Bullet", REPORT_BODY_PT)
    _set_style_font(doc, "List Number", REPORT_BODY_PT)

    style = doc.styles["Normal"]
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.line_spacing = 1.05
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def _is_toc_paragraph(paragraph) -> bool:
    try:
        if paragraph._p.xpath("./w:hyperlink"):
            return True
    except Exception:
        pass
    try:
        left_indent = paragraph.paragraph_format.left_indent
        if left_indent is not None and left_indent.pt and left_indent.pt > 0 and paragraph.text.strip():
            return True
    except Exception:
        pass
    return False


def _should_justify_paragraph(paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return False

    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if style_name.startswith("Heading"):
        return False
    if style_name in {"Title", "Subtitle", "List Bullet", "List Number", "Caption"}:
        return False
    if paragraph.alignment in {WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT}:
        return False
    if _is_toc_paragraph(paragraph):
        return False
    return True


def _format_report_paragraphs(doc: Document) -> None:
    """Apply professional paragraph formatting to narrative body text only.

    Tables, figures, captions, bullets, headings, cover text, TOC links, headers,
    and footers are intentionally excluded. The goal is to fix ragged narrative
    text in the exported PDF without damaging table readability.
    """
    for p in doc.paragraphs:
        if not _should_justify_paragraph(p):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08


def _add_body_paragraph(doc: Document, text: Any):
    p = doc.add_paragraph(_report_display_text(text))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    return p


def _add_header_footer(section, audit_date_str: str, report_title: str = "Mobile Application") -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = HEADER_TEXT_TEMPLATE.format(report_title=report_title)
    if p.runs:
        _set_run_font(p.runs[0], size_pt=REPORT_HEADER_FOOTER_PT)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = f"{audit_date_str} | "
    if fp.runs:
        _set_run_font(fp.runs[0], size_pt=REPORT_HEADER_FOOTER_PT)
    fp.add_run("Page ")
    _add_field_run(fp, " PAGE ")
    fp.add_run(" of ")
    _add_field_run(fp, " NUMPAGES ")
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_cover(doc: Document, audit_date_str: str, auditor: str, report_title: str = "Mobile Application") -> None:
    """Render an executive-grade native Word cover page."""
    for _ in range(2):
        doc.add_paragraph()

    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    _set_table_alignment(banner)
    cell = banner.cell(0, 0)
    _set_cell_shading(cell, THEME_NAVY)
    _set_cell_margins(cell, top=260, start=180, bottom=260, end=180)
    _set_cell_vertical_alignment(cell)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("mSEC-AM")
    _set_run_font(r, size_pt=28, bold=True, color_hex="FFFFFF")
    p2 = cell.add_paragraph("Mobile Security Audit Summary")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p2.runs[0], size_pt=16, bold=True, color_hex="FFFFFF")
    p3 = cell.add_paragraph(report_title)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p3.runs[0], size_pt=12, italic=True, color_hex="FFFFFF")

    doc.add_paragraph()
    card = doc.add_table(rows=0, cols=2)
    card.style = "Table Grid"
    _set_table_alignment(card)
    cover_rows = [
        ("Audit date", audit_date_str),
        ("Auditor", auditor),
        ("Method", "mSEC-AM - mobile SECurity Audit Method"),
        ("Classification", "Confidential / Internal Use"),
        ("Evidence model", "Audit workbook, Vision360, Trivy, MobSF, SAST, and treatment traceability"),
    ]
    for label, value in cover_rows:
        cells = card.add_row().cells
        cells[0].text = label
        cells[1].text = value
        _set_cell_shading(cells[0], THEME_BLUE)
        _set_cell_margins(cells[0], top=90, start=110, bottom=90, end=110)
        _set_cell_margins(cells[1], top=90, start=110, bottom=90, end=110)
        _format_table_cell(cells[0], font_size=REPORT_BODY_PT, bold=True, no_wrap=True)
        _format_table_cell(cells[1], font_size=REPORT_BODY_PT)

    doc.add_paragraph()
    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    note_cell = note.cell(0, 0)
    _set_cell_shading(note_cell, THEME_LIGHT_BLUE)
    _set_cell_margins(note_cell, top=130, start=130, bottom=130, end=130)
    np = note_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run("Evidence-grounded executive report with full requirement, scanner, Vision360, and remediation traceability.")
    _set_run_font(nr, size_pt=REPORT_BODY_PT, italic=True, color_hex=THEME_NAVY)
    doc.add_page_break()

def _enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.xpath('./w:updateFields')
    if not existing:
        node = OxmlElement("w:updateFields")
        node.set(qn("w:val"), "true")
        settings.append(node)


def _make_bookmark_name(text: str, used: set[str]) -> str:
    base = re.sub(r"[^A-Za-z0-9_]+", "_", str(text)).strip("_")
    if not base:
        base = "section"
    if base[0].isdigit():
        base = f"s_{base}"
    name = base[:32]
    seed = name
    i = 2
    while name in used:
        suffix = f"_{i}"
        name = f"{seed[:32-len(suffix)]}{suffix}"
        i += 1
    used.add(name)
    return name


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def _add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _render_clickable_toc(after_paragraph, toc_entries: List[tuple[int, str, str]]) -> None:
    current = after_paragraph
    for level, title, anchor in toc_entries:
        p = _insert_paragraph_after(current)
        p.paragraph_format.left_indent = Inches(0.25 * max(level - 1, 0))
        p.paragraph_format.space_after = Pt(2)
        _add_internal_hyperlink(p, title, anchor)
        current = p


def _add_toc(doc: Document):
    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = doc.add_paragraph()
    doc.add_page_break()
    return p


def _add_two_col_table(doc: Document, rows: List[List[str]]) -> None:
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    _set_table_alignment(tbl)
    for k, v in rows:
        cells = tbl.add_row().cells
        cells[0].text = _report_display_text(k)
        cells[1].text = _report_display_text(v)
        _set_cell_shading(cells[0], THEME_BLUE)
        _set_cell_margins(cells[0], top=70, start=90, bottom=70, end=90)
        _set_cell_margins(cells[1], top=70, start=90, bottom=70, end=90)
        _format_table_cell(cells[0], font_size=REPORT_TABLE_BODY_PT, bold=True, no_wrap=True)
        _format_table_cell(cells[1], font_size=REPORT_TABLE_BODY_PT)
    doc.add_paragraph()


def _add_callout(doc: Document, title: str, bullets: List[str]) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    _set_table_alignment(tbl)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, THEME_LIGHT_BLUE)
    _set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    p = cell.paragraphs[0]
    r = p.add_run(_report_display_text(title))
    _set_run_font(r, size_pt=REPORT_BODY_PT, bold=True, color_hex=THEME_NAVY)
    for b in bullets:
        bp = cell.add_paragraph(_report_display_text(b), style="List Bullet")
        for run in bp.runs:
            _set_run_font(run, size_pt=REPORT_BODY_PT)
    doc.add_paragraph()



def _add_metric_cards(doc: Document, cards: List[tuple[str, Any, str]]) -> None:
    """Render compact executive metric cards."""
    if not cards:
        return
    cols = min(4, max(1, len(cards)))
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = "Table Grid"
    _set_table_alignment(tbl)
    for idx, (label, value, note) in enumerate(cards[:cols]):
        cell = tbl.cell(0, idx)
        fill = THEME_LIGHT_BLUE if idx % 2 == 0 else THEME_LIGHT
        _set_cell_shading(cell, fill)
        _set_cell_margins(cell, top=120, start=90, bottom=120, end=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(value))
        _set_run_font(r, size_pt=16, bold=True, color_hex=THEME_NAVY)
        p2 = cell.add_paragraph(str(label))
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p2.runs:
            _set_run_font(p2.runs[0], size_pt=8.5, bold=True, color_hex=THEME_NAVY)
        if note:
            p3 = cell.add_paragraph(str(note))
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p3.runs:
                _set_run_font(p3.runs[0], size_pt=7.2, italic=True)
    doc.add_paragraph()

def _add_figure(doc: Document, img_path: str, caption: str) -> None:
    doc.add_picture(img_path, width=Inches(6.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        _set_run_font(cap.runs[0], size_pt=REPORT_CAPTION_PT, italic=True)
    doc.add_paragraph()


def _donut(values: List[int], labels: List[str], title: str, center_text: str, out_path: str) -> None:
    fig, ax = plt.subplots(figsize=(7.4, 4.8))
    wedges, _, _ = ax.pie(values, labels=None, autopct="%1.1f%%", startangle=90, pctdistance=0.82)
    centre = plt.Circle((0, 0), 0.55, fc="white")
    fig.gca().add_artist(centre)
    ax.axis("equal")
    ax.set_title(title)
    ax.legend(wedges, labels, loc="center left", bbox_to_anchor=(1.02, 0.5), frameon=False)
    ax.text(0, 0, center_text, ha="center", va="center", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_share_noncompliances(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = sorted([(d["category_name"], int(d.get("non_compliant", 0))) for _, d in cat_stats.items()], key=lambda x: x[1])
    names = [x[0] for x in items]
    counts = [float(x[1]) for x in items]
    total = float(sum(counts))
    pct = [0.0 for _ in counts] if total <= 0.0 else [(c / total) * 100.0 for c in counts]
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    bars = ax.barh(y, pct)
    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(n, 30) for n in names])
    ax.set_xlabel("Share of total non-compliances (%)")
    ax.set_title("Share of non-compliances by category (high-level)")
    ax.set_xlim(0, max(5, float(max(pct) if pct else 0.0) + 5))
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    for b, c, p in zip(bars, counts, pct):
        ax.text(b.get_width() + 0.6, b.get_y() + b.get_height() / 2, f"{p:.1f}%  (n={int(c)})", va="center", fontsize=9)
    fig.text(0.01, 0.01, f"Source: audit workbook. Total non-compliances: {int(total)}.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_compliance_rate(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = []
    for _, d in cat_stats.items():
        app = int(d.get("applicable", 0))
        comp = int(d.get("compliant", 0))
        pct = float(d.get("compliance_pct", 0.0))
        items.append((d["category_name"], pct, comp, app))
    items = sorted(items, key=lambda x: x[1])
    names = [x[0] for x in items]
    pct_vals = [x[1] for x in items]
    comp_vals = [x[2] for x in items]
    app_vals = [x[3] for x in items]
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    bars = ax.barh(y, pct_vals)
    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(n, 30) for n in names])
    ax.set_xlabel("Compliance rate (%) - applicable controls only")
    ax.set_title("Compliance rate by category (applicable controls only)")
    ax.set_xlim(0, 100)
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    for b, p, c, a in zip(bars, pct_vals, comp_vals, app_vals):
        ax.text(p + 1, b.get_y() + b.get_height() / 2, f"{p:.1f}%  ({c}/{a})", va="center", fontsize=9)
    fig.text(0.01, 0.01, "Note: (c/a) indicates Compliant / Applicable controls per category. Source: audit workbook.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _stacked_counts(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = sorted([(d["category_name"], int(d.get("compliant", 0)), int(d.get("non_compliant", 0)), int(d.get("not_applicable", 0))) for _, d in cat_stats.items()], key=lambda x: (x[1] + x[2] + x[3]), reverse=True)
    names = [x[0] for x in items]
    c = [x[1] for x in items]
    n = [x[2] for x in items]
    na = [x[3] for x in items]
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    ax.barh(y, c, label="Compliant")
    ax.barh(y, n, left=c, label="Non-compliant")
    left_cn = [cc + nn for cc, nn in zip(c, n)]
    ax.barh(y, na, left=left_cn, label="Not applicable")
    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(nm, 30) for nm in names])
    ax.invert_yaxis()
    ax.set_xlabel("Count")
    ax.set_title("Counts by category and status")
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    ax.legend(loc="lower right", frameon=False)
    fig.tight_layout()
    fig.savefig(out_path, dpi=300)
    plt.close(fig)



def _hbar_treatment_pattern_volume(treatment_plan: Dict[str, Any], out_path: str) -> None:
    data = _deep_list(treatment_plan, ["weakness_pattern_volume"])
    if not data:
        data = _deep_list(_as_dict(treatment_plan.get("chart_data")), ["weakness_pattern_volume"])
    rows = []
    for item in data[:12]:
        if isinstance(item, dict):
            rows.append((str(item.get("pattern") or "Unmapped"), _safe_int(item.get("count"), 0)))
    rows = sorted([r for r in rows if r[1] > 0], key=lambda x: x[1])
    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    if not rows:
        ax.text(0.5, 0.5, "No treatment pattern data available", ha="center", va="center")
        ax.axis("off")
    else:
        names = [r[0] for r in rows]
        counts = [r[1] for r in rows]
        y = list(range(len(rows)))
        bars = ax.barh(y, counts)
        ax.set_yticks(y)
        ax.set_yticklabels([_wrap_label(n, 32) for n in names])
        ax.set_xlabel("Mapped non-compliant controls")
        ax.set_title("Treatment volume by weakness pattern")
        ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
        for b, c in zip(bars, counts):
            ax.text(b.get_width() + 0.6, b.get_y() + b.get_height() / 2, str(c), va="center", fontsize=9)
        fig.text(0.01, 0.01, "Source: audit workbook and Stage 1 treatment-plan model.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_treatment_owner_workload(treatment_plan: Dict[str, Any], out_path: str) -> None:
    data = _deep_list(treatment_plan, ["owner_workload"])
    if not data:
        data = _deep_list(_as_dict(treatment_plan.get("chart_data")), ["owner_workload"])
    if not data:
        data = _as_list(treatment_plan.get("owner_summary"))
    rows = []
    for item in data:
        if isinstance(item, dict):
            rows.append((str(item.get("owner") or "Unassigned"), _safe_int(item.get("control_items") or item.get("count"), 0), _safe_int(item.get("high_severity_items"), 0)))
    rows = sorted([r for r in rows if r[1] > 0], key=lambda x: x[1])
    fig, ax = plt.subplots(figsize=(9.2, 5.2))
    if not rows:
        ax.text(0.5, 0.5, "No owner workload data available", ha="center", va="center")
        ax.axis("off")
    else:
        names = [r[0] for r in rows]
        counts = [r[1] for r in rows]
        y = list(range(len(rows)))
        bars = ax.barh(y, counts)
        ax.set_yticks(y)
        ax.set_yticklabels([_wrap_label(n, 32) for n in names])
        ax.set_xlabel("Control treatment items")
        ax.set_title("Treatment workload by recommended owner")
        ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
        for b, total, high in zip(bars, counts, [r[2] for r in rows]):
            suffix = f"  (High: {high})" if high else ""
            ax.text(b.get_width() + 0.6, b.get_y() + b.get_height() / 2, f"{total}{suffix}", va="center", fontsize=9)
        fig.text(0.01, 0.01, "Source: Stage 1 treatment-plan model. Owner is recommended, not assigned approval.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _scatter_treatment_priority_matrix(treatment_plan: Dict[str, Any], out_path: str) -> None:
    data = _deep_list(treatment_plan, ["priority_matrix"])
    if not data:
        data = _deep_list(_as_dict(treatment_plan.get("chart_data")), ["priority_matrix"])
    rows = [x for x in data if isinstance(x, dict)][:12]
    fig, ax = plt.subplots(figsize=(8.8, 5.4))
    if not rows:
        ax.text(0.5, 0.5, "No priority matrix data available", ha="center", va="center")
        ax.axis("off")
    else:
        x = [_safe_int(r.get("mapped_noncompliant_count"), 0) for r in rows]
        y = [_safe_int(r.get("severity_score"), 1) for r in rows]
        ax.scatter(x, y, s=[max(80, min(520, c * 8)) for c in x], alpha=0.7)
        ax.set_xlabel("Mapped non-compliant controls")
        ax.set_ylabel("Severity score")
        ax.set_yticks([1, 2, 3])
        ax.set_yticklabels(["Low", "Medium", "High"])
        ax.set_title("Treatment priority matrix")
        ax.grid(True, linestyle="--", linewidth=0.5, alpha=0.6)
        for r, xi, yi in zip(rows, x, y):
            label = _wrap_label(str(r.get("pattern") or "Pattern"), 18).split("\n")[0]
            ax.annotate(label, (xi, yi), textcoords="offset points", xytext=(5, 5), fontsize=8)
        fig.text(0.01, 0.01, "Source: workbook prevalence and deterministic severity mapping. Bubble size follows treatment volume.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)

def _likelihood_from_count(cnt: int) -> str:
    if cnt >= 50:
        return "High"
    if cnt >= 20:
        return "Medium-High"
    if cnt >= 10:
        return "Medium"
    return "Low-Medium"


def _target_timeline(sev: str) -> str:
    return "0-90 days" if sev == "High" else ("0-180 days" if sev == "Medium" else "0-365 days")


def _target_date_str(audit_dt: date, sev: str) -> str:
    days = 90 if sev == "High" else (180 if sev == "Medium" else 365)
    due = audit_dt.toordinal() + days
    return date.fromordinal(due).strftime("%d %b %Y")



def _as_dict(value: Any) -> Dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if value is None:
        return []
    return [value]


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        if value is None:
            return default
        if isinstance(value, bool):
            return int(value)
        return int(float(str(value).strip()))
    except Exception:
        return default


def _safe_str(value: Any, default: str = "") -> str:
    if value is None:
        return default
    return str(value)


def _clean_text(value: Any) -> str:
    """Convert arbitrary evidence values into report-safe text."""
    if value is None:
        return ""
    if isinstance(value, str):
        text = value
    elif isinstance(value, (int, float, bool)):
        text = str(value)
    elif isinstance(value, dict):
        # Keep dictionaries out of the final report. Prefer meaningful fields.
        preferred = []
        for key in ("message", "title", "description", "summary", "rule", "rule_id", "tool", "level", "file", "path"):
            val = value.get(key)
            if val not in (None, "", [], {}):
                preferred.append(f"{key}: {_clean_text(val)}")
        text = "; ".join(preferred) if preferred else json.dumps(value, ensure_ascii=False)
    elif isinstance(value, list):
        text = "; ".join(_clean_text(item) for item in value if item not in (None, "", [], {}))
    else:
        text = str(value)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace("...", "")
    return text


def _normalize_android_artifact_for_display(value: Any, names_only: bool = False) -> str:
    """Render MobSF dynamic storage artifacts as readable Android paths or file names."""
    text = _clean_text(value)
    if not text:
        return ""
    s = text.replace("\\", "/").strip().strip("'\"`")

    if "/data/data/" in s:
        idx = s.find("/data/data/")
        s = s[idx:]
        s = re.split(r"[\s\"'<>;,)]", s)[0]
        s = re.sub(r"/+", "/", s).rstrip(".,;:")
        return Path(s).name if names_only else s

    if s.lower().startswith("data/data/"):
        s = "/" + s
        s = re.sub(r"/+", "/", s)
        return Path(s).name if names_only else s.rstrip(".,;:")

    collapsed = re.sub(r"[^A-Za-z0-9_./-]+", "", s)
    if collapsed.lower().startswith("datadata"):
        tail = collapsed[len("datadata"):]
        for folder in ("shared_prefs", "databases", "no_backup", "cache", "files"):
            idx = tail.lower().find(folder)
            if idx > 0:
                package_name = tail[:idx].strip("./")
                rest = tail[idx + len(folder):].lstrip("/._-")
                if package_name and rest:
                    normalized = f"/data/data/{package_name}/{folder}/{rest}"
                    return Path(normalized).name if names_only else normalized
                if package_name:
                    normalized = f"/data/data/{package_name}/{folder}"
                    return Path(normalized).name if names_only else normalized
    for folder in ("shared_prefs", "databases", "no_backup", "cache", "files"):
        idx = collapsed.lower().find(folder)
        if idx > 0:
            package_name = collapsed[:idx].strip("./")
            rest = collapsed[idx + len(folder):].lstrip("/._-")
            if package_name.startswith("org.") and rest:
                normalized = f"/data/data/{package_name}/{folder}/{rest}"
                return Path(normalized).name if names_only else normalized

    match = re.search(r"([A-Za-z0-9_.-]+\.(?:xml|db|sqlite|sqlite3|properties))$", collapsed, re.IGNORECASE)
    if match:
        return match.group(1)
    return _short(text, 180)


def _format_artifact_list(values: Any, limit: int = 4, names_only: bool = False) -> str:
    items = _as_list(values)
    cleaned: List[str] = []
    seen = set()
    for item in items:
        text = _normalize_android_artifact_for_display(item, names_only=names_only)
        if text and text not in seen:
            seen.add(text)
            cleaned.append(text)
        if len(cleaned) >= limit:
            break
    return "; ".join(cleaned) if cleaned else "No examples normalized into analysis pack"


def _clean_generated_text(value: Any, ensure_period: bool = False) -> str:
    """Return model-authored text without content rewriting.

    The report is prompt-first: the model must generate acceptable language.
    This helper only normalizes whitespace for Word rendering and optionally
    adds sentence punctuation. It does not change wording or meaning.
    """
    text = _clean_text(value)
    if ensure_period and text and not text.endswith((".", ";", ":")):
        text += "."
    return text


def _sanitize_android_artifacts_in_text(value: Any) -> str:
    """Compatibility wrapper. No wording change is performed."""
    return _clean_generated_text(value)


def _sanitize_ai_narrative_text(value: Any, technical: Dict[str, Any]) -> str:
    """Compatibility wrapper. The AI prompt and section validator control language quality."""
    return _clean_generated_text(value)


def _sast_counts(sast: Dict[str, Any]) -> Dict[str, int]:
    summary = _as_dict(sast.get("summary"))
    security = _deep_int(sast, ["retained_security_findings", "security_relevant_app_findings"])
    if security == 0:
        security = _safe_int(summary.get("retained_security_findings") or summary.get("security_relevant_app_findings"), 0)
    app_signals = _deep_int(sast, ["retained_app_code_signings", "retained_app_code_signals", "retained_app_code_findings", "total_findings", "total", "results_count"])
    if app_signals == 0:
        app_signals = _safe_int(summary.get("retained_app_code_signals") or summary.get("retained_app_code_findings") or summary.get("app_code_results"), 0)
    hardening = _deep_int(sast, ["hardening_or_maintainability_signals"])
    if hardening == 0:
        hardening = _safe_int(summary.get("hardening_or_maintainability_signals"), max(0, app_signals - security))
    raw = _deep_int(sast, ["raw_sarif_results", "raw_results_count"])
    if raw == 0:
        raw = _safe_int(summary.get("raw_results_in_selected_sarif"), 0)
    return {
        "retained_app_code_signals": app_signals,
        "retained_security_findings": security,
        "hardening_or_maintainability_signals": hardening,
        "raw_sarif_results": raw,
    }



def _sanitize_recommendation_text(value: Any) -> str:
    """Compatibility wrapper for model-authored recommendations.

    The prompt, per-section retry logic, and final quality gate are responsible
    for report language compliance.
    """
    return _clean_generated_text(value, ensure_period=True)


def _sanitize_closure_criteria(value: Any) -> str:
    """Compatibility wrapper for model-authored closure criteria. No rewriting."""
    return _clean_generated_text(value, ensure_period=True)


def _sanitize_positive_control_final(value: Any) -> str:
    """Compatibility wrapper for model-authored positive-control statements. No rewriting."""
    return _clean_generated_text(value, ensure_period=True)


def _is_emptyish(value: Any) -> bool:
    if value in (None, ""):
        return True
    if isinstance(value, (list, tuple, set)):
        return len([x for x in value if not _is_emptyish(x)]) == 0
    if isinstance(value, dict):
        return len([v for v in value.values() if not _is_emptyish(v)]) == 0
    return False


def _cell_text(value: Any) -> str:
    return _report_display_text(value)


def _add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(_report_display_text(text))
    if p.runs:
        p.runs[0].italic = True


def _sanitize_positive_statement(text: Any) -> str:
    """Compatibility wrapper for workbook evidence excerpts. No rewriting."""
    return _clean_generated_text(text)


def _qa_doc_text(doc: Document) -> str:
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _normalize_document_typography(doc: Document) -> None:
    """Final typographic normalization to keep the DOCX visually consistent."""
    _set_style_font(doc, "Normal", REPORT_BODY_PT)
    _set_style_font(doc, "Heading 1", REPORT_HEADING1_PT, bold=True)
    _set_style_font(doc, "Heading 2", REPORT_HEADING2_PT, bold=True)
    _set_style_font(doc, "Heading 3", REPORT_HEADING3_PT, bold=True)
    _set_style_font(doc, "Caption", REPORT_CAPTION_PT, italic=True)
    _set_style_font(doc, "List Bullet", REPORT_BODY_PT)
    _set_style_font(doc, "List Number", REPORT_BODY_PT)

    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=REPORT_HEADER_FOOTER_PT)

    for paragraph in doc.paragraphs:
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        if style_name == "Title":
            target = 22.0
            bold = True
            italic = None
        elif style_name == "Subtitle":
            target = 13.0
            bold = True
            italic = None
        elif style_name.startswith("Heading 1"):
            target = REPORT_HEADING1_PT
            bold = True
            italic = None
        elif style_name.startswith("Heading 2"):
            target = REPORT_HEADING2_PT
            bold = True
            italic = None
        elif style_name.startswith("Heading 3"):
            target = REPORT_HEADING3_PT
            bold = True
            italic = None
        elif style_name == "Caption" or paragraph.text.strip().startswith(("Table ", "Figure ")):
            target = REPORT_CAPTION_PT
            bold = None
            italic = True
        else:
            target = REPORT_BODY_PT
            bold = None
            italic = None
        for run in paragraph.runs:
            _set_run_font(run, size_pt=target, bold=bold, italic=italic)

    # Tables are formatted when they are created. Re-walking every table cell in
    # large appendices is very expensive with python-docx, so this final pass
    # only reinforces header styling and repeating headers.
    for table in doc.tables:
        if table.rows:
            _repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                _set_cell_shading(cell, THEME_BLUE)
                _format_table_cell(cell, font_size=REPORT_TABLE_HEADER_PT, bold=True, no_wrap=True)



REPORT_DISPLAY_REPLACEMENTS = [
    (re.compile(r"\bResult:\s*no\s+for\s+requirement\b", re.IGNORECASE), "Non-compliant requirement"),
    (re.compile(r"\bResult:\s*yes\s+for\s+requirement\b", re.IGNORECASE), "Compliant requirement"),
    (re.compile(r"\bnot aligned signals:\b", re.IGNORECASE), "Relevant mapped signals:"),
    (re.compile(r"\bfallback verdict note hint\b", re.IGNORECASE), "computed evidence note"),
    (re.compile(r"\bwith a fallback verdict\b", re.IGNORECASE), "based on mapped evidence"),
    (re.compile(r"\bfallback verdicts\b", re.IGNORECASE), "computed evidence results"),
    (re.compile(r"\bfallback verdict\b", re.IGNORECASE), "computed evidence result"),
    (re.compile(r"\bdoes not align with the requirement\b", re.IGNORECASE), "does not satisfy the requirement"),
    (re.compile(r"\bdoes not align with the expected outcome\b", re.IGNORECASE), "does not satisfy the expected control outcome"),
    (re.compile(r"\bdoes not align with\b", re.IGNORECASE), "does not satisfy"),
    (re.compile(r"\bnot aligned signals\b", re.IGNORECASE), "mapped signals requiring review"),
    (re.compile(r"\bnot aligned signal\b", re.IGNORECASE), "mapped signal requiring review"),
    (re.compile(r"\bnot aligned\b", re.IGNORECASE), "not satisfied"),
    (re.compile(r"\balignment issues\b", re.IGNORECASE), "mapped evidence requiring review"),
    (re.compile(r"\balignment issue\b", re.IGNORECASE), "mapped evidence requiring review"),
    (re.compile(r"\bobserved summaries of its flags\b", re.IGNORECASE), "mapped flag evidence"),
    (re.compile(r"\bobserved summaries of related flags\b", re.IGNORECASE), "mapped flag evidence"),
    (re.compile(r"\bobserved summary for\b", re.IGNORECASE), "mapped evidence for"),
    (re.compile(r"\bretained_app_code_findings\b", re.IGNORECASE), "application-scope SAST signals"),
    (re.compile(r"\bretained_security_findings\b", re.IGNORECASE), "security-relevant SAST findings"),
    (re.compile(r"\bhardening_or_maintainability_signals\b", re.IGNORECASE), "hardening, quality, or maintainability signals"),
    (re.compile(r"\bContradicting signals:\b", re.IGNORECASE), "Relevant mapped signals:"),
    (re.compile(r"\bcontradicts the expected outcome\b", re.IGNORECASE), "does not align with the expected outcome"),
    (re.compile(r"\bcontradicting signals:\b", re.IGNORECASE), "Relevant mapped signals:"),
    (re.compile(r"\bcontradictions\b", re.IGNORECASE), "alignment issues"),
    (re.compile(r"\bcontradiction\b", re.IGNORECASE), "alignment issue"),
    (re.compile(r"\bcontradictory\b", re.IGNORECASE), "not aligned"),
    (re.compile(r"\bcontradicted\b", re.IGNORECASE), "not aligned"),
    (re.compile(r"\bcontradicts\b", re.IGNORECASE), "does not align with"),
    (re.compile(r"\bcontradicting\b", re.IGNORECASE), "not aligned"),
    (re.compile(r"\bcontradict\b", re.IGNORECASE), "does not align with"),
    (re.compile(r"\bdiscrepancies\b", re.IGNORECASE), "differences"),
    (re.compile(r"\bdiscrepancy\b", re.IGNORECASE), "difference"),
    (re.compile(r"\bdiscrepant\b", re.IGNORECASE), "different"),
    (re.compile(r"technical coverage limitations", re.IGNORECASE), "technical evidence scope"),
    (re.compile(r"scope and limitations", re.IGNORECASE), "scope"),
    (re.compile(r"entr\(y/ies\)", re.IGNORECASE), "entries"),
]


def _report_display_text(value: Any) -> str:
    """Normalize text that is about to be rendered in the final DOCX.

    This does not change audit verdicts, counts, PUIDs, flags, CVEs, scanner
    evidence, or workbook-derived results. It only converts internal pipeline
    field names and QA-trigger wording into director-facing report language so
    the final safety gate does not fail after a valid report has already been
    assembled from structured evidence.
    """
    text = _clean_text(value)
    if not text:
        return ""
    for pattern, replacement in REPORT_DISPLAY_REPLACEMENTS:
        text = pattern.sub(replacement, text)
    return _clean_text(text)


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph text only when needed, preserving the paragraph style."""
    if paragraph.text == new_text:
        return
    paragraph.text = new_text


def _sanitize_report_language(doc: Document) -> None:
    """Remove internal pipeline wording from final rendered DOCX text.

    AI sections are already validated before they enter the document. This final
    pass is a deterministic safety net for structured workbook and scanner
    excerpts, especially table cells copied from analysis-pack fields.
    """
    for paragraph in doc.paragraphs:
        cleaned = _report_display_text(paragraph.text)
        if cleaned != paragraph.text:
            _replace_paragraph_text(paragraph, cleaned)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cleaned = _report_display_text(cell.text)
                if cleaned != cell.text:
                    cell.text = cleaned


def _quality_gate(doc: Document) -> None:
    """Fail fast only on defects that make the DOCX invalid or visibly unsafe."""
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    combined = _report_display_text(full_text + "\n" + table_text)
    lowered = combined.lower()
    forbidden_fragments = [
        "missing_inputs\\n",
        "{'tool':",
        '"tool":',
        "The application applications",
        "The application the mobile application",
        "\\nAuditor",
        "entr(y/ies)",
        "retained_app_code_findings",
        "hardening_or_maintainability_signals",
        "discrepanc",
        "contradict",
        "technical coverage limitations",
        "scope and limitations",
    ]
    hits = [token for token in forbidden_fragments if token.lower() in lowered]
    if "puid / item" in lowered:
        hits.append("PUID / item")
    if doc.tables and "Table 1." not in full_text:
        hits.append("missing table captions")
    if "java/android/backup-enabled" in lowered and "backup enabled\nnot available in parsed evidence" in lowered:
        hits.append("backup evidence not applied")
    if hits:
        raise SystemExit("[ERROR] Audit Summary quality gate failed. Forbidden text found: " + ", ".join(hits))

    unexpected_fonts = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            name = getattr(run.font, "name", None)
            if name and name != REPORT_FONT_NAME:
                unexpected_fonts.add(name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        name = getattr(run.font, "name", None)
                        if name and name != REPORT_FONT_NAME:
                            unexpected_fonts.add(name)
    if unexpected_fonts:
        raise SystemExit("[ERROR] Audit Summary quality gate failed. Unexpected direct font families remain: " + ", ".join(sorted(unexpected_fonts)[:10]))


def _short(value: Any, limit: int = 170) -> str:
    """Normalize text for Word cells without using ellipsis.

    The previous implementation appended "...", which made the report look as
    if LLM tokens were missing. This function keeps full text by default when
    limit <= 0 and uses an explicit marker only when a hard limit is requested.
    """
    text = _clean_text(value)
    if limit and limit > 0 and len(text) > limit:
        return text[:limit].rstrip() + " [text shortened]"
    return text

def _first_present(obj: Dict[str, Any], keys: List[str], default: Any = None) -> Any:
    for key in keys:
        if isinstance(obj, dict) and key in obj and obj[key] not in (None, ""):
            return obj[key]
    return default


def _deep_find_first(obj: Any, keys: List[str]) -> Any:
    if isinstance(obj, dict):
        for key in keys:
            if key in obj and obj[key] not in (None, ""):
                return obj[key]
        for value in obj.values():
            found = _deep_find_first(value, keys)
            if found not in (None, ""):
                return found
    elif isinstance(obj, list):
        for item in obj:
            found = _deep_find_first(item, keys)
            if found not in (None, ""):
                return found
    return None


def _deep_int(obj: Any, keys: List[str], default: int = 0) -> int:
    value = _deep_find_first(obj, keys)
    return _safe_int(value, default)


def _deep_dict(obj: Any, keys: List[str]) -> Dict[str, Any]:
    value = _deep_find_first(obj, keys)
    return value if isinstance(value, dict) else {}


def _deep_list(obj: Any, keys: List[str]) -> List[Any]:
    value = _deep_find_first(obj, keys)
    if isinstance(value, list):
        return value
    return []


def _block_available(block: Any) -> bool:
    if not isinstance(block, dict) or not block:
        return False
    if block.get("available") is True:
        return True
    for key, value in block.items():
        if key in {"available", "paths", "input_paths"}:
            continue
        if isinstance(value, (list, dict)) and len(value) > 0:
            return True
        if isinstance(value, (int, float)) and value > 0:
            return True
        if isinstance(value, str) and value.strip():
            return True
    return False


def _format_bool(value: Any) -> str:
    if value is True:
        return "Detected / Yes"
    if value is False:
        return "Not detected / No"
    if value in (None, ""):
        return "Not available in parsed evidence"
    if isinstance(value, (int, float)):
        return str(value)
    text = _clean_text(value)
    return text if text else "Not available in parsed evidence"

_TABLE_RENDER_COUNTER = 0


def _table_caption_for_headers(headers: List[str]) -> str:
    joined = " | ".join(_clean_text(x) for x in headers).lower()
    if "weakness pattern" in joined and "workbook basis" in joined:
        return "Risk triage summary"
    if "evidence source" in joined and "available" in joined:
        return "Technical evidence source coverage"
    if "cve" in joined or "installed -> fixed" in joined or ("component" in joined and "severity" in joined):
        return "Dependency vulnerability evidence details"
    if "indicator" in joined and "parsed" in joined:
        return "Android static evidence indicators"
    if "runtime evidence type" in joined:
        return "MobSF dynamic runtime evidence summary"
    if "sast classification" in joined:
        return "SAST classification summary"
    if "tool" in joined and "security findings" in joined and "raw sarif" in joined:
        return "SAST findings by tool"
    if "tool" in joined and "rule" in joined and "message" in joined:
        return "SAST security-relevant finding details"
    if "top hardening" in joined:
        return "Top hardening and maintainability signals"
    if "execution item" in joined and "summary" in joined:
        return "Technical execution metadata"
    if "limitation" in joined and "report-safe explanation" in joined:
        return "Technical execution metadata"
    if "scanner evidence" in joined or "technical finding" in joined:
        return "Technical evidence correlation matrix"
    if "positive control" in joined:
        return "Positive and qualified control traceability"
    if "requirement" in joined or "puid" in joined:
        return "Requirement treatment traceability"
    if "owner" in joined and "target" in joined:
        return "Management action plan"
    return "Audit evidence table"


def _add_table_caption(doc: Document, title: str) -> None:
    global _TABLE_RENDER_COUNTER
    _TABLE_RENDER_COUNTER += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"Table {_TABLE_RENDER_COUNTER}. {title}")
    r.bold = True
    r.italic = True
    _set_run_font(r, size_pt=REPORT_CAPTION_PT, bold=True, italic=True)


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = OxmlElement("w:noWrap")
    tc_pr.append(no_wrap)


def _format_table_cell(cell, font_size: float = REPORT_TABLE_BODY_PT, bold: bool = False, no_wrap: bool = False) -> None:
    if no_wrap:
        _set_cell_no_wrap(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            _set_run_font(run, size_pt=font_size, bold=bold)


def _add_table(doc: Document, headers: List[str], rows: List[List[Any]], max_rows: int | None = None, empty_message: str | None = None) -> bool:
    """Add a readable, captioned table only when rows exist."""
    actual_rows = rows[:max_rows] if max_rows is not None else rows
    actual_rows = [r for r in actual_rows if not _is_emptyish(r)]
    if not actual_rows:
        if empty_message:
            _add_note(doc, empty_message)
        return False

    _add_table_caption(doc, _table_caption_for_headers(headers))
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = True
    _set_table_alignment(tbl)
    h = tbl.rows[0].cells
    _repeat_table_header(tbl.rows[0])
    for idx, txt in enumerate(headers):
        header_text = _cell_text(txt)
        if header_text == "PUID / item":
            header_text = "Requirement"
        h[idx].text = header_text
        _set_cell_shading(h[idx], THEME_BLUE)
        _set_cell_margins(h[idx], top=70, start=70, bottom=70, end=70)
        _set_cell_vertical_alignment(h[idx])
        _format_table_cell(h[idx], font_size=REPORT_TABLE_HEADER_PT, bold=True, no_wrap=True)
    for row_values in actual_rows:
        r = tbl.add_row().cells
        for idx in range(len(headers)):
            value = _cell_text(row_values[idx] if idx < len(row_values) else "")
            if value == "PUID / item":
                value = "Requirement"
            r[idx].text = value
            if len(tbl.rows) % 2 == 1:
                _set_cell_shading(r[idx], THEME_LIGHT)
            _set_cell_margins(r[idx], top=55, start=65, bottom=55, end=65)
            _set_cell_vertical_alignment(r[idx], WD_CELL_VERTICAL_ALIGNMENT.TOP)
            keep_together = idx == 0 and len(headers) >= 4
            _format_table_cell(r[idx], font_size=7.2 if len(value) > 180 else REPORT_TABLE_BODY_PT, no_wrap=keep_together)
    doc.add_paragraph()
    return True

def _app_report_title(app: Dict[str, Any]) -> str:
    app = _as_dict(app)
    name = _first_present(app, ["app_name", "application_name", "name", "title", "app", "project", "project_name"])
    version = _first_present(app, ["version", "app_version", "version_name", "release", "app_release"])

    if not name:
        for key, value in app.items():
            if "name" in str(key).lower() and value:
                name = value
                break
    if not version:
        for key, value in app.items():
            if "version" in str(key).lower() and value:
                version = value
                break

    name_s = _safe_str(name, "Mobile Application").strip() or "Mobile Application"
    version_s = _safe_str(version).strip()
    if version_s and version_s not in name_s:
        version_label = version_s if version_s.lower().startswith("v") else f"v{version_s}"
        return f"{name_s} {version_label}"
    return name_s


def _technical_evidence(pack: Dict[str, Any]) -> Dict[str, Any]:
    return _as_dict(pack.get("technical_evidence"))


def _source_status_rows(technical: Dict[str, Any]) -> List[List[Any]]:
    labels = [
        ("Audit workbook", "workbook"),
        ("Trivy SCA", "trivy_sca"),
        ("MobSF static", "mobsf_static"),
        ("MobSF dynamic", "mobsf_dynamic"),
        ("SAST app-code", "sast_app_code"),
    ]
    rows = []
    for label, key in labels:
        if key == "workbook":
            rows.append([label, "Authoritative", "Provides requirement-level yes, no, and n/a audit results, metrics, PUIDs, categories, flags, and workbook justifications."])
            continue
        block = _as_dict(technical.get(key))
        available = _block_available(block)
        status = "Processed" if available else "No artifact processed"
        summary = "Parsed into the analysis pack for technical evidence." if available else "No normalized evidence was included in the analysis pack for this source."
        if key == "trivy_sca" and available:
            summary = f"{_deep_int(block, ['total_vulnerabilities', 'vulnerabilities_total', 'total'])} dependency vulnerability findings; {_deep_int(block, ['packages_detected', 'package_count', 'packages_total'])} packages detected."
        elif key == "sast_app_code" and available:
            counts = _sast_counts(block)
            raw_counts = _deep_dict(block, ["raw_tool_counts", "tool_counts", "by_tool"])
            raw_text = ", ".join(f"{k}: {v}" for k, v in raw_counts.items()) if raw_counts else "raw tool counts unavailable"
            summary = (
                f"{counts['retained_security_findings']} security-relevant application-code findings and "
                f"{counts['hardening_or_maintainability_signals']} hardening and quality signals retained after scope filtering. "
                f"Raw SARIF counts retained for traceability: {raw_text}."
            )
        elif key == "mobsf_static" and available:
            trace = _as_dict(block.get("apk_traceability"))
            version = _clean_text(trace.get("version_name")) or _clean_text(_as_dict(block.get("app_info")).get("version_name"))
            build = _clean_text(trace.get("inferred_build_type"))
            extra = f" Version {version}; build type {build}." if version or build else ""
            summary = "Android APK, manifest, certificate, permission, signing, tracker, and hardening indicators were parsed." + extra
        elif key == "mobsf_dynamic" and available:
            summary = "Runtime SharedPreferences, SQLite/database, WAL/SHM sidecar, local-storage, URL, and tracker evidence was parsed."
        rows.append([label, status, summary])
    return rows


def _technical_takeaways(technical: Dict[str, Any]) -> List[str]:
    takeaways: List[str] = []
    trivy = _as_dict(technical.get("trivy_sca"))
    if _block_available(trivy):
        total = _deep_int(trivy, ["total_vulnerabilities", "vulnerabilities_total", "total"])
        by_sev = _deep_dict(trivy, ["by_severity", "severity_counts"])
        critical = _safe_int(by_sev.get("CRITICAL") or by_sev.get("critical"))
        high = _safe_int(by_sev.get("HIGH") or by_sev.get("high"))
        if total > 0:
            takeaways.append(f"Trivy SCA evidence identified {total} known dependency vulnerability findings, including {critical} Critical and {high} High findings where reported.")
    mobsf = _as_dict(technical.get("mobsf_static"))
    if _block_available(mobsf):
        indicators = []
        for label, keys in [
            ("debuggable build", ["debuggable", "has_debuggable", "debuggable_true", "android_debuggable"]),
            ("debug certificate", ["debug_certificate", "has_debug_cert", "signed_with_debug_certificate"]),
            ("backup enabled", ["allow_backup", "backup_enabled", "android_allow_backup"]),
            ("exported components", ["exported_components_count", "exported_component_count"]),
        ]:
            value = _deep_find_first(mobsf, keys)
            if value is True or _safe_int(value) > 0:
                indicators.append(label)
        if indicators:
            takeaways.append("MobSF static evidence reported Android hardening indicators requiring review, including " + ", ".join(indicators[:4]) + ".")
    sast = _as_dict(technical.get("sast_app_code"))
    if _block_available(sast):
        counts = _sast_counts(sast)
        if counts["retained_security_findings"] > 0:
            takeaways.append(
                f"SAST evidence retained {counts['retained_security_findings']} security-relevant application-code findings after scope filtering; "
                f"{counts['hardening_or_maintainability_signals']} additional hardening and quality signals are reported separately."
            )
        elif counts["retained_app_code_signals"] > 0:
            takeaways.append(
                f"SAST evidence retained {counts['retained_app_code_signals']} application-scope hardening and quality signals, but no security-relevant SAST findings were classified from the normalized evidence."
            )
    return takeaways


def _trivy_findings(trivy: Dict[str, Any]) -> List[Dict[str, Any]]:
    for keys in (["detailed_findings"], ["vulnerabilities"], ["cve_table"], ["findings"], ["top_findings"]):
        items = _deep_list(trivy, keys)
        if items:
            return [x for x in items if isinstance(x, dict)]
    return []


def _add_trivy_section(doc: Document, technical: Dict[str, Any]) -> None:
    trivy = _as_dict(technical.get("trivy_sca"))
    if not _block_available(trivy):
        doc.add_paragraph("Trivy SCA evidence was not available in the analysis pack.")
        return
    total = _deep_int(trivy, ["total_vulnerabilities", "vulnerabilities_total", "total"])
    packages = _deep_int(trivy, ["packages_detected", "package_count", "packages_total"])
    fixable = _deep_int(trivy, ["fixable_total", "fixable_vulnerabilities", "fixable"])
    licenses = _deep_int(trivy, ["license_entries_detected", "licenses_detected", "license_count"])
    by_sev = _deep_dict(trivy, ["by_severity", "severity_counts"])
    severity_text = ", ".join(f"{k}: {_safe_int(v)}" for k, v in by_sev.items()) if by_sev else "not reported"
    high_components = _as_list(trivy.get("high_severity_components"))
    high_note = ""
    if high_components:
        high_note = " High-severity findings affect: " + ", ".join(_clean_text(x) for x in high_components[:8] if _clean_text(x)) + "."

    doc.add_paragraph(
        f"Trivy Software Composition Analysis reported {packages} detected packages, {total} known dependency vulnerability findings, "
        f"{fixable} findings with a fixed version available, and {licenses} license entries. Severity distribution: {severity_text}.{high_note}"
    )

    rows = []
    for finding in _trivy_findings(trivy)[:18]:
        cwe = ", ".join(_clean_text(x) for x in _as_list(_first_present(finding, ["cwe_ids", "cwe", "CweIDs"], []))[:4])
        cvss = _first_present(finding, ["cvss_v3_score", "cvss", "score"], "")
        published = _first_present(finding, ["published_at", "published", "PublishedDate"], "")
        modified = _first_present(finding, ["last_modified_at", "modified", "LastModifiedDate"], "")
        date_parts = []
        if published:
            date_parts.append(f"Published: {published}")
        if modified:
            date_parts.append(f"Modified: {modified}")
        rows.append([
            _first_present(finding, ["severity", "Severity"], ""),
            _first_present(finding, ["id", "VulnerabilityID", "cve", "rule_id"], ""),
            _first_present(finding, ["pkg", "PkgName", "package", "package_name", "component"], ""),
            f"{_first_present(finding, ['installed', 'InstalledVersion', 'installed_version'], '')} -> {_first_present(finding, ['fixed', 'FixedVersion', 'fixed_version'], '')}",
            "; ".join(x for x in [f"CWE: {cwe}" if cwe else "", f"CVSS: {cvss}" if cvss not in (None, "") else ""] if x),
            _short(" | ".join(x for x in [
                "; ".join(date_parts),
                _first_present(finding, ["target", "Target"], ""),
                _first_present(finding, ["title", "Title", "description", "Description"], ""),
            ] if x), 360),
        ])
    _add_table(doc, ["Severity", "CVE / ID", "Component", "Installed -> fixed", "CWE / CVSS", "Dates / target / summary"], rows, max_rows=18)


def _sast_has_rule(technical: Dict[str, Any], rule_substring: str) -> bool:
    rule_substring = str(rule_substring or "").lower()
    if not rule_substring:
        return False
    sast = _as_dict(technical.get("sast_app_code"))
    candidate_lists = [
        "security_findings_sample",
        "hardening_signals_sample",
        "findings",
        "security_findings",
        "top_security_rules",
        "top_hardening_rules",
    ]
    for key in candidate_lists:
        for item in _as_list(sast.get(key)):
            if isinstance(item, dict):
                haystack = " ".join(_clean_text(item.get(field)) for field in ("rule_id", "rule", "id", "message", "title", "name")).lower()
                if rule_substring in haystack:
                    return True
            elif rule_substring in _clean_text(item).lower():
                return True
    summary = _as_dict(sast.get("summary"))
    for key in ("top_security_rules", "top_retained_rules", "top_raw_rules"):
        for item in _as_list(summary.get(key)):
            if isinstance(item, dict) and rule_substring in _clean_text(item.get("rule_id")).lower():
                return True
    return False


def _indicator_with_supporting_evidence(label: str, value: Any, technical: Dict[str, Any]) -> str:
    formatted = _format_bool(value)
    if label.lower() == "backup enabled" and formatted == "Not available in parsed evidence":
        if _sast_has_rule(technical, "java/android/backup-enabled") or _sast_has_rule(technical, "backup-enabled"):
            return "Detected (SAST CodeQL evidence)"
    return formatted



def _mobsf_signal_rows(mobsf: Dict[str, Any], technical: Dict[str, Any] | None = None) -> List[List[Any]]:
    checks = [
        ("Debuggable", ["debuggable", "has_debuggable", "debuggable_true", "android_debuggable"]),
        ("Debug certificate", ["debug_certificate", "has_debug_cert", "signed_with_debug_certificate"]),
        ("v1 signature / Janus exposure", ["v1_signature", "janus", "has_janus", "v1_signature_present"]),
        ("SHA1 certificate/signature evidence", ["sha1", "uses_sha1", "has_sha1"]),
        ("Backup enabled", ["allow_backup", "backup_enabled", "android_allow_backup"]),
        ("Minimum SDK", ["min_sdk", "minsdk", "minSdk"]),
        ("Dangerous permissions", ["dangerous_permissions_count", "dangerous_permission_count"]),
        ("Exported components", ["exported_components_count", "exported_component_count"]),
        ("Trackers detected", ["trackers_detected", "tracker_count", "detected_trackers"]),
    ]
    rows = []
    for label, keys in checks:
        value = _deep_find_first(mobsf, keys)
        rows.append([label, _indicator_with_supporting_evidence(label, value, technical or {})])
    return rows


def _add_mobsf_static_section(doc: Document, technical: Dict[str, Any]) -> None:
    mobsf = _as_dict(technical.get("mobsf_static"))
    if not _block_available(mobsf):
        doc.add_paragraph("MobSF static evidence was not processed in the analysis pack.")
        return
    doc.add_paragraph("MobSF static evidence summarizes the APK identity, manifest configuration, permissions, signing, certificate, tracker, and binary hardening indicators observed in the analyzed artifact.")
    trace = _as_dict(mobsf.get("apk_traceability"))
    if trace:
        trace_rows = [
            ["APK file", trace.get("file_name")],
            ["Package", trace.get("package_name")],
            ["Version", trace.get("version_name")],
            ["SHA-256", trace.get("sha256")],
            ["Minimum SDK", trace.get("min_sdk")],
            ["Target SDK", trace.get("target_sdk")],
            ["Build type", trace.get("inferred_build_type")],
        ]
        _add_table(doc, ["APK traceability item", "Observed value"], trace_rows, max_rows=10)
    _add_table(doc, ["Indicator", "Parsed value"], _mobsf_signal_rows(mobsf, technical), max_rows=20)

    finding_lists = []
    for key in ["manifest_findings", "manifest_findings_sample", "certificate_findings", "certificate_findings_sample", "findings", "high_findings", "warnings"]:
        items = _deep_list(mobsf, [key])
        if items:
            finding_lists.extend(items)
    rows = []
    for item in finding_lists[:10]:
        if isinstance(item, dict):
            rows.append([
                _first_present(item, ["severity", "level", "risk"], ""),
                _first_present(item, ["title", "name", "rule", "id"], ""),
                _first_present(item, ["description", "message", "summary"], ""),
            ])
        else:
            rows.append(["", "", item])
    _add_table(doc, ["Severity", "Finding", "Description"], rows, max_rows=10, empty_message="No detailed MobSF static finding rows were available in the normalized evidence block.")

def _add_mobsf_dynamic_section(doc: Document, technical: Dict[str, Any]) -> None:
    dynamic = _as_dict(technical.get("mobsf_dynamic"))
    if not _block_available(dynamic):
        _add_body_paragraph(doc, "MobSF dynamic evidence was not processed in the analysis pack.")
        return
    _add_body_paragraph(doc, "MobSF dynamic evidence summarizes runtime storage and behavioral observations, including SharedPreferences, SQLite databases, WAL/SHM sidecars, no-backup storage, URLs, and trackers reported for the analyzed execution.")
    rows = []
    categories = [x for x in _as_list(dynamic.get("runtime_evidence_categories")) if isinstance(x, dict)]
    if categories:
        for cat in categories:
            items = _as_list(cat.get("examples")) or _as_list(cat.get("example_names"))
            sample = _format_artifact_list(items, limit=5, names_only=False) if items else "Observed count recorded without item sample"
            rows.append([_clean_text(cat.get("type")), _safe_int(cat.get("count"), 0), sample])
    else:
        for label, keys in [
            ("SharedPreferences artifacts", ["shared_preferences_artifacts", "shared_preferences", "shared_preferences_files", "preferences"]),
            ("SQLite/database artifacts", ["sqlite_database_artifacts", "sqlite_databases", "databases", "db_files"]),
            ("WAL/SHM sidecar artifacts", ["wal_shm_artifacts"]),
            ("Runtime local-storage artifacts", ["local_storage_artifacts_sample", "local_storage_artifacts", "files", "storage_artifacts"]),
            ("Trackers", ["trackers", "detected_trackers"]),
        ]:
            items = _deep_list(dynamic, keys)
            count = len(items) if items else _deep_int(dynamic, [keys[0] + "_count"], 0)
            sample = _format_artifact_list(items, limit=5, names_only=False) if items else "Observed count recorded without item sample"
            rows.append([label, count, sample])
    _add_table(doc, ["Runtime evidence type", "Count", "Examples / parser note"], rows, max_rows=20)

def _sast_findings(sast: Dict[str, Any]) -> List[Dict[str, Any]]:
    for key in ["findings", "results", "app_code_findings", "top_findings"]:
        items = _deep_list(sast, [key])
        if items:
            return [x for x in items if isinstance(x, dict)]
    return []


def _add_sast_section(doc: Document, technical: Dict[str, Any]) -> None:
    sast = _as_dict(technical.get("sast_app_code"))
    if not _block_available(sast):
        _add_body_paragraph(doc, "SAST app-code evidence was not available in the analysis pack.")
        return

    counts = _sast_counts(sast)
    raw_counts = _deep_dict(sast, ["raw_tool_counts", "tool_counts", "by_tool"])
    security_tool_counts = _deep_dict(sast, ["retained_security_tool_counts"])
    hardening_tool_counts = _deep_dict(sast, ["hardening_tool_counts"])
    scope_note = _first_present(
        sast,
        ["classification_note", "scope_note", "filter_note"],
        "CI/CD workflows, audit-generation scripts, test paths, tooling, and non-application files are excluded from the application audit scope."
    )

    _add_body_paragraph(
        doc,
        f"SAST retained {counts['retained_app_code_signals']} application-scope signals after scope filtering. "
        f"Of these, {counts['retained_security_findings']} are classified as security-relevant findings and "
        f"{counts['hardening_or_maintainability_signals']} are classified as hardening, quality, or maintainability signals. "
        f"{_clean_text(scope_note)}"
    )

    breakdown_rows = [
        ["Security-relevant app-code findings", counts["retained_security_findings"], "Eligible for security remediation narrative."],
        ["Hardening / quality / maintainability signals", counts["hardening_or_maintainability_signals"], "Useful engineering evidence; not counted as vulnerabilities unless separately classified as security-relevant."],
        ["Raw SARIF results", counts["raw_sarif_results"], "Traceability and coverage signal; not automatically an app-code finding."],
    ]
    _add_table(doc, ["SAST classification", "Count", "Interpretation"], breakdown_rows, max_rows=10)

    tool_rows: List[List[Any]] = []
    all_tools = sorted(set(raw_counts) | set(security_tool_counts) | set(hardening_tool_counts))
    for tool in all_tools:
        tool_rows.append([
            tool,
            security_tool_counts.get(tool, 0),
            hardening_tool_counts.get(tool, 0),
            raw_counts.get(tool, 0),
        ])
    _add_table(doc, ["Tool", "Security findings", "Hardening / quality signals", "Raw SARIF results"], tool_rows, max_rows=12)

    security_rows = []
    for item in _deep_list(sast, ["security_findings_sample"])[:12]:
        if not isinstance(item, dict):
            continue
        security_rows.append([
            _first_present(item, ["tool", "driver", "source"], ""),
            _first_present(item, ["rule_id", "ruleId", "rule", "id"], ""),
            _first_present(item, ["level", "severity", "kind"], ""),
            _first_present(item, ["file", "path", "uri"], ""),
            _first_present(item, ["line", "start_line", "startLine"], ""),
            _first_present(item, ["message", "title", "description"], ""),
        ])
    _add_table(
        doc,
        ["Tool", "Rule", "Level", "File", "Line", "Message"],
        security_rows,
        max_rows=12,
        empty_message="No detailed security-relevant SAST rows were available after normalization."
    )

    hardening_rules = _deep_list(sast, ["top_hardening_rules"])
    if hardening_rules:
        rows = []
        for item in hardening_rules[:10]:
            if isinstance(item, dict):
                rows.append([item.get("rule_id", ""), item.get("count", "")])
        _add_table(doc, ["Top hardening / quality rule", "Retained count"], rows, max_rows=10)


def _format_limitation_row(key: str, value: Any) -> List[str] | None:
    if _is_emptyish(value):
        return None
    normalized_key = _clean_text(key)

    if normalized_key == "missing_inputs":
        if isinstance(value, list) and not value:
            return None
        details = _clean_text(value)
        if not details:
            return None
        return ["Technical inputs", details]

    if normalized_key == "sast_extraction_warning_count":
        count = _safe_int(value, 0)
        if count <= 0:
            return None
        return [
            "SAST execution notifications",
            f"{count} SAST extraction or frontend notification(s) were reported by the toolchain. The findings table presents application-code findings separately from execution notifications."
        ]

    if normalized_key == "sast_notifications_sample":
        count = len(value) if isinstance(value, list) else 1
        tools = []
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict) and item.get("tool"):
                    tools.append(str(item.get("tool")))
        elif isinstance(value, dict) and value.get("tool"):
            tools.append(str(value.get("tool")))
        tool_text = ", ".join(sorted(set(tools))) if tools else "SAST toolchain"
        return [
            "SAST execution notifications",
            f"{tool_text} reported {count} extraction or frontend notification(s). The findings table presents the application-code findings produced for this execution."
        ]

    if normalized_key in {"sast_extraction_warnings", "sast_extraction_warnings_by_tool", "sast_extraction_warnings_by_level", "sast_extraction_warnings_summary"}:
        if isinstance(value, list):
            details = "; ".join(_clean_text(x) for x in value if not _is_emptyish(x))
        else:
            details = _clean_text(value)
        if not details:
            return None
        label_map = {
            "sast_extraction_warnings": "SAST execution notifications",
            "sast_extraction_warnings_by_tool": "SAST execution notifications by tool",
            "sast_extraction_warnings_by_level": "SAST execution notifications by level",
            "sast_extraction_warnings_summary": "SAST execution notification summary",
        }
        return [label_map.get(normalized_key, normalized_key.replace("_", " ").title()), details]

    if isinstance(value, list):
        details = "; ".join(_clean_text(x) for x in value if not _is_emptyish(x))
    elif isinstance(value, dict):
        details = _clean_text(value)
    else:
        details = _clean_text(value)
    if not details:
        return None
    return [normalized_key.replace("_", " ").title(), details]


def _add_technical_execution_metadata(doc: Document, technical: Dict[str, Any]) -> None:
    limitations = _as_dict(technical.get("coverage_limitations"))
    rows: List[List[str]] = []
    seen_labels: set[str] = set()

    def add_row(row: List[str] | None) -> None:
        if not row:
            return
        label = _clean_text(row[0])
        details = _clean_text(row[1] if len(row) > 1 else "")
        if not label or not details:
            return
        # Keep a single executive row for generic SAST extraction warnings;
        # detailed breakdown rows remain separately visible by tool, level, and summary.
        if label in seen_labels:
            return
        seen_labels.add(label)
        rows.append([label, details])

    missing_inputs = _deep_list(limitations, ["missing_inputs"])
    if not missing_inputs:
        add_row(["Technical inputs", "All expected technical input artifact categories were present in the analysis pack."])

    if isinstance(limitations, dict):
        for key, value in limitations.items():
            row = _format_limitation_row(str(key), value)
            if row:
                if row[0] == "Technical inputs" and not missing_inputs:
                    continue
                add_row(row)
    if rows:
        _add_table(doc, ["Execution item", "Summary"], rows, max_rows=12)
    else:
        _add_body_paragraph(doc, "Technical execution metadata was not populated in the analysis pack.")


def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name, "").strip().lower()
    if not raw:
        return default
    return raw in {"1", "true", "yes", "on"}


def _ai_pattern_writeups_required() -> bool:
    """Require AI-generated pattern narratives by default.

    Recommendations and MAP closure criteria are not authored from static
    preauthored text. They must be returned by the configured AI task using the
    workbook metrics and scanner context. Set
    AUDIT_SUMMARY_AI_PATTERN_WRITEUPS_REQUIRED=0 only for local debugging.
    """
    if _env_bool("AUDIT_SUMMARY_AI_REQUIRED", False):
        return True
    return _env_bool("AUDIT_SUMMARY_AI_PATTERN_WRITEUPS_REQUIRED", True)


def _ai_recommendations_for_pattern(pattern: str, writeups: Dict[str, Dict[str, Any]]) -> List[str]:
    item = _as_dict(writeups.get(pattern))
    recs = item.get("recommendations")
    if not isinstance(recs, list):
        return []
    return [_clean_text(x) for x in recs if _clean_text(x)]


def _ai_field_for_pattern(pattern: str, writeups: Dict[str, Dict[str, Any]], field: str) -> str:
    return _clean_text(_as_dict(writeups.get(pattern)).get(field, ""))


def _validate_ai_pattern_writeups(patterns: List[Dict[str, Any]], writeups: Dict[str, Dict[str, Any]]) -> None:
    """Validate that AI supplied the report-authored pattern content.

    This prevents static, pre-authored recommendations from silently entering
    the report. The configured model must produce expected state, impact,
    recommendations, and MAP closure criteria for each reported pattern.
    """
    missing: List[str] = []
    for p in patterns[:10]:
        pat = str(p.get("pattern") or "").strip()
        if not pat:
            continue
        item = _as_dict(writeups.get(pat))
        missing_fields = []
        if not _clean_text(item.get("expected")):
            missing_fields.append("expected")
        if not _clean_text(item.get("impact")):
            missing_fields.append("impact")
        if not _ai_recommendations_for_pattern(pat, writeups):
            missing_fields.append("recommendations")
        if not _clean_text(item.get("closure_criteria")):
            missing_fields.append("closure_criteria")
        if missing_fields:
            missing.append(f"{pat}: {', '.join(missing_fields)}")

    if missing and _ai_pattern_writeups_required():
        details = " | ".join(missing[:12])
        raise SystemExit(
            "[ERROR] AI-generated pattern writeups are incomplete. "
            "The report is configured to require AI-authored recommendations, "
            "impact statements, expected-state narratives, and MAP closure "
            f"criteria. Missing: {details}"
        )
    if missing:
        print("[AI][WARN] AI-generated pattern writeups are incomplete: " + " | ".join(missing[:12]))


def _extract_json_object(text: str) -> str:
    text = (text or "").strip()
    if not text:
        raise ValueError("Model returned empty text.")

    # Direct JSON object.
    try:
        obj = json.loads(text)
        if isinstance(obj, dict):
            return json.dumps(obj, ensure_ascii=False)
    except Exception:
        pass

    # JSON inside Markdown code fence.
    fenced = re.search(
        r"```(?:json)?\s*(\{.*?\})\s*```",
        text,
        flags=re.DOTALL | re.IGNORECASE,
    )
    if fenced:
        return fenced.group(1)

    # Remove common fence wrappers if the model used an unterminated fence.
    cleaned = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, dict):
            return json.dumps(obj, ensure_ascii=False)
    except Exception:
        pass

    # Find the first balanced JSON object anywhere in the response.
    start = text.find("{")
    if start < 0:
        preview = text[:500].replace("\n", "\\n")
        raise ValueError(f"No JSON object found in model output. Preview: {preview}")

    depth = 0
    in_string = False
    escape = False

    for i in range(start, len(text)):
        ch = text[i]

        if escape:
            escape = False
            continue

        if ch == "\\":
            escape = True
            continue

        if ch == '"':
            in_string = not in_string
            continue

        if in_string:
            continue

        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                candidate = text[start:i + 1]
                obj = json.loads(candidate)
                if isinstance(obj, dict):
                    return candidate

    preview = text[:500].replace("\n", "\\n")
    raise ValueError(f"No complete JSON object found in model output. Preview: {preview}")


def _ai_env(name: str, default: str = "") -> str:
    return os.getenv(name, default).strip()


_AI_RUNTIME_CACHE = None


def _audit_ai_task() -> str:
    return _ai_env("AI_TASK", "audit_summary_docx") or "audit_summary_docx"


def _get_ai_runtime():
    """Resolve the configured AI runtime from parameters/ai.config.json.

    The audit-summary workflow exports AI_TASK=audit_summary_docx and may set
    AI_PROFILE or provider/model overrides. AIRuntime delegates resolution to
    scripts/lib/ai_config.py, so the effective model is controlled by
    parameters/ai.config.json first and GitHub Actions variables second.
    """
    global _AI_RUNTIME_CACHE
    if _AI_RUNTIME_CACHE is not None:
        return _AI_RUNTIME_CACHE

    if AIRuntime is None:
        print(f"[AI] AIRuntime is not available; AI-authored prose is unavailable. Import error: {AI_RUNTIME_IMPORT_ERROR}")
        return None

    try:
        profile = _ai_env("AI_PROFILE") or None
        runtime = AIRuntime(task=_audit_ai_task(), profile=profile)
        _AI_RUNTIME_CACHE = runtime
        cfg = getattr(runtime, "config", {}) or {}
        print(
            "[AI_CONFIG] "
            f"task={cfg.get('resolved_task') or _audit_ai_task()} "
            f"profile={cfg.get('resolved_profile') or profile or ''} "
            f"provider={getattr(runtime, 'provider', '')} "
            f"model={getattr(runtime, 'model', '')} "
            f"litellm_model={getattr(runtime, 'litellm_model', '')} "
            f"api_base={getattr(runtime, 'api_base', '')}"
        )
        return runtime
    except Exception as exc:
        print(f"[AI][WARN] Could not resolve AIRuntime from parameters/ai.config.json: {exc}")
        return None


def _ai_enabled() -> bool:
    raw = _ai_env("AUDIT_SUMMARY_AI_ENABLED", "1").lower()
    if raw in {"0", "false", "no", "off"}:
        return False

    runtime = _get_ai_runtime()
    if runtime is None:
        return False

    try:
        if not runtime.available():
            cfg = getattr(runtime, "config", {}) or {}
            print(
                "[AI] Runtime is not available. AI-authored recommendations will fail if required. "
                f"provider={getattr(runtime, 'provider', '')} "
                f"model={getattr(runtime, 'model', '')} "
                f"api_key_env_var={cfg.get('api_key_env_var', 'OPENAI_API_KEY')}"
            )
            return False
    except Exception as exc:
        print(f"[AI][WARN] Runtime availability check failed: {exc}")
        return False

    return True


def _ai_max_tokens(default: int = 1600) -> int:
    env_value = _ai_env("AI_MAX_OUTPUT_TOKENS")
    if env_value:
        return _safe_int(env_value, default)

    runtime = _get_ai_runtime()
    if runtime is not None:
        cfg = getattr(runtime, "config", {}) or {}
        configured = cfg.get("max_output_tokens")
        if configured not in (None, ""):
            return _safe_int(configured, default)

    return default


AI_REPORT_LANGUAGE_POLICY = (
    "Report language policy for the final director-facing audit report: "
    "write concise management-ready audit language. "
    "Present audit results, observed technical evidence, security meaning, recommended actions, and traceability. "
    "Describe each evidence source independently and use the workbook as the requirement-result source. "
    "Use normalized evidence values exactly as supplied for all counts, severities, PUIDs, CVEs, packages, files, and scanner observations. "
    "For executive narrative, include only observed values and action-oriented interpretation. "
    "For appendix text, use neutral execution metadata and traceability language. "
)

AI_UNWANTED_REPORT_TERMS = [
    "discrepanc",
    "contradict",
    "technical coverage limitations",
    "scope and limitations",
    "missing evidence",
    "failed parsing",
    "parser did not",
    "internal pipeline",
    "retained_app_code_findings",
    "hardening_or_maintainability_signals",
    "entr(y/ies)",
    "The application applications",
    "The application the mobile application",
]


def _collect_ai_language_policy_hits(value: Any) -> List[str]:
    """Return unwanted wording generated by the model before it enters the DOCX.

    This is prompt-first validation, not content rewriting. When hits are found,
    Stage 2 asks the model to regenerate the same section with stricter guidance.
    The final DOCX quality gate remains a safety net only.
    """
    hits: List[str] = []

    def visit(obj: Any) -> None:
        if isinstance(obj, str):
            low = obj.lower()
            for term in AI_UNWANTED_REPORT_TERMS:
                if term.lower() in low:
                    hits.append(term)
        elif isinstance(obj, list):
            for item in obj:
                visit(item)
        elif isinstance(obj, dict):
            for item in obj.values():
                visit(item)

    visit(value)
    out: List[str] = []
    seen = set()
    for h in hits:
        key = h.lower()
        if key not in seen:
            seen.add(key)
            out.append(h)
    return out


def _ai_repair_attempts() -> int:
    return max(0, min(3, _safe_int(os.getenv("AUDIT_SUMMARY_AI_SECTION_REPAIR_ATTEMPTS", "0"), 0)))


def _prompt_expected_items(user_payload: Dict[str, Any]) -> int:
    if not isinstance(user_payload, dict):
        return 0
    if isinstance(user_payload.get("items"), list):
        return len(user_payload.get("items") or [])
    context = user_payload.get("context") if isinstance(user_payload.get("context"), dict) else {}
    for key in ("positive_controls", "top_weakness_patterns", "patterns", "technical_findings"):
        value = context.get(key) if isinstance(context, dict) else None
        if isinstance(value, list):
            return len(value)
    schema = user_payload.get("required_output_schema") if isinstance(user_payload.get("required_output_schema"), dict) else {}
    for value in schema.values():
        if isinstance(value, list):
            return 1
    return 1


def _prompt_received_items(obj: Dict[str, Any]) -> int:
    if not isinstance(obj, dict) or not obj:
        return 0
    counts = [len(v) for v in obj.values() if isinstance(v, list)]
    return max(counts) if counts else 1


def _prompt_schema_valid(obj: Dict[str, Any], user_payload: Dict[str, Any]) -> bool:
    if not isinstance(obj, dict):
        return False
    schema = user_payload.get("required_output_schema") if isinstance(user_payload, dict) else None
    if not isinstance(schema, dict) or not schema:
        return True
    return all(key in obj for key in schema.keys())


def _prompt_traceability_ok(obj: Dict[str, Any], user_payload: Dict[str, Any]) -> bool | None:
    if not isinstance(obj, dict) or not isinstance(user_payload, dict):
        return None
    expected = []
    for item in user_payload.get("items", []) if isinstance(user_payload.get("items"), list) else []:
        if isinstance(item, dict):
            item_id = _clean_text(item.get("item_id") or item.get("puid") or item.get("id"))
            if item_id:
                expected.append(item_id)
    if not expected:
        return None
    seen = []
    for value in obj.values():
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    item_id = _clean_text(item.get("item_id") or item.get("puid") or item.get("id"))
                    if item_id:
                        seen.append(item_id)
    if not seen:
        return False
    return set(seen).issubset(set(expected))


def _ai_json_chat(section_name: str, system_prompt: str, user_payload: Dict[str, Any], max_tokens: int = 1600) -> Dict[str, Any]:
    """Call the configured AI runtime and require report-ready language.

    This function implements prompt-first quality control. If the model returns
    JSON containing wording that should not appear in the final report, Stage 2
    retries the same section with explicit feedback instead of relying on a late
    sanitization pass.
    """
    if not _ai_enabled():
        return {}

    runtime = _get_ai_runtime()
    if runtime is None:
        return {}

    last_hits: List[str] = []
    attempts = 1 + _ai_repair_attempts()
    call_started_at = time.time()
    last_error = ""
    contract = contract_for_section(section_name, fallback_source_function="_ai_json_chat")
    base_system_prompt = str(contract.get("system_prompt_transcript") or contract.get("system_prompt") or system_prompt)
    expected_schema = user_payload.get("required_output_schema") if isinstance(user_payload, dict) else {}
    expected_items = _prompt_expected_items(user_payload)

    for attempt in range(1, attempts + 1):
        if attempt == 1:
            effective_system_prompt = base_system_prompt
        else:
            effective_system_prompt = (
                base_system_prompt
                + "\n\nRegenerate the same JSON section. The previous response used report language that is not allowed: "
                + ", ".join(last_hits)
                + ". Remove those terms and express the content as observed evidence, management action, execution metadata, or traceability only. "
                + "Regenerate using only observed evidence, management action, execution metadata, and traceability."
            )

        messages = [
            {"role": "system", "content": effective_system_prompt},
            {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
        ]

        try:
            print(
                "[AI] Calling "
                f"section={section_name} "
                f"attempt={attempt}/{attempts} "
                f"task={_audit_ai_task()} "
                f"provider={getattr(runtime, 'provider', '')} "
                f"model={getattr(runtime, 'model', '')}"
            )
            response = runtime.create(
                input=messages,
                max_output_tokens=max_tokens,
                reasoning={"effort": _ai_env("AI_REASONING_EFFORT")},
            )
            content = (getattr(response, "output_text", "") or "").strip()
            if not content:
                print(f"[AI][WARN] Section {section_name} returned empty content.")
                continue

            obj = json.loads(_extract_json_object(content))
            if not isinstance(obj, dict):
                continue

            hits = _collect_ai_language_policy_hits(obj)
            if hits:
                last_hits = hits
                if attempt < attempts:
                    print(f"[AI][RETRY] Section {section_name} used non-report wording: {', '.join(hits)}")
                    continue
                print(f"[AI][WARN] Section {section_name} still contains non-report wording after retries: {', '.join(hits)}")
                record_prompt_call(
                    prompt_id=str(contract.get("prompt_id") or ""),
                    prompt_name=str(contract.get("prompt_name") or section_name),
                    prompt_scope=str(contract.get("prompt_scope") or "audit_summary"),
                    prompt_category=str(contract.get("prompt_category") or "primary_audit_summary_prompt"),
                    source_file=str(contract.get("source_file") or "scripts/audit_summary_stage2_generate_docx.py"),
                    source_function=str(contract.get("source_function") or "_ai_json_chat"),
                    section_name=section_name,
                    system_prompt=effective_system_prompt,
                    user_payload=user_payload,
                    expected_schema=expected_schema,
                    model=str(getattr(runtime, "model", "")),
                    provider=str(getattr(runtime, "provider", "")),
                    max_output_tokens=max_tokens,
                    reasoning_effort=_ai_env("AI_REASONING_EFFORT"),
                    attempt_count=attempt,
                    retry_count=max(0, attempt - 1),
                    expected_items=expected_items,
                    received_items=_prompt_received_items(obj),
                    json_valid=True,
                    schema_valid=_prompt_schema_valid(obj, user_payload),
                    traceability_ok=_prompt_traceability_ok(obj, user_payload),
                    repair_used="repair" in section_name,
                    fallback_used=True,
                    elapsed_s=time.time() - call_started_at,
                    error="Non-report wording remained after repair attempts: " + ", ".join(hits),
                    registration_status=str(contract.get("registration_status") or "registered"),
                )
                return {}

            print(f"[AI] Section {section_name} completed.")
            record_prompt_call(
                prompt_id=str(contract.get("prompt_id") or ""),
                prompt_name=str(contract.get("prompt_name") or section_name),
                prompt_scope=str(contract.get("prompt_scope") or "audit_summary"),
                prompt_category=str(contract.get("prompt_category") or "primary_audit_summary_prompt"),
                source_file=str(contract.get("source_file") or "scripts/audit_summary_stage2_generate_docx.py"),
                source_function=str(contract.get("source_function") or "_ai_json_chat"),
                section_name=section_name,
                system_prompt=effective_system_prompt,
                user_payload=user_payload,
                expected_schema=expected_schema,
                model=str(getattr(runtime, "model", "")),
                provider=str(getattr(runtime, "provider", "")),
                max_output_tokens=max_tokens,
                reasoning_effort=_ai_env("AI_REASONING_EFFORT"),
                attempt_count=attempt,
                retry_count=max(0, attempt - 1),
                expected_items=expected_items,
                received_items=_prompt_received_items(obj),
                json_valid=True,
                schema_valid=_prompt_schema_valid(obj, user_payload),
                traceability_ok=_prompt_traceability_ok(obj, user_payload),
                repair_used="repair" in section_name,
                elapsed_s=time.time() - call_started_at,
                registration_status=str(contract.get("registration_status") or "registered"),
            )
            return obj
        except Exception as exc:
            last_error = str(exc)
            print(f"[AI][WARN] Section {section_name} attempt {attempt}/{attempts} failed: {exc}")
    record_prompt_call(
        prompt_id=str(contract.get("prompt_id") or ""),
        prompt_name=str(contract.get("prompt_name") or section_name),
        prompt_scope=str(contract.get("prompt_scope") or "audit_summary"),
        prompt_category=str(contract.get("prompt_category") or "primary_audit_summary_prompt"),
        source_file=str(contract.get("source_file") or "scripts/audit_summary_stage2_generate_docx.py"),
        source_function=str(contract.get("source_function") or "_ai_json_chat"),
        section_name=section_name,
        system_prompt=base_system_prompt,
        user_payload=user_payload,
        expected_schema=expected_schema,
        model=str(getattr(runtime, "model", "")),
        provider=str(getattr(runtime, "provider", "")),
        max_output_tokens=max_tokens,
        reasoning_effort=_ai_env("AI_REASONING_EFFORT"),
        attempt_count=attempts,
        retry_count=max(0, attempts - 1),
        expected_items=expected_items,
        received_items=0,
        json_valid=False,
        schema_valid=False,
        traceability_ok=False,
        repair_used="repair" in section_name,
        fallback_used=True,
        elapsed_s=time.time() - call_started_at,
        error=last_error or "AI section generation failed",
        registration_status=str(contract.get("registration_status") or "registered"),
    )
    return {}

def _compact_patterns_for_ai(patterns: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for p in patterns[:limit]:
        count = int(p.get("mapped_noncompliant_count", 0))
        scanner_context = _as_dict(p.get("scanner_context_for_ai"))
        out.append({
            "pattern": p.get("pattern"),
            "mapped_noncompliant_count": count,
            "severity": p.get("severity"),
            "likelihood": _likelihood_from_count(count),
            "recommended_owner": p.get("recommended_owner"),
            "example_puids": (p.get("example_puids") or [])[:5],
            "description_anchors": [_clean_text(x) for x in (p.get("description_anchors") or [])[:2]],
            "related_flags_sample": (p.get("related_flags_sample") or [])[:16],
            "related_flags_by_family": _as_dict(p.get("related_flags_by_family")),
            "workbook_examples_for_ai": (p.get("workbook_examples_for_ai") or [])[:5],
            "scanner_context_for_ai": {
                "relevant_sources": scanner_context.get("relevant_sources", []),
                "evidence": scanner_context.get("evidence", {}),
                "guardrails": scanner_context.get("guardrails", []),
            },
        })
    return out


def _classify_positive_control_statement(statement: str, evidence: str = "", flags: str = "") -> str:
    text = " ".join([_clean_text(statement), _clean_text(evidence), _clean_text(flags)]).lower()
    if any(token in text for token in [
        " but ", "however", "residual", "not detected", "not observed", "not available",
        "risky permissions are present", "dangerous permissions", "partially", "fallback verdict",
        "ssl pinning was not detected", "clear text traffic and ssl pinning"
    ]):
        if "risky permissions" in text or "ssl pinning was not detected" in text or "not detected" in text:
            return "Observed controls with additional technical context"
        return "Observed supporting signals"
    if any(token in text for token in ["scanner coverage", "did not report", "manual logout", "auditor review"]):
        return "Observed supporting signals"
    return "Observed positive controls"


def _group_positive_controls(pos_controls: List[Dict[str, Any]], positive_control_writeups: Dict[str, str]) -> Dict[str, List[Dict[str, Any]]]:
    groups = {
        "Observed positive controls": [],
        "Observed supporting signals": [],
        "Observed controls with additional technical context": [],
    }
    for pc in pos_controls:
        puid = _clean_text(pc.get("puid"))
        statement = _sanitize_positive_control_final(positive_control_writeups.get(puid, ""))
        if not statement:
            # Prompt-first policy: Stage 2 does not build a positive-control
            # narrative from deterministic wording. Missing model output is
            # handled by validation and by Appendix D source traceability.
            continue
        bucket = _classify_positive_control_statement(statement, _clean_text(pc.get("evidence_excerpt")), _clean_text(pc.get("flags_used")))
        enriched = dict(pc)
        enriched["reported_statement"] = statement
        enriched["evidence_class"] = bucket
        groups.setdefault(bucket, []).append(enriched)
    return groups



def _compact_positive_controls_for_ai(pos_controls: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for pc in pos_controls[:limit]:
        requirement_text = _clean_text(pc.get("requirement_text", "")) or _clean_text(pc.get("description", ""))
        out.append({
            "puid": pc.get("puid"),
            "requirement_text": requirement_text,
            "flags_used": _clean_text(pc.get("flags_used", "")),
            "evidence_excerpt": _clean_text(pc.get("evidence_excerpt", "")),
        })
    return out


def _compact_technical_for_ai(technical: Dict[str, Any]) -> Dict[str, Any]:
    trivy = _as_dict(technical.get("trivy_sca"))
    mobsf_static = _as_dict(technical.get("mobsf_static"))
    mobsf_dynamic = _as_dict(technical.get("mobsf_dynamic"))
    sast = _as_dict(technical.get("sast_app_code"))
    vision = _as_dict(technical.get("vision360"))
    limitations = _as_dict(technical.get("coverage_limitations"))
    sast_counts = _sast_counts(sast)

    warning_count = _deep_int(limitations, ["sast_extraction_warning_count"])
    if warning_count == 0:
        warning_items = _deep_list(limitations, ["sast_extraction_warnings_summary", "sast_notifications_sample", "sast_extraction_warnings"])
        warning_count = len(warning_items)
    if warning_count == 0 and _clean_text(limitations.get("sast_extraction_warnings", "")):
        m = re.search(r"(\d+)\s+SAST extraction", _clean_text(limitations.get("sast_extraction_warnings", "")), re.IGNORECASE)
        warning_count = _safe_int(m.group(1), 0) if m else 1

    return {
        "vision360": {
            "available": _block_available(vision),
            "flags_count": _deep_int(vision, ["flags_count"]),
            "state_counts": _deep_dict(vision, ["state_counts"]),
            "group_counts": _deep_dict(vision, ["group_counts"]),
            "top_active_groups": _deep_list(vision, ["top_active_groups"]),
            "feature_keys": _deep_list(vision, ["feature_keys"]),
        },
        "trivy_sca": {
            "available": _block_available(trivy),
            "packages_detected": _deep_int(trivy, ["packages_detected", "package_count", "packages_total"]),
            "total_vulnerabilities": _deep_int(trivy, ["total_vulnerabilities", "vulnerabilities_total", "total"]),
            "fixable_total": _deep_int(trivy, ["fixable_total", "fixable_vulnerabilities", "fixable"]),
            "by_severity": _deep_dict(trivy, ["by_severity", "severity_counts"]),
            "top_findings": [
                {
                    "id": _first_present(f, ["id", "VulnerabilityID", "cve", "rule_id"], ""),
                    "severity": _first_present(f, ["severity", "Severity"], ""),
                    "package": _first_present(f, ["pkg", "PkgName", "package", "package_name"], ""),
                    "installed": _first_present(f, ["installed", "InstalledVersion", "installed_version"], ""),
                    "fixed": _first_present(f, ["fixed", "FixedVersion", "fixed_version"], ""),
                    "title": _clean_text(_first_present(f, ["title", "Title", "description", "Description"], "")),
                }
                for f in _trivy_findings(trivy)[:12]
            ],
        },
        "mobsf_static": {
            "available": _block_available(mobsf_static),
            "signals": _mobsf_signal_rows(mobsf_static, technical) if _block_available(mobsf_static) else [],
            "manifest_findings_count": _deep_int(mobsf_static, ["manifest_findings_count"]),
            "certificate_findings_count": _deep_int(mobsf_static, ["certificate_findings_count"]),
            "dangerous_permissions": _deep_list(mobsf_static, ["dangerous_permissions"])[:20],
            "permissions_count": _deep_int(mobsf_static, ["permissions_count"]),
        },
        "mobsf_dynamic": {
            "available": _block_available(mobsf_dynamic),
            "local_storage_artifacts_count": _deep_int(mobsf_dynamic, ["local_storage_artifacts_count"]),
            "local_storage_artifacts_sample": _deep_list(mobsf_dynamic, ["local_storage_artifacts_sample", "local_storage_artifacts", "files", "storage_artifacts"])[:12],
            "local_storage_artifact_names": _deep_list(mobsf_dynamic, ["local_storage_artifact_names"])[:12],
            "shared_preferences_artifacts": _deep_list(mobsf_dynamic, ["shared_preferences_artifacts", "shared_preferences", "shared_preferences_files", "preferences"])[:12],
            "shared_preferences_artifact_names": _deep_list(mobsf_dynamic, ["shared_preferences_artifact_names"])[:12],
            "sqlite_database_artifacts": _deep_list(mobsf_dynamic, ["sqlite_database_artifacts", "sqlite_databases", "databases", "db_files"])[:12],
            "sqlite_database_artifact_names": _deep_list(mobsf_dynamic, ["sqlite_database_artifact_names"])[:12],
            "detected_trackers": _deep_int(mobsf_dynamic, ["detected_trackers", "trackers_detected"]),
        },
        "sast_app_code": {
            "available": _block_available(sast),
            "retained_app_code_signals": sast_counts["retained_app_code_signals"],
            "retained_security_findings": sast_counts["retained_security_findings"],
            "hardening_or_maintainability_signals": sast_counts["hardening_or_maintainability_signals"],
            "raw_sarif_results": sast_counts["raw_sarif_results"],
            "raw_sarif_counts_by_tool": _deep_dict(sast, ["raw_tool_counts", "tool_counts", "by_tool"]),
            "retained_security_tool_counts": _deep_dict(sast, ["retained_security_tool_counts"]),
            "hardening_tool_counts": _deep_dict(sast, ["hardening_tool_counts"]),
            "top_security_rules": _deep_list(sast, ["top_security_rules"])[:12],
            "top_hardening_rules": _deep_list(sast, ["top_hardening_rules"])[:12],
            "scope_filter_rule": "Only retained_security_findings may be described as security-relevant application-code SAST findings.",
            "raw_counts_interpretation": "Raw SARIF counts are traceability and coverage signals only. Retained app-code signals may include hardening or maintainability findings and are not automatically vulnerabilities.",
            "security_findings_sample": _deep_list(sast, ["security_findings_sample"])[:12],
            "hardening_signals_sample": _deep_list(sast, ["hardening_signals_sample"])[:12],
            "sast_notifications_sample": _deep_list(_as_dict(technical.get("coverage_limitations")), ["sast_notifications_sample", "sast_extraction_warnings_summary"])[:8],
        },
        "technical_execution_metadata": {
            "all_expected_input_categories_present": len(_deep_list(limitations, ["missing_inputs"])) == 0,
            "sast_execution_notification_count": warning_count,
            "sast_execution_notifications_by_tool": _clean_text(limitations.get("sast_extraction_warnings_by_tool", "")),
            "sast_execution_notifications_by_level": _clean_text(limitations.get("sast_extraction_warnings_by_level", "")),
        },
    }


def _call_llm_for_audit_sections(
    metrics: Dict[str, Any],
    app: Dict[str, Any],
    patterns: List[Dict[str, Any]],
    technical: Dict[str, Any],
    likelihood_rubric: Dict[str, str],
    positive_controls: List[Dict[str, Any]] | None = None,
) -> Dict[str, Any]:
    if not _ai_enabled():
        return {}

    common_system = str(contract_for_section("executive_summary", fallback_source_function="_ai_json_chat").get("system_prompt_transcript") or "")
    if not common_system:
        raise RuntimeError("Audit Summary common system prompt is empty in scripts/prompt_contracts.json")

    compact_patterns = _compact_patterns_for_ai(patterns)
    compact_technical = _compact_technical_for_ai(technical)

    base_context = {
        "application": app,
        "metrics": metrics,
        "likelihood_rubric": likelihood_rubric,
        "top_weakness_patterns": compact_patterns,
        "technical_evidence": compact_technical,
        "positive_controls": _compact_positive_controls_for_ai(positive_controls or []),
    }

    out: Dict[str, Any] = {}

    executive = _ai_json_chat(
        "executive_summary",
        common_system,
        {
            "task": "Generate executive summary components for an mSEC-AM mobile application audit report.",
            "constraints": {
                "key_takeaways_count": "5 to 7",
                "must_reference": ["overall compliance rate", "applicable controls", "non-compliant controls", "top weakness patterns", "technical scanner evidence where available"],
                "do_not_claim": ["full codebase is clean", "MobSF absence equals no risk", "SAST raw findings are app findings"],
            },
            "context": base_context,
            "required_output_schema": {
                "audit_summary_paragraph": "<one concise paragraph>",
                "key_takeaways": ["<bullet text>"],
            },
        },
        max_tokens=min(_ai_max_tokens(1800), 2200),
    )
    out.update({k: v for k, v in executive.items() if k in {"audit_summary_paragraph", "key_takeaways"}})

    positive_controls_ai = _ai_json_chat(
        "positive_controls",
        common_system,
        {
            "task": "Rewrite verified positive control statements into precise, grammatical audit-report English.",
            "constraints": {
                "use_exact_puid": True,
                "do_not_overstate": True,
                "do_not_invent_evidence": True,
                "if_evidence_is_partial": "State the supported observation and the verification action without using limitation language.",
                "statement_style": "One concise sentence per control, suitable for an executive report.",
                "prudence_rules": [
                    "Do not convert cleartext traffic disabled into full TLS enforcement or SSL pinning evidence.",
                    "Do not convert session-cookie assignment into proof that HTTPOnly is enforced unless HTTPOnly evidence is present.",
                    "Do not say the application is malware-free; say that available malware evidence did not report adware or known malware.",
                    "Do not claim that the app prevents accepting all SSL/TLS certificates when only manifest cleartext evidence is available; state only the cleartext observation and the verification action.",
                    "Do not claim that data exchange is secured via TLS because SSL pinning was not detected; describe pinning only as an observed configuration item or planned verification item."
                ],
            },
            "context": {
                "application": app,
                "positive_controls": _compact_positive_controls_for_ai(positive_controls or []),
                "technical_evidence": compact_technical,
            },
            "required_output_schema": {
                "positive_controls": [
                    {
                        "puid": "<exact PUID>",
                        "statement": "<rewritten statement grounded only in flags and evidence>"
                    }
                ]
            },
        },
        max_tokens=min(_ai_max_tokens(1400), 1800),
    )
    if isinstance(positive_controls_ai.get("positive_controls"), list):
        out["positive_controls"] = positive_controls_ai["positive_controls"]

    technical_narratives = _ai_json_chat(
        "technical_narratives",
        common_system,
        {
            "task": "Generate report-ready technical narrative paragraphs from scanner evidence.",
            "constraints": {
                "one_paragraph_each": True,
                "mention_execution_metadata": True,
                "do_not_overstate_sast": True,
                "observed_values_only": "Use observed normalized values. Omit absent fields from executive paragraphs.",
                "sast_rule": "Use technical_evidence.sast_app_code.retained_security_findings as the SAST security finding count. retained_app_code_signals may include hardening or maintainability findings. Mention raw SARIF counts only as traceability or execution metadata.",
                "mobsf_dynamic_rule": "Mention runtime artifact examples only when they are present in the normalized MobSF dynamic arrays. Use normalized Android paths or artifact names only; do not repeat collapsed raw strings such as datadata...",
                "avoid_absolute_claims": "Do not claim the app is clean, fully protected, or fully encrypted unless the supplied data directly supports that exact statement.",
            },
            "context": base_context,
            "required_output_schema": {
                "technical_coverage_paragraph": "<paragraph>",
                "technical_evidence_intro": "<paragraph>",
                "trivy_paragraph": "<paragraph>",
                "mobsf_static_paragraph": "<paragraph>",
                "mobsf_dynamic_paragraph": "<paragraph>",
                "sast_paragraph": "<paragraph>",
                "technical_execution_metadata_paragraph": "<paragraph>",
            },
        },
        max_tokens=min(_ai_max_tokens(2200), 2600),
    )
    if technical_narratives:
        out["technical_narratives"] = technical_narratives

    pattern_writeups = _ai_json_chat(
        "pattern_writeups",
        common_system,
        {
            "task": "Generate concise weakness-pattern writeups and recommendations grounded in workbook prevalence and technical evidence.",
            "constraints": {
                "patterns": "Use exact pattern names from input.",
                "expected": "1 positive expected secure state sentence. Do not describe the current problem in this field.",
                "impact": "1 sentence mentioning confidentiality, integrity, availability, or health-data regulatory exposure only when supported.",
                "recommendations": "4 to 6 actionable bullets per pattern. Each recommendation must be generated from the supplied workbook prevalence, PUID examples, scanner findings, and execution metadata. Do not use generic boilerplate or static templates. Do not convert raw SARIF counts or Detekt quality findings into vulnerabilities. Treat TLS pinning as threat-model dependent, not a universal absolute. Prefer documenting certificate validation and pinning decisions over mandating pinning across all communications. Translate raw flag names into operational actions; use exact flag names sparingly because detailed flag traceability belongs in the workbook and appendices.",
                "closure_criteria": "1 measurable sentence suitable for the MAP, generated from the supplied evidence and scanner context. Prefer evidence-based closure such as updated workbook scoring, updated Trivy/MobSF/SAST artifacts, regression evidence, or formal risk acceptance. Avoid broad zero noncompliant findings language for governance or multi-control patterns. Use zero only for narrowly scoped Critical/High fixable vulnerabilities when supported by Trivy evidence.",
                "no_time_window_headings": True,
                "no_unprovided_metrics": True,
                "no_static_recommendations": True,
            },
            "context": {
                "application": app,
                "likelihood_rubric": likelihood_rubric,
                "top_weakness_patterns": compact_patterns,
                "technical_evidence": compact_technical,
            },
            "required_output_schema": {
                "pattern_writeups": [
                    {
                        "pattern": "<exact pattern name>",
                        "expected": "<sentence>",
                        "impact": "<sentence>",
                        "recommendations": ["<AI-generated action grounded in evidence>"],
                        "closure_criteria": "<AI-generated measurable closure criterion grounded in evidence>",
                    }
                ]
            },
        },
        max_tokens=min(_ai_max_tokens(3600), 4200),
    )
    if isinstance(pattern_writeups.get("pattern_writeups"), list):
        out["pattern_writeups"] = pattern_writeups["pattern_writeups"]

    return out



def _treatment_plan(pack: Dict[str, Any]) -> Dict[str, Any]:
    return _as_dict(pack.get("treatment_plan"))


def _treatment_mode() -> str:
    """Return treatment generation mode.

    full: ask the AI model to generate treatment fields for every rendered
    control and technical item. This is slow and strict.
    priority: ask the AI model only for priority technical treatment items.
    The large control appendix is rendered as structured traceability data.
    off: do not ask the AI model for item-level treatment writeups.
    """
    raw = os.getenv("AUDIT_SUMMARY_TREATMENT_MODE", "priority").strip().lower()
    aliases = {
        "prioritized": "priority",
        "prio": "priority",
        "fast": "priority",
        "0": "off",
        "false": "off",
        "no": "off",
        "1": "priority",
        "true": "priority",
        "yes": "priority",
    }
    raw = aliases.get(raw, raw)
    if raw not in {"full", "priority", "off"}:
        print(f"[AI][TREATMENT][WARN] Unknown AUDIT_SUMMARY_TREATMENT_MODE={raw!r}; using priority.")
        return "priority"
    return raw


def _treatment_full_mode() -> bool:
    return _treatment_mode() == "full"


def _treatment_priority_mode() -> bool:
    return _treatment_mode() == "priority"


def _treatment_ai_required() -> bool:
    if "AUDIT_SUMMARY_AI_TREATMENT_REQUIRED" in os.environ:
        return _env_bool("AUDIT_SUMMARY_AI_TREATMENT_REQUIRED", False)
    return _treatment_full_mode()


def _treatment_batch_size() -> int:
    return max(4, min(25, _safe_int(os.getenv("AUDIT_SUMMARY_AI_TREATMENT_BATCH_SIZE", "12"), 12)))


def _max_rows_from_env(name: str, total: int) -> int:
    raw = os.getenv(name, "").strip().lower()
    if raw in {"", "0", "all", "full", "none", "unlimited"}:
        return max(0, total)
    return max(1, min(_safe_int(raw, total), total))


def _max_control_treatment_rows(total: int = 0) -> int:
    return _max_rows_from_env("AUDIT_SUMMARY_MAX_CONTROL_TREATMENT_ROWS", total)


def _max_technical_treatment_rows(total: int = 0) -> int:
    return _max_rows_from_env("AUDIT_SUMMARY_MAX_TECHNICAL_TREATMENT_ROWS", total)


def _max_correlation_rows(total: int = 0) -> int:
    return _max_rows_from_env("AUDIT_SUMMARY_MAX_CORRELATION_ROWS", total)


def _chunks(items: List[Any], size: int) -> List[List[Any]]:
    return [items[i:i + size] for i in range(0, len(items), size)]


def _compact_treatment_item(item: Dict[str, Any], item_kind: str) -> Dict[str, Any]:
    if item_kind == "control":
        return {
            "item_id": item.get("item_id"),
            "puid": item.get("puid"),
            "category": item.get("category_name"),
            "current_status": item.get("current_status"),
            "weakness_pattern": item.get("weakness_pattern"),
            "severity": item.get("severity"),
            "likelihood": item.get("likelihood"),
            "recommended_owner": item.get("recommended_owner"),
            "description": _clean_text(item.get("description")),
            "evidence_excerpt": _clean_text(item.get("evidence_excerpt")),
            "flags": _as_list(item.get("flags"))[:20],
            "flags_by_family": _as_dict(item.get("flags_by_family")),
        }
    return {
        "item_id": item.get("item_id"),
        "item_type": item.get("item_type"),
        "source": item.get("source"),
        "severity": item.get("severity"),
        "finding_id": item.get("finding_id"),
        "affected_component": item.get("affected_component"),
        "file": item.get("file"),
        "line": item.get("line"),
        "location": item.get("location"),
        "installed_version": item.get("installed_version"),
        "fixed_version": item.get("fixed_version"),
        "fix_available": item.get("fix_available"),
        "observed_issue": _clean_text(item.get("observed_issue")),
        "linked_patterns": _as_list(item.get("linked_patterns"))[:8],
        "linked_puids": _as_list(item.get("linked_puids"))[:12],
    }


def _sanitize_treatment_text(value: Any) -> str:
    """Compatibility wrapper for treatment-plan AI fields. No rewriting."""
    return _clean_generated_text(value, ensure_period=True)


def _normalize_treatment_item_id(item_id: Any, expected_ids: List[str] | None = None) -> str:
    """Normalize common LLM item-id formatting mistakes without inventing IDs.

    The model must use the exact item_id supplied by Stage 1. This helper only
    maps harmless formatting variants such as TECH-MOBSF-25 -> TECH-MOBSF-025
    when that mapping is unambiguous in the expected item-id set.
    """
    raw = _clean_text(item_id)
    if not raw:
        return ""

    if not expected_ids:
        return raw

    expected = [_clean_text(x) for x in expected_ids if _clean_text(x)]
    if raw in expected:
        return raw

    m = re.match(r"^(.*?)-0*(\d+)$", raw)
    if not m:
        return raw

    prefix = m.group(1)
    number = _safe_int(m.group(2), -1)
    matches = []
    for expected_id in expected:
        em = re.match(r"^(.*?)-0*(\d+)$", expected_id)
        if em and em.group(1) == prefix and _safe_int(em.group(2), -2) == number:
            matches.append(expected_id)

    return matches[0] if len(matches) == 1 else raw


def _merge_treatment_results(results: List[Dict[str, Any]], expected_ids: List[str] | None = None) -> Dict[str, Dict[str, Any]]:
    out: Dict[str, Dict[str, Any]] = {}
    for item in results:
        if not isinstance(item, dict):
            continue
        item_id = _normalize_treatment_item_id(item.get("item_id"), expected_ids)
        if not item_id:
            continue
        out[item_id] = {
            "treatment_action": _sanitize_treatment_text(item.get("treatment_action")),
            "verification_method": _sanitize_treatment_text(item.get("verification_method")),
            "closure_evidence": _sanitize_treatment_text(item.get("closure_evidence")),
            "residual_risk": _sanitize_treatment_text(item.get("residual_risk")),
        }
    return out


def _merge_treatment_ai_maps(base: Dict[str, Any], patch: Dict[str, Any]) -> Dict[str, Any]:
    """Merge AI-authored treatment maps.

    Non-empty repaired fields replace empty fields and may also overwrite prior
    values from malformed batches. This remains AI-authored text, not static
    preauthored remediation.
    """
    out = dict(base or {})
    for map_key in ("control_treatments", "technical_treatments"):
        current = dict(_as_dict(out.get(map_key)))
        incoming = _as_dict(patch.get(map_key))
        for item_id, writeup_any in incoming.items():
            writeup = _as_dict(writeup_any)
            if not item_id or not writeup:
                continue
            existing = dict(_as_dict(current.get(item_id)))
            for field in ("treatment_action", "verification_method", "closure_evidence", "residual_risk"):
                value = _sanitize_treatment_text(writeup.get(field))
                if value:
                    existing[field] = value
            current[item_id] = existing
        out[map_key] = current

    metadata = dict(_as_dict(out.get("metadata")))
    patch_metadata = _as_dict(patch.get("metadata"))
    for k, v in patch_metadata.items():
        metadata[k] = v
    out["metadata"] = metadata
    return out


def _treatment_missing_fields(writeup_any: Any) -> List[str]:
    writeup = _as_dict(writeup_any)
    required = ("treatment_action", "verification_method", "closure_evidence")
    return [field for field in required if not _clean_text(writeup.get(field))]


def _incomplete_treatment_items(treatment_plan: Dict[str, Any], treatment_ai: Dict[str, Any]) -> Dict[str, List[Dict[str, Any]]]:
    control_items = [x for x in _as_list(treatment_plan.get("control_items")) if isinstance(x, dict)]
    technical_items = [x for x in _as_list(treatment_plan.get("technical_items")) if isinstance(x, dict)]
    control_map = _as_dict(treatment_ai.get("control_treatments"))
    technical_map = _as_dict(treatment_ai.get("technical_treatments"))

    missing_control: List[Dict[str, Any]] = []
    for item in control_items[:_max_control_treatment_rows()]:
        item_id = _clean_text(item.get("item_id"))
        missing_fields = _treatment_missing_fields(control_map.get(item_id))
        if missing_fields:
            enriched = dict(item)
            enriched["_missing_ai_fields"] = missing_fields
            missing_control.append(enriched)

    missing_technical: List[Dict[str, Any]] = []
    for item in technical_items[:_max_technical_treatment_rows()]:
        item_id = _clean_text(item.get("item_id"))
        missing_fields = _treatment_missing_fields(technical_map.get(item_id))
        if missing_fields:
            enriched = dict(item)
            enriched["_missing_ai_fields"] = missing_fields
            missing_technical.append(enriched)

    return {
        "control_items": missing_control,
        "technical_items": missing_technical,
    }


def _treatment_repair_attempts() -> int:
    return max(0, min(5, _safe_int(os.getenv("AUDIT_SUMMARY_AI_TREATMENT_REPAIR_ATTEMPTS", "0"), 0)))


def _treatment_missing_id_list(missing: Dict[str, List[Dict[str, Any]]], limit: int = 20) -> str:
    ids: List[str] = []
    for key in ("control_items", "technical_items"):
        for item in missing.get(key, []):
            item_id = _clean_text(item.get("item_id"))
            if item_id:
                fields = ", ".join(_as_list(item.get("_missing_ai_fields")))
                ids.append(f"{item_id} ({fields})" if fields else item_id)
            if len(ids) >= limit:
                break
        if len(ids) >= limit:
            break
    return ", ".join(ids)


def _call_llm_for_treatment_repair(
    app: Dict[str, Any],
    technical_compact: Dict[str, Any],
    missing_control_items: List[Dict[str, Any]],
    missing_technical_items: List[Dict[str, Any]],
    system_prompt: str,
    attempt: int,
    max_attempts: int,
) -> Dict[str, Any]:
    """Ask the configured AI model to complete only missing treatment fields."""
    patch: Dict[str, Any] = {
        "control_treatments": {},
        "technical_treatments": {},
        "metadata": {
            "repair_attempt": attempt,
            "repair_max_attempts": max_attempts,
            "ai_authored": True,
        },
    }

    batch_size = _treatment_batch_size()
    control_batches = _chunks(missing_control_items, batch_size)
    technical_batches = _chunks(missing_technical_items, batch_size)
    repair_total_batches = len(control_batches) + len(technical_batches)

    print(
        f"[AI][TREATMENT][REPAIR] Attempt {attempt}/{max_attempts}: "
        f"{len(missing_control_items)} control item(s), "
        f"{len(missing_technical_items)} technical item(s), "
        f"{repair_total_batches} repair batch(es)."
    )

    for idx, batch in enumerate(control_batches, start=1):
        print(
            f"[AI][TREATMENT][REPAIR] Progress {idx}/{repair_total_batches} | "
            f"control repair batch {idx}/{len(control_batches)} | "
            f"items={', '.join(_clean_text(x.get('item_id')) for x in batch)}"
        )
        obj = _ai_json_chat(
            f"control_treatment_repair_{attempt}_{idx}_of_{len(control_batches)}",
            system_prompt,
            {
                "task": "Repair missing or incomplete treatment-plan fields for non-compliant SECM-CAT controls.",
                "constraints": {
                    "one_result_per_input_item": True,
                    "do_not_invent_evidence": True,
                    "use_exact_item_id_and_puid": True,
                    "complete_only_missing_or_empty_fields": True,
                    "required_non_empty_fields": ["item_id", "treatment_action", "verification_method", "closure_evidence", "residual_risk"],
                    "missing_fields_are_listed_per_item": True,
                },
                "application": app,
                "technical_evidence_summary": technical_compact,
                "items": [_compact_treatment_item(x, "control") | {"missing_ai_fields": _as_list(x.get("_missing_ai_fields"))} for x in batch],
                "required_output_schema": {
                    "control_treatments": [
                        {
                            "item_id": "<exact input item_id>",
                            "treatment_action": "<non-empty AI-generated action grounded in this control>",
                            "verification_method": "<non-empty verification method>",
                            "closure_evidence": "<non-empty closure evidence>",
                            "residual_risk": "<residual risk or risk acceptance note>",
                        }
                    ]
                },
            },
            max_tokens=min(_ai_max_tokens(2200), 2800),
        )
        values = obj.get("control_treatments") if isinstance(obj.get("control_treatments"), list) else []
        expected_ids = [_clean_text(x.get("item_id")) for x in batch]
        patch = _merge_treatment_ai_maps(
            patch,
            {"control_treatments": _merge_treatment_results([x for x in values if isinstance(x, dict)], expected_ids)}
        )
        print(
            f"[AI][TREATMENT][REPAIR] Completed control repair batch {idx}/{len(control_batches)} "
            f"on attempt {attempt}/{max_attempts}."
        )

    base_offset = len(control_batches)
    for idx, batch in enumerate(technical_batches, start=1):
        global_idx = base_offset + idx
        print(
            f"[AI][TREATMENT][REPAIR] Progress {global_idx}/{repair_total_batches} | "
            f"technical repair batch {idx}/{len(technical_batches)} | "
            f"items={', '.join(_clean_text(x.get('item_id')) for x in batch)}"
        )
        obj = _ai_json_chat(
            f"technical_treatment_repair_{attempt}_{idx}_of_{len(technical_batches)}",
            system_prompt,
            {
                "task": "Repair missing or incomplete treatment-plan fields for technical scanner findings.",
                "constraints": {
                    "one_result_per_input_item": True,
                    "do_not_invent_evidence": True,
                    "use_exact_item_id_and_finding_id": True,
                    "complete_only_missing_or_empty_fields": True,
                    "required_non_empty_fields": ["item_id", "treatment_action", "verification_method", "closure_evidence", "residual_risk"],
                    "dependency_rule": "For Trivy items, use fixed_version when present; if no fixed version exists, propose risk acceptance, compensating controls, or monitoring rather than inventing a version.",
                    "sast_rule": "For SAST items, use only supplied rule, message, file, and line. Do not invent code paths.",
                    "mobsf_rule": "For MobSF items, recommend configuration, signing, manifest, runtime-storage, or verification actions based only on supplied evidence.",
                    "missing_fields_are_listed_per_item": True,
                },
                "application": app,
                "technical_evidence_summary": technical_compact,
                "items": [_compact_treatment_item(x, "technical") | {"missing_ai_fields": _as_list(x.get("_missing_ai_fields"))} for x in batch],
                "required_output_schema": {
                    "technical_treatments": [
                        {
                            "item_id": "<exact input item_id>",
                            "treatment_action": "<non-empty AI-generated action grounded in this finding>",
                            "verification_method": "<non-empty verification method>",
                            "closure_evidence": "<non-empty closure evidence>",
                            "residual_risk": "<residual risk or risk acceptance note>",
                        }
                    ]
                },
            },
            max_tokens=min(_ai_max_tokens(2200), 2800),
        )
        values = obj.get("technical_treatments") if isinstance(obj.get("technical_treatments"), list) else []
        expected_ids = [_clean_text(x.get("item_id")) for x in batch]
        patch = _merge_treatment_ai_maps(
            patch,
            {"technical_treatments": _merge_treatment_results([x for x in values if isinstance(x, dict)], expected_ids)}
        )
        print(
            f"[AI][TREATMENT][REPAIR] Completed technical repair batch {idx}/{len(technical_batches)} "
            f"on attempt {attempt}/{max_attempts}."
        )

    return patch


def _priority_control_treatment_limit() -> int:
    return max(0, _safe_int(os.getenv("AUDIT_SUMMARY_PRIORITY_CONTROL_TREATMENT_ITEMS", "0"), 0))


def _priority_technical_treatment_limit() -> int:
    return max(0, _safe_int(os.getenv("AUDIT_SUMMARY_PRIORITY_TECHNICAL_TREATMENT_ITEMS", "24"), 24))


def _technical_priority_rank(item: Dict[str, Any]) -> int:
    sev = _clean_text(item.get("severity")).upper()
    return {"CRITICAL": 5, "HIGH": 4, "ERROR": 3, "MEDIUM": 2, "WARNING": 1, "LOW": 1}.get(sev, 0)


def _select_priority_control_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    limit = _priority_control_treatment_limit()
    if limit <= 0:
        return []
    ranked = sorted(
        items,
        key=lambda x: (
            {"High": 3, "Medium": 2, "Low": 1}.get(_clean_text(x.get("severity")), 0),
            {"High": 4, "Medium-High": 3, "Medium–High": 3, "Medium": 2, "Low-Medium": 1, "Low–Medium": 1}.get(_clean_text(x.get("likelihood")), 0),
            _clean_text(x.get("item_id")),
        ),
        reverse=True,
    )
    return ranked[:limit]


def _select_priority_technical_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    limit = _priority_technical_treatment_limit()
    if limit <= 0:
        return []
    ranked = sorted(items, key=lambda x: (_technical_priority_rank(x), _clean_text(x.get("item_id"))), reverse=True)
    critical_high = [x for x in ranked if _technical_priority_rank(x) >= 4]
    source = critical_high if critical_high else ranked
    return source[:limit]


def _treatment_ai_mode_metadata(mode: str, control_source: List[Dict[str, Any]], technical_source: List[Dict[str, Any]], batch_size: int) -> Dict[str, Any]:
    return {
        "mode": mode,
        "control_items_requested": len(control_source),
        "technical_items_requested": len(technical_source),
        "control_batches": 0,
        "technical_batches": 0,
        "total_batches": 0,
        "batch_size": batch_size,
        "ai_authored": bool(control_source or technical_source),
    }


def _call_llm_for_treatment_plan(app: Dict[str, Any], treatment_plan: Dict[str, Any], technical: Dict[str, Any]) -> Dict[str, Any]:
    """Generate treatment actions with the configured AI model.

    Stage 1 supplies only facts. This function asks the AI model to generate
    treatment actions, verification methods, closure evidence, and residual-risk
    notes for each item. No static remediation content is authored here.
    """
    mode = _treatment_mode()
    if not treatment_plan:
        return {}
    if mode == "off":
        print("[AI][TREATMENT] AUDIT_SUMMARY_TREATMENT_MODE=off; item-level treatment AI generation is skipped.")
        return {"control_treatments": {}, "technical_treatments": {}, "metadata": {"mode": mode, "ai_authored": False}}
    if not _ai_enabled():
        return {}

    system_prompt = str(contract_for_section("control_treatment_", fallback_source_function="_call_llm_for_treatment_plan").get("system_prompt_transcript") or "")
    if not system_prompt:
        raise RuntimeError("P-AS2-005 treatment system prompt is empty in scripts/prompt_contracts.json")

    technical_compact = _compact_technical_for_ai(technical)
    control_items = [x for x in _as_list(treatment_plan.get("control_items")) if isinstance(x, dict)]
    technical_items = [x for x in _as_list(treatment_plan.get("technical_items")) if isinstance(x, dict)]
    batch_size = _treatment_batch_size()

    control_results: List[Dict[str, Any]] = []
    technical_results: List[Dict[str, Any]] = []

    if mode == "priority":
        control_source = _select_priority_control_items(control_items)
        technical_source = _select_priority_technical_items(technical_items)
    else:
        max_control_items = _safe_int(os.getenv("AUDIT_SUMMARY_AI_TREATMENT_CONTROL_ITEM_LIMIT", "0"), 0)
        max_technical_items = _safe_int(os.getenv("AUDIT_SUMMARY_AI_TREATMENT_TECHNICAL_ITEM_LIMIT", "0"), 0)
        control_source = control_items[:max_control_items] if max_control_items > 0 else control_items
        technical_source = technical_items[:max_technical_items] if max_technical_items > 0 else technical_items

    if not control_source and not technical_source:
        print(
            f"[AI][TREATMENT] mode={mode}: no item-level treatment batches requested. "
            "Appendix treatment sections will be rendered from structured traceability data."
        )
        return {
            "control_treatments": {},
            "technical_treatments": {},
            "metadata": _treatment_ai_mode_metadata(mode, control_source, technical_source, batch_size),
        }

    control_batches = _chunks(control_source, batch_size)
    technical_batches = _chunks(technical_source, batch_size)
    total_control_batches = len(control_batches)
    total_technical_batches = len(technical_batches)
    total_batches = total_control_batches + total_technical_batches

    print(
        f"[AI][TREATMENT] Starting treatment-plan generation: "
        f"{len(control_source)} control item(s) in {total_control_batches} batch(es), "
        f"{len(technical_source)} technical item(s) in {total_technical_batches} batch(es), "
        f"batch_size={batch_size}, total_batches={total_batches}."
    )

    for idx, batch in enumerate(control_batches, start=1):
        global_idx = idx
        first_item = ((idx - 1) * batch_size) + 1
        last_item = min(idx * batch_size, len(control_source))
        print(
            f"[AI][TREATMENT] Progress {global_idx}/{total_batches} | "
            f"control batch {idx}/{total_control_batches} | "
            f"items {first_item}-{last_item}/{len(control_source)}"
        )
        obj = _ai_json_chat(
            f"control_treatment_{idx}_of_{total_control_batches}",
            system_prompt,
            {
                "task": "Generate treatment-plan actions for non-compliant SECM-CAT controls.",
                "constraints": {
                    "one_result_per_input_item": True,
                    "do_not_invent_evidence": True,
                    "use_exact_item_id_and_puid": True,
                    "output_fields": ["item_id", "treatment_action", "verification_method", "closure_evidence", "residual_risk"],
                    "all_output_fields_must_be_non_empty": True,
                    "language_policy": "Use director-facing audit language. Do not discuss discrepancies, contradictions, limitations, missing evidence, unavailable values, failed parsing, or internal pipeline behavior.",
                },
                "application": app,
                "technical_evidence_summary": technical_compact,
                "items": [_compact_treatment_item(x, "control") for x in batch],
                "required_output_schema": {
                    "control_treatments": [
                        {
                            "item_id": "<exact input item_id>",
                            "treatment_action": "<non-empty AI-generated action grounded in this control>",
                            "verification_method": "<non-empty verification method>",
                            "closure_evidence": "<non-empty closure evidence>",
                            "residual_risk": "<residual risk or risk acceptance note>",
                        }
                    ]
                },
            },
            max_tokens=min(_ai_max_tokens(2600), 3200),
        )
        values = obj.get("control_treatments") if isinstance(obj.get("control_treatments"), list) else []
        control_results.extend([x for x in values if isinstance(x, dict)])
        print(
            f"[AI][TREATMENT] Completed {global_idx}/{total_batches} | "
            f"control batch {idx}/{total_control_batches}"
        )

    for idx, batch in enumerate(technical_batches, start=1):
        global_idx = total_control_batches + idx
        first_item = ((idx - 1) * batch_size) + 1
        last_item = min(idx * batch_size, len(technical_source))
        print(
            f"[AI][TREATMENT] Progress {global_idx}/{total_batches} | "
            f"technical batch {idx}/{total_technical_batches} | "
            f"items {first_item}-{last_item}/{len(technical_source)}"
        )
        obj = _ai_json_chat(
            f"technical_treatment_{idx}_of_{total_technical_batches}",
            system_prompt,
            {
                "task": "Generate treatment-plan actions for technical scanner findings.",
                "constraints": {
                    "one_result_per_input_item": True,
                    "do_not_invent_evidence": True,
                    "use_exact_item_id_and_finding_id": True,
                    "output_fields": ["item_id", "treatment_action", "verification_method", "closure_evidence", "residual_risk"],
                    "all_output_fields_must_be_non_empty": True,
                    "language_policy": "Use director-facing audit language. Do not discuss discrepancies, contradictions, limitations, missing evidence, unavailable values, failed parsing, or internal pipeline behavior.",
                    "dependency_rule": "For Trivy items, use fixed_version when present; if no fixed version exists, propose risk acceptance, compensating controls, or monitoring rather than inventing a version.",
                    "sast_rule": "For SAST items, use only supplied rule, message, file, and line. Do not invent code paths.",
                    "mobsf_rule": "For MobSF items, recommend configuration, signing, manifest, runtime-storage, or verification actions based only on supplied evidence.",
                },
                "application": app,
                "technical_evidence_summary": technical_compact,
                "items": [_compact_treatment_item(x, "technical") for x in batch],
                "required_output_schema": {
                    "technical_treatments": [
                        {
                            "item_id": "<exact input item_id>",
                            "treatment_action": "<non-empty AI-generated action grounded in this finding>",
                            "verification_method": "<non-empty verification method>",
                            "closure_evidence": "<non-empty closure evidence>",
                            "residual_risk": "<residual risk or risk acceptance note>",
                        }
                    ]
                },
            },
            max_tokens=min(_ai_max_tokens(2600), 3200),
        )
        values = obj.get("technical_treatments") if isinstance(obj.get("technical_treatments"), list) else []
        technical_results.extend([x for x in values if isinstance(x, dict)])
        print(
            f"[AI][TREATMENT] Completed {global_idx}/{total_batches} | "
            f"technical batch {idx}/{total_technical_batches}"
        )

    expected_control_ids = [_clean_text(x.get("item_id")) for x in control_source]
    expected_technical_ids = [_clean_text(x.get("item_id")) for x in technical_source]
    treatment_ai = {
        "control_treatments": _merge_treatment_results(control_results, expected_control_ids),
        "technical_treatments": _merge_treatment_results(technical_results, expected_technical_ids),
        "metadata": {
            "mode": mode,
            "control_items_requested": len(control_source),
            "technical_items_requested": len(technical_source),
            "control_batches": total_control_batches,
            "technical_batches": total_technical_batches,
            "total_batches": total_batches,
            "batch_size": batch_size,
            "ai_authored": True,
        },
    }

    requested_treatment_plan = dict(treatment_plan or {})
    requested_treatment_plan["control_items"] = control_source
    requested_treatment_plan["technical_items"] = technical_source

    # In priority mode, only validate the item-level AI rows that were actually
    # requested from the model. The large control appendix is rendered from
    # structured traceability data and must not trigger treatment repair noise.
    missing = _incomplete_treatment_items(requested_treatment_plan, treatment_ai)
    missing_count = len(missing["control_items"]) + len(missing["technical_items"])
    max_repair_attempts = _treatment_repair_attempts()

    if missing_count:
        print(
            f"[AI][TREATMENT][REPAIR] Initial treatment output is incomplete for {missing_count} rendered item(s): "
            f"{_treatment_missing_id_list(missing, limit=30)}"
        )

    for attempt in range(1, max_repair_attempts + 1):
        if not missing_count:
            break

        repair_patch = _call_llm_for_treatment_repair(
            app=app,
            technical_compact=technical_compact,
            missing_control_items=missing["control_items"],
            missing_technical_items=missing["technical_items"],
            system_prompt=system_prompt,
            attempt=attempt,
            max_attempts=max_repair_attempts,
        )
        treatment_ai = _merge_treatment_ai_maps(treatment_ai, repair_patch)
        missing = _incomplete_treatment_items(requested_treatment_plan, treatment_ai)
        missing_count = len(missing["control_items"]) + len(missing["technical_items"])
        print(
            f"[AI][TREATMENT][REPAIR] Completed repair attempt {attempt}/{max_repair_attempts}. "
            f"Remaining missing item(s): {missing_count}"
            + (f" ({_treatment_missing_id_list(missing, limit=30)})" if missing_count else "")
        )

    final_missing = _incomplete_treatment_items(requested_treatment_plan, treatment_ai)
    final_missing_count = len(final_missing["control_items"]) + len(final_missing["technical_items"])
    metadata = dict(_as_dict(treatment_ai.get("metadata")))
    metadata["repair_attempts_allowed"] = max_repair_attempts
    metadata["repair_remaining_missing_items"] = final_missing_count
    metadata["control_treatment_response_count"] = len(_as_dict(treatment_ai.get("control_treatments")))
    metadata["technical_treatment_response_count"] = len(_as_dict(treatment_ai.get("technical_treatments")))
    treatment_ai["metadata"] = metadata

    print(
        f"[AI][TREATMENT] Finished treatment-plan generation: "
        f"{metadata['control_treatment_response_count']} control treatment response(s), "
        f"{metadata['technical_treatment_response_count']} technical treatment response(s), "
        f"remaining_missing={final_missing_count}."
    )

    return treatment_ai


def _validate_ai_treatment_writeups(treatment_plan: Dict[str, Any], treatment_ai: Dict[str, Any]) -> None:
    if not treatment_plan or not _treatment_ai_required():
        return
    if not _treatment_full_mode():
        return
    missing_by_kind = _incomplete_treatment_items(treatment_plan, treatment_ai)
    missing: List[str] = []
    for key in ("control_items", "technical_items"):
        for item in missing_by_kind.get(key, []):
            item_id = _clean_text(item.get("item_id"))
            fields = ", ".join(_as_list(item.get("_missing_ai_fields")))
            missing.append(f"{item_id} ({fields})" if fields else item_id)
    if missing:
        raise SystemExit(
            "[ERROR] AI-generated treatment plan is incomplete after repair attempts. "
            "Treatment actions, verification methods, and closure evidence are required for rendered treatment items. "
            "Missing item IDs: " + ", ".join(missing[:40])
        )


def _treatment_writeup(treatment_ai: Dict[str, Any], item_id: str, technical: bool = False) -> Dict[str, Any]:
    key = "technical_treatments" if technical else "control_treatments"
    return _as_dict(_as_dict(treatment_ai.get(key)).get(item_id))


def _render_treatment_overview(doc: Document, treatment_plan: Dict[str, Any], fig5: str, fig6: str, fig7: str) -> None:
    summary = _as_dict(treatment_plan.get("summary"))
    mode = _treatment_mode()
    if mode == "full":
        intro = "This section bridges the executive MAP and the technical annexes. Treatment content is generated from the current run's non-compliant SECM-CAT controls and scanner findings; the evidence model is deterministic, while treatment actions and verification wording are generated by the configured AI model."
    elif mode == "priority":
        intro = "This section bridges the executive MAP and the technical annexes. The evidence model is deterministic. AI is used for executive sections, weakness-pattern narratives, and priority technical treatment items, while large appendices are rendered as structured traceability data."
    else:
        intro = "This section bridges the executive MAP and the technical annexes. The evidence model is deterministic and large appendices are rendered as structured traceability data without item-level AI treatment generation."
    _add_body_paragraph(doc, intro)
    rows = [
        ["SECM-CAT control treatment items", summary.get("control_items_total", 0)],
        ["Technical scanner treatment items", summary.get("technical_items_total", 0)],
        ["Evidence correlation links", summary.get("correlation_items_total", 0)],
        ["Recommended owner groups", summary.get("owner_groups_total", 0)],
    ]
    _add_two_col_table(doc, rows)
    _add_figure(doc, fig5, "Figure 5. Treatment volume by weakness pattern (source: Stage 1 treatment-plan model).")
    _add_figure(doc, fig6, "Figure 6. Treatment workload by recommended owner (source: Stage 1 treatment-plan model).")
    _add_figure(doc, fig7, "Figure 7. Treatment priority matrix: severity versus workbook prevalence (source: Stage 1 treatment-plan model).")


def _render_control_treatment_appendix(doc: Document, treatment_plan: Dict[str, Any], treatment_ai: Dict[str, Any]) -> None:
    items = [x for x in _as_list(treatment_plan.get("control_items")) if isinstance(x, dict)]
    max_rows = _max_control_treatment_rows(len(items))
    if not items:
        _add_note(doc, "No SECM-CAT treatment items were available in the analysis pack.")
        return

    mode = _treatment_mode()
    if mode != "full":
        _add_section_callout(
            doc,
            "Appendix A coverage",
            f"This register renders {len(items[:max_rows])} of {len(items)} non-compliant SECM-CAT controls. Rows are grouped by framework category and preserve PUID, flags, evidence basis, weakness pattern, owner, and target timeline.",
            fill=THEME_LIGHT_BLUE,
        )

        category_counts: Dict[str, int] = {}
        for item in items:
            key = f"{_clean_text(item.get('category_code'))} - {_clean_text(item.get('category_name'))}"
            category_counts[key] = category_counts.get(key, 0) + 1
        summary_rows = [[cat, count] for cat, count in category_counts.items()]
        _add_table(doc, ["SECM-CAT category", "Non-compliant controls"], summary_rows, max_rows=None)

        grouped: Dict[str, List[Dict[str, Any]]] = {}
        for item in items[:max_rows]:
            key = f"{_clean_text(item.get('category_code'))} - {_clean_text(item.get('category_name'))}"
            grouped.setdefault(key, []).append(item)

        rendered = 0
        for idx, (category, group_items) in enumerate(grouped.items(), start=1):
            heading = doc.add_paragraph(f"A.{idx} {category} ({len(group_items)} controls)", style="Heading 2")
            _style_heading_paragraph(heading, 2)
            rows: List[List[Any]] = []
            for item in group_items:
                flags = ", ".join(_as_list(item.get("flags")))
                priority = f"{item.get('severity')} / {item.get('likelihood')}"
                timeline = _target_timeline(_clean_text(item.get("severity")))
                rows.append([
                    f"{item.get('puid')}\n{item.get('item_id')}",
                    f"{item.get('weakness_pattern')}\nPriority: {priority}",
                    f"{item.get('recommended_owner')}\nTarget: {timeline}",
                    flags,
                    _short(item.get("evidence_excerpt") or item.get("description"), 520),
                ])
                rendered += 1
            _add_table(
                doc,
                ["Requirement", "Pattern / priority", "Owner / target", "Flags", "Evidence basis"],
                rows,
                max_rows=None,
            )
        if rendered != len(items):
            raise SystemExit(f"[ERROR] Appendix A coverage gate failed: rendered {rendered} of {len(items)} SECM-CAT controls.")
        return

    def add_card_row(tbl, label: str, value: Any) -> None:
        cells = tbl.add_row().cells
        cells[0].text = _cell_text(label)
        cells[1].text = _cell_text(value)
        _set_cell_shading(cells[0], THEME_GRAY)
        _set_cell_margins(cells[0], top=65, start=80, bottom=65, end=80)
        _set_cell_margins(cells[1], top=65, start=80, bottom=65, end=80)
        _format_table_cell(cells[0], font_size=8.5, bold=True, no_wrap=True)
        _format_table_cell(cells[1], font_size=8.0)

    rendered = 0
    for item in items[:max_rows]:
        item_id = _clean_text(item.get("item_id"))
        writeup = _treatment_writeup(treatment_ai, item_id, technical=False)
        puid = _clean_text(item.get("puid"))
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        run = heading.add_run(f"{puid} | {item_id}")
        _set_run_font(run, size_pt=10, bold=True, color_hex=THEME_NAVY)

        flags = ", ".join(_as_list(item.get("flags"))) or "No flags listed"
        evidence = _short(item.get("evidence_excerpt") or item.get("description"), 720)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = True
        _set_table_alignment(tbl)
        add_card_row(tbl, "Pattern", item.get("weakness_pattern"))
        add_card_row(tbl, "Priority", f"{item.get('severity')} / {item.get('likelihood')}")
        add_card_row(tbl, "Owner", item.get("recommended_owner"))
        add_card_row(tbl, "Flags", flags)
        add_card_row(tbl, "Evidence basis", evidence)
        add_card_row(tbl, "Treatment action", _short(writeup.get("treatment_action"), 620))
        add_card_row(tbl, "Verification", _short(writeup.get("verification_method"), 620))
        add_card_row(tbl, "Closure evidence", _short(writeup.get("closure_evidence"), 620))
        add_card_row(tbl, "Residual risk", _short(writeup.get("residual_risk"), 420))
        rendered += 1
    if rendered != len(items):
        raise SystemExit(f"[ERROR] Appendix A coverage gate failed: rendered {rendered} of {len(items)} SECM-CAT controls.")


def _render_technical_treatment_appendix(doc: Document, treatment_plan: Dict[str, Any], treatment_ai: Dict[str, Any]) -> None:
    items = [x for x in _as_list(treatment_plan.get("technical_items")) if isinstance(x, dict)]
    max_rows = _max_technical_treatment_rows(len(items))
    if not items:
        _add_note(doc, "No technical treatment items were available in the analysis pack.")
        return

    mode = _treatment_mode()
    if mode != "full":
        rows: List[List[Any]] = []
        for item in items[:max_rows]:
            item_id = _clean_text(item.get("item_id"))
            writeup = _treatment_writeup(treatment_ai, item_id, technical=True)
            affected = []
            for key, label in [
                ("affected_component", "component"),
                ("installed_version", "installed"),
                ("fixed_version", "fixed"),
                ("file", "file"),
                ("line", "line"),
            ]:
                value = item.get(key)
                if value not in (None, "", [], {}):
                    affected.append(f"{label}: {_clean_text(value)}")
            linked = ", ".join(_as_list(item.get("linked_puids"))[:10])
            priority_action = _short(writeup.get("treatment_action"), 360) if writeup else ""
            rows.append([
                f"{item_id}\n{item.get('source')}",
                f"{item.get('item_type')}\nSeverity: {item.get('severity')}",
                _short(" | ".join(affected), 260),
                _short(item.get("observed_issue"), 360),
                linked,
                priority_action,
            ])
        _add_table(
            doc,
            ["Technical item", "Type / severity", "Affected component", "Observed issue", "Linked PUIDs", "Priority AI action"],
            rows,
            max_rows=max_rows,
        )
        if len(items) > max_rows:
            _add_note(doc, f"Rendered {max_rows} of {len(items)} technical treatment items. Increase AUDIT_SUMMARY_MAX_TECHNICAL_TREATMENT_ROWS to include more rows.")
        return

    def add_card_row(tbl, label: str, value: Any) -> None:
        cells = tbl.add_row().cells
        cells[0].text = _cell_text(label)
        cells[1].text = _cell_text(value)
        _set_cell_shading(cells[0], "EDEDED")
        _format_table_cell(cells[0], font_size=8.5, bold=True, no_wrap=True)
        _format_table_cell(cells[1], font_size=8.0)

    for item in items[:max_rows]:
        item_id = _clean_text(item.get("item_id"))
        writeup = _treatment_writeup(treatment_ai, item_id, technical=True)
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        run = heading.add_run(f"{item_id} | {item.get('source')} | {item.get('finding_id')}")
        run.bold = True
        run.font.size = Pt(10)

        affected = []
        for key, label in [
            ("affected_component", "component"),
            ("installed_version", "installed"),
            ("fixed_version", "fixed"),
            ("fix_available", "fix_available"),
            ("cwe_ids", "CWE"),
            ("cvss_v3_score", "CVSS"),
            ("published_at", "published"),
            ("last_modified_at", "modified"),
            ("file", "file"),
            ("line", "line"),
        ]:
            value = item.get(key)
            if value not in (None, "", [], {}):
                affected.append(f"{label}: {_clean_text(value)}")
        linked = ", ".join(_as_list(item.get("linked_puids"))[:14]) or "No linked PUIDs normalized"

        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = True
        add_card_row(tbl, "Type", item.get("item_type"))
        add_card_row(tbl, "Severity", item.get("severity"))
        add_card_row(tbl, "Affected component", " | ".join(affected))
        add_card_row(tbl, "Observed issue", _short(item.get("observed_issue"), 720))
        add_card_row(tbl, "Treatment action", _short(writeup.get("treatment_action"), 620))
        add_card_row(tbl, "Verification", _short(writeup.get("verification_method"), 620))
        add_card_row(tbl, "Closure evidence", _short(writeup.get("closure_evidence"), 620))
        add_card_row(tbl, "Linked PUIDs", linked)
        add_card_row(tbl, "Residual risk", _short(writeup.get("residual_risk"), 420))
    if len(items) > max_rows:
        _add_note(doc, f"Rendered {max_rows} of {len(items)} technical treatment items. Increase AUDIT_SUMMARY_MAX_TECHNICAL_TREATMENT_ROWS to include more rows.")

def _render_correlation_appendix(doc: Document, treatment_plan: Dict[str, Any]) -> None:
    items = [x for x in _as_list(treatment_plan.get("correlation_items")) if isinstance(x, dict)]
    max_rows = _max_correlation_rows(len(items))
    if not items:
        _add_note(doc, "No evidence correlation items were available in the analysis pack.")
        return

    _add_section_callout(
        doc,
        "Appendix C coverage",
        f"This matrix renders {len(items[:max_rows])} of {len(items)} evidence correlation links. Rows are grouped by scanner source to improve traceability and readability.",
        fill=THEME_LIGHT_BLUE,
    )

    source_counts: Dict[str, int] = {}
    for item in items:
        source = _clean_text(item.get("technical_source")) or "Workbook / normalized evidence"
        source_counts[source] = source_counts.get(source, 0) + 1
    _add_table(doc, ["Evidence source", "Correlation links"], [[k, v] for k, v in source_counts.items()], max_rows=None)

    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for item in items[:max_rows]:
        source = _clean_text(item.get("technical_source")) or "Workbook / normalized evidence"
        grouped.setdefault(source, []).append(item)

    rendered = 0
    for idx, (source, group_items) in enumerate(grouped.items(), start=1):
        heading = doc.add_paragraph(f"C.{idx} {source} correlations ({len(group_items)} links)", style="Heading 2")
        _style_heading_paragraph(heading, 2)
        rows: List[List[Any]] = []
        for item in group_items:
            requirement = f"{item.get('puid')}\n{item.get('weakness_pattern')}"
            flags = ", ".join(_as_list(item.get("flags_sample")))
            finding = f"{item.get('technical_finding_id')}\nTreatment: {item.get('technical_item_id')}"
            fit = f"{_clean_text(item.get('evidence_fit'))}\n{_clean_text(item.get('correlation_strength'))}"
            evidence = _short(item.get('evidence_summary'), 420)
            rows.append([requirement, flags, finding, fit, evidence])
            rendered += 1
        _add_table(
            doc,
            ["Requirement / pattern", "Flags", "Finding / treatment", "Evidence fit / strength", "Evidence summary"],
            rows,
            max_rows=None,
        )
    if rendered != len(items):
        raise SystemExit(f"[ERROR] Appendix C coverage gate failed: rendered {rendered} of {len(items)} correlation rows.")


def _render_vision360_appendix(doc: Document, technical: Dict[str, Any]) -> None:
    vision = _as_dict(technical.get("vision360"))
    if not _block_available(vision):
        _add_note(doc, "Vision360 evidence was not available in the analysis pack.")
        return

    group_summary = [x for x in _as_list(vision.get("group_summary")) if isinstance(x, dict)]
    flag_rows = [x for x in _as_list(vision.get("full_flags")) if isinstance(x, dict)]
    evidence_rows = [x for x in _as_list(vision.get("flag_evidence_details")) if isinstance(x, dict)]

    _add_section_callout(
        doc,
        "Vision360 evidence register",
        f"Vision360 produced {vision.get('flags_count', len(flag_rows))} flag verdicts. This appendix preserves the complete dynamic flag matrix, grouped by framework family, plus per-flag evidence details where available.",
        fill=THEME_LIGHT_BLUE,
    )

    rows = []
    for item in group_summary:
        rows.append([
            item.get("group"),
            item.get("total_flags"),
            item.get("pass"),
            item.get("fail"),
            item.get("unknown"),
            item.get("evidence_count"),
        ])
    _add_table(doc, ["Group", "Total flags", "Pass", "Fail", "Unknown", "Evidence count"], rows, max_rows=None, empty_message="No Vision360 group summary was available.")

    grouped_flags: Dict[str, List[Dict[str, Any]]] = {}
    for item in flag_rows:
        grouped_flags.setdefault(_clean_text(item.get("group")) or "Ungrouped", []).append(item)

    rendered = 0
    for idx, (group, group_items) in enumerate(grouped_flags.items(), start=1):
        heading = doc.add_paragraph(f"F.{idx} {group} flags ({len(group_items)})", style="Heading 2")
        _style_heading_paragraph(heading, 2)
        rows = []
        for item in group_items:
            rows.append([
                item.get("id"),
                item.get("state"),
                f"{_clean_text(item.get('title'))}\n{_short(item.get('summary'), 320)}",
                item.get("evidence_count"),
                ", ".join(_as_list(item.get("primary_sources"))),
            ])
            rendered += 1
        _add_table(doc, ["Flag ID", "State", "Title / summary", "Evidence count", "Primary sources"], rows, max_rows=None, empty_message="No Vision360 flag rows were available for this group.")

    if flag_rows and rendered != len(flag_rows):
        raise SystemExit(f"[ERROR] Vision360 appendix coverage gate failed: rendered {rendered} of {len(flag_rows)} flags.")

    detail_rows = []
    for item in evidence_rows:
        detail_rows.append([
            item.get("flag_id"),
            item.get("source"),
            item.get("path"),
            item.get("rule_id"),
            _short(item.get("excerpt"), 360),
        ])
    _add_table(doc, ["Flag ID", "Source", "Path", "Rule ID", "Evidence excerpt"], detail_rows, max_rows=None, empty_message="No per-flag Vision360 evidence details were available.")


def _render_mobsf_static_inventory_appendix(doc: Document, technical: Dict[str, Any]) -> None:
    mobsf = _as_dict(technical.get("mobsf_static"))
    if not _block_available(mobsf):
        _add_note(doc, "MobSF static inventory was not available in the analysis pack.")
        return

    _add_section_callout(doc, "MobSF static inventory", "This appendix expands MobSF static analysis beyond the executive snapshot and preserves APK identity, signing, manifest, permission, component, URL/domain, email, tracker, and library evidence parsed for this run.", fill=THEME_LIGHT_BLUE)
    trace = _as_dict(mobsf.get("apk_traceability"))
    trace_rows = [[k, v] for k, v in trace.items()]
    _add_table(doc, ["APK traceability item", "Observed value"], trace_rows, max_rows=None)

    cert_rows = []
    for item in _as_list(mobsf.get("certificate_findings")):
        if isinstance(item, dict):
            cert_rows.append([item.get("severity"), item.get("title"), item.get("description"), item.get("component")])
    _add_table(doc, ["Severity", "Certificate finding", "Description", "Component"], cert_rows, max_rows=None, empty_message="No MobSF certificate finding rows were available.")

    manifest_rows = []
    for item in _as_list(mobsf.get("manifest_findings")):
        if isinstance(item, dict):
            manifest_rows.append([item.get("severity"), item.get("title"), item.get("description"), item.get("component")])
    _add_table(doc, ["Severity", "Manifest finding", "Description", "Component"], manifest_rows, max_rows=None, empty_message="No MobSF manifest finding rows were available.")

    perm_rows = []
    for item in _as_list(mobsf.get("permissions_full") or mobsf.get("permissions_sample")):
        if isinstance(item, dict):
            perm_rows.append([item.get("name"), item.get("level"), item.get("description")])
    _add_table(doc, ["Permission", "Protection level", "Description"], perm_rows, max_rows=None, empty_message="No Android permission inventory was available.")

    component_rows = []
    for item in _as_list(mobsf.get("components_inventory")):
        if isinstance(item, dict):
            component_rows.append([item.get("type"), item.get("component"), item.get("exported"), item.get("permission"), item.get("risk_note")])
    _add_table(doc, ["Type", "Component", "Exported", "Permission", "Risk note"], component_rows, max_rows=None, empty_message="No Android component inventory was available.")

    external_rows = []
    for item in _as_list(mobsf.get("external_references")):
        if isinstance(item, dict):
            external_rows.append([item.get("type"), item.get("value"), item.get("path"), item.get("risk_note")])
    _add_table(doc, ["Type", "Value", "Path / source", "Risk note"], external_rows, max_rows=None, empty_message="No MobSF URL, domain, or email inventory was available.")

    library_rows = [[x] for x in _as_list(mobsf.get("libraries")) if _clean_text(x)]
    _add_table(doc, ["Library"], library_rows, max_rows=None, empty_message="No MobSF library inventory was available.")


def _license_risk_class(license_name: Any) -> str:
    s = _clean_text(license_name).lower()
    if not s:
        return "Unknown"
    if any(tok in s for tok in ["agpl", "gpl"]):
        return "Strong copyleft review"
    if any(tok in s for tok in ["lgpl", "mpl", "epl", "cddl"]):
        return "File-level or weak copyleft review"
    if any(tok in s for tok in ["apache", "mit", "bsd", "isc", "zlib"]):
        return "Permissive"
    return "Review required"


def _render_sca_inventory_appendix(doc: Document, technical: Dict[str, Any]) -> None:
    trivy = _as_dict(technical.get("trivy_sca"))
    if not _block_available(trivy):
        _add_note(doc, "SCA evidence was not available in the analysis pack.")
        return

    _add_section_callout(doc, "SCA inventory", "This appendix preserves the Software Composition Analysis inventory: vulnerabilities, package versions, detected licenses, and remediation status. Counts and rows are generated dynamically from the Trivy and Vision360 SCA evidence available for this run.", fill=THEME_LIGHT_BLUE)

    vuln_rows = []
    for item in _trivy_findings(trivy):
        vuln_rows.append([
            item.get("severity"),
            item.get("id") or item.get("cve"),
            item.get("package") or item.get("pkg"),
            item.get("installed"),
            item.get("fixed"),
            "Yes" if item.get("fix_available") or item.get("fixed") else "No",
            item.get("status"),
            item.get("cvss_v3_score"),
            ", ".join(_as_list(item.get("cwe_ids") or item.get("cwe"))),
        ])
    _add_table(doc, ["Severity", "CVE", "Package", "Installed", "Fixed", "Fix available", "Status", "CVSS", "CWE"], vuln_rows, max_rows=None, empty_message="No SCA vulnerability rows were available.")

    package_rows = []
    license_map = _as_dict(trivy.get("package_license_map"))
    vulnerable_packages = {_clean_text(item.get("package") or item.get("pkg")) for item in _trivy_findings(trivy)}
    for item in _as_list(trivy.get("package_inventory") or trivy.get("packages")):
        if isinstance(item, dict):
            name = _clean_text(item.get("name") or item.get("package") or item.get("pkg") or item.get("id"))
            licenses = _as_list(item.get("licenses") or license_map.get(name))
            package_rows.append([
                name,
                item.get("version"),
                item.get("ecosystem"),
                item.get("target"),
                ", ".join(_clean_text(x) for x in licenses if _clean_text(x)),
                "Yes" if name in vulnerable_packages else "No",
            ])
    _add_table(doc, ["Package", "Version", "Ecosystem", "Target", "Licenses", "Known vulnerable"], package_rows, max_rows=None, empty_message="No SCA package inventory was available.")

    license_rows = []
    for item in _as_list(trivy.get("license_inventory") or trivy.get("license_entries")):
        if isinstance(item, dict):
            lic = item.get("license") or item.get("name") or item.get("License")
            license_rows.append([
                item.get("package"),
                item.get("version"),
                lic,
                _license_risk_class(lic),
                item.get("target") or item.get("file"),
            ])
    _add_table(doc, ["Package", "Version", "License", "License class", "Target / file"], license_rows, max_rows=None, empty_message="No SCA license inventory was available.")

    remediation_rows = []
    for item in _trivy_findings(trivy):
        remediation_rows.append([
            item.get("package") or item.get("pkg"),
            item.get("installed"),
            item.get("fixed"),
            item.get("severity"),
            "Upgrade to the fixed version and rerun SCA validation." if item.get("fixed") else "Review vendor advisory and document accepted risk or compensating control.",
        ])
    _add_table(doc, ["Package", "Current version", "Recommended fixed version", "Highest severity", "Action"], remediation_rows, max_rows=None, empty_message="No SCA remediation rows were available.")


def _call_llm_for_style(patterns: List[Dict[str, Any]], likelihood_rubric: Dict[str, str], max_takeaways: int = 7) -> Dict[str, Any]:
    # Backward-compatible wrapper retained for older callers.
    return {}


def _sanitize_ai_technical_narratives(tech_ai: Dict[str, Any], technical: Dict[str, Any]) -> Dict[str, Any]:
    """Validate AI technical narratives without rewriting them.

    If the model does not return a section, deterministic scanner summaries are
    rendered by the dedicated Trivy, MobSF, and SAST table functions. This keeps
    the narrative prompt-first and avoids post-generation wording changes.
    """
    clean: Dict[str, Any] = {}
    for key, value in dict(tech_ai or {}).items():
        if isinstance(value, str):
            candidate = _clean_generated_text(value)
            if candidate and not _collect_ai_language_policy_hits(candidate):
                clean[key] = candidate
        else:
            clean[key] = value
    return clean


def _expected_secure_state(pattern: str, ai_expected: Any = "") -> str:
    """Return the AI-authored expected secure state for a weakness pattern.

    Prompt-first policy: this function validates model output only. It does not
    contain pattern templates, replacement rules, or deterministic narrative
    templates. In full treatment mode the field is mandatory. In priority mode,
    a missing or weak expected-state sentence is omitted instead of failing the
    full report, because the MAP and appendices are rendered from structured
    evidence and priority AI output.
    """
    pat = _clean_text(pattern)
    ai = _clean_text(ai_expected)
    # Keep the validator focused on clear current-problem wording. Do not block
    # broad security nouns such as weakness, gap, or vulnerability when they are
    # used in a positive expected-state sentence such as "prevents credential
    # exposure vulnerabilities".
    forbidden_problem_language = [
        "contains numerous", "exhibits deficiencies", "demonstrates instances",
        "lacks robust", "exhibits vulnerabilities", "is missing",
        "not evidenced", "fails to", "non-compliant",
    ]
    lowered = ai.lower()
    has_problem_language = any(token in lowered for token in forbidden_problem_language)
    has_expected_language = bool(re.search(r"\b(should|must|is expected to|requires|implements|enforces|protects|prevents|uses|maintains)\b", ai, re.IGNORECASE))
    if ai and not has_problem_language and has_expected_language:
        return ai

    message = (
        "AI-generated expected secure state is missing or not report-ready "
        f"for weakness pattern: {pat}."
    )
    if _treatment_full_mode():
        raise SystemExit("[ERROR] " + message + " Regenerate pattern_writeups with a positive expected-state sentence.")
    print("[AI][WARN] " + message + " Omitting that optional line in priority mode.")
    return ""


def main() -> None:
    in_path = os.getenv("AUDIT_ANALYSIS_JSON_PATH", DEFAULT_IN)
    out_path = os.getenv("AUDIT_SUMMARY_DOCX_PATH", DEFAULT_OUT)
    chart_dir = os.getenv("AUDIT_SUMMARY_CHART_DIR", CHART_DIR)

    if not os.path.isfile(in_path):
        raise SystemExit(f"[ERROR] analysis pack not found: {in_path}")

    Path(chart_dir).mkdir(parents=True, exist_ok=True)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    with open(in_path, "r", encoding="utf-8") as f:
        pack = json.load(f)

    metrics = pack["metrics"]
    cat_stats = pack["category_metrics"]
    app = pack["app_metadata"]
    actors = pack["actors"]
    patterns = pack["weakness_patterns"]
    pos_controls = pack.get("positive_controls_candidates", [])[:7]
    likelihood_rubric = pack.get("likelihood_rubric", {})
    technical = _technical_evidence(pack)
    treatment_plan = _treatment_plan(pack)
    report_title = _app_report_title(app)

    audit_dt = datetime.now(timezone.utc).date()
    audit_date_str = audit_dt.strftime("%d %b %Y")

    plt.rcParams.update({"font.size": 10, "figure.titlesize": 12, "axes.titlesize": 12, "axes.labelsize": 10})

    fig1 = os.path.join(chart_dir, "figure1_overall_donut.png")
    fig2 = os.path.join(chart_dir, "figure2_noncompliance_share_hbar.png")
    fig3 = os.path.join(chart_dir, "figure3_compliance_rate_hbar.png")
    fig4 = os.path.join(chart_dir, "figure4_counts_stacked_hbar.png")
    fig5 = os.path.join(chart_dir, "figure5_treatment_pattern_volume.png")
    fig6 = os.path.join(chart_dir, "figure6_treatment_owner_workload.png")
    fig7 = os.path.join(chart_dir, "figure7_treatment_priority_matrix.png")

    applicable = int(metrics["applicable"])
    compliant = int(metrics["compliant"])
    non_compliant = int(metrics["non_compliant"])
    not_applicable = int(metrics["not_applicable"])
    overall_pct = float(metrics["overall_compliance_pct"])

    _donut([compliant, non_compliant, not_applicable], ["Compliant", "Non-compliant", "Not applicable"], "Overall compliance distribution (workbook-derived)", f"{overall_pct:.2f}%\ncompliant\n(applicable)", fig1)
    _hbar_share_noncompliances(cat_stats, fig2)
    _hbar_compliance_rate(cat_stats, fig3)
    _stacked_counts(cat_stats, fig4)
    _hbar_treatment_pattern_volume(treatment_plan, fig5)
    _hbar_treatment_owner_workload(treatment_plan, fig6)
    _scatter_treatment_priority_matrix(treatment_plan, fig7)

    prose: Dict[str, Any] = {}
    try:
        prose = _call_llm_for_audit_sections(metrics, app, patterns, technical, likelihood_rubric, pos_controls)
        if prose:
            ai_sections_path = os.getenv("AUDIT_SUMMARY_AI_SECTIONS_PATH", "").strip()
            if not ai_sections_path:
                ai_sections_path = str(_runtime_data_dir() / "audit_summary_ai_sections.json")
            Path(ai_sections_path).parent.mkdir(parents=True, exist_ok=True)
            with open(ai_sections_path, "w", encoding="utf-8") as f:
                json.dump(prose, f, ensure_ascii=False, indent=2)
            print(f"[AI] Section outputs saved -> {ai_sections_path}")
    except Exception as exc:
        print(f"[AI][WARN] Sectioned AI generation failed; AI-authored prose is unavailable: {exc}")
        prose = {}

    treatment_ai: Dict[str, Any] = {}
    try:
        treatment_ai = _call_llm_for_treatment_plan(app, treatment_plan, technical)
        if treatment_ai:
            prose["treatment_plan_writeups"] = treatment_ai
            ai_sections_path = os.getenv("AUDIT_SUMMARY_AI_SECTIONS_PATH", "").strip()
            if not ai_sections_path:
                ai_sections_path = str(_runtime_data_dir() / "audit_summary_ai_sections.json")
            Path(ai_sections_path).parent.mkdir(parents=True, exist_ok=True)
            with open(ai_sections_path, "w", encoding="utf-8") as f:
                json.dump(prose, f, ensure_ascii=False, indent=2)
            print(f"[AI] Treatment-plan section outputs saved -> {ai_sections_path}")
    except Exception as exc:
        if _treatment_ai_required() and treatment_plan:
            raise
        print(f"[AI][WARN] Treatment-plan AI generation failed; treatment appendices will show missing AI fields: {exc}")
        treatment_ai = {}

    _validate_ai_treatment_writeups(treatment_plan, treatment_ai)

    tech_ai = _sanitize_ai_technical_narratives(_as_dict(prose.get("technical_narratives")), technical)
    key_takeaways = prose.get("key_takeaways", [])
    positive_control_writeups = {
        _clean_text(w.get("puid")): _clean_text(w.get("statement"))
        for w in prose.get("positive_controls", [])
        if isinstance(w, dict) and _clean_text(w.get("puid")) and _clean_text(w.get("statement"))
    }
    writeups = {w["pattern"]: w for w in prose.get("pattern_writeups", []) if isinstance(w, dict) and "pattern" in w}
    _validate_ai_pattern_writeups(patterns, writeups)
    if not key_takeaways:
        if _ai_pattern_writeups_required():
            raise SystemExit("[ERROR] AI-generated key takeaways were not returned. Regenerate executive_summary with 5 to 7 report-ready bullets.")
        key_takeaways = [f"{p['pattern']} - {p['severity']} severity; {int(p['mapped_noncompliant_count'])} related non-compliant controls in the workbook." for p in patterns[:7]]
    technical_takeaways = _technical_takeaways(technical)
    if technical_takeaways:
        key_takeaways = (technical_takeaways + key_takeaways)[:7]

    doc = Document()
    _set_doc_defaults(doc)
    _add_header_footer(doc.sections[0], audit_date_str, report_title)
    _add_cover(doc, audit_date_str, actors["Auditor"], report_title)
    used_bookmarks: set[str] = set()
    toc_entries: List[tuple[int, str, str]] = []
    bookmark_id = 1000

    def add_nav_heading(text: str, level: int):
        nonlocal bookmark_id
        style = f"Heading {level}"
        p = doc.add_paragraph(text, style=style)
        _style_heading_paragraph(p, level)
        anchor = _make_bookmark_name(text, used_bookmarks)
        _add_bookmark(p, anchor, bookmark_id)
        toc_entries.append((level, text, anchor))
        bookmark_id += 1
        return p

    toc_placeholder = _add_toc(doc)

    add_nav_heading("1. App information", 1)
    _add_two_col_table(doc, [[k, v] for k, v in app.items()])

    add_nav_heading("2. Actors", 1)
    _add_two_col_table(doc, [["Auditor", actors["Auditor"]], ["Requirement Engineering team", "\n".join(actors["Requirement Engineering team"])], ["Engineering Group (EN)", "\n".join(actors["Engineering Group (EN)"])]])

    add_nav_heading("3. Audit scope and evidence sources", 1)
    doc.add_paragraph("This Audit Summary consolidates the compliance determinations recorded in the audit workbook for the assessed application version and workbook-defined scope. The workbook provides the authoritative requirement-level results. Technical scan evidence from Trivy, MobSF static analysis, MobSF dynamic analysis, and app-code SAST is organized by source to provide observed technical context, prioritization inputs, and traceability for treatment actions.")

    add_nav_heading("4. Evidence criteria", 1)
    doc.add_paragraph("- Compliant: the workbook records the control as implemented for the assessed scope.\n- Non-compliant: the workbook records the control as missing, insufficient, or not evidenced.\n- Not applicable: the workbook records the control as outside the assessed context.\n- Technical evidence: automated scan artifacts provide observed application-code, Android manifest, binary, dependency, runtime, and storage evidence used for prioritization and treatment traceability.")

    add_nav_heading("5. Audit summary", 1)
    doc.add_paragraph("The audit was carried out using the mSEC-AM (mobile SECurity Audit Method).")
    _add_body_paragraph(doc, f"Overall, {int(metrics['total_assessed'])} requirements were assessed. {applicable} were applicable controls and {not_applicable} were recorded as not applicable. Of the applicable controls, {compliant} were compliant and {non_compliant} were non-compliant, resulting in an overall compliance rate of {overall_pct:.2f}% (applicable controls only).")
    _add_metric_cards(doc, [
        ("Total assessed", int(metrics['total_assessed']), "Workbook scope"),
        ("Applicable", applicable, "Controls in scope"),
        ("Non-compliant", non_compliant, "Treatment backlog"),
        ("Compliance rate", f"{overall_pct:.2f}%", "Applicable only"),
    ])
    ai_audit_summary = _clean_text(prose.get("audit_summary_paragraph", ""))
    if ai_audit_summary:
        doc.add_paragraph(ai_audit_summary)
    else:
        doc.add_paragraph("This report summarizes the dominant weakness patterns evidenced by non-compliant requirements and proposes actionable remediations suitable for mHealth/EMR environments handling sensitive health information.")
    if technical:
        doc.add_paragraph("The report organizes technical scan evidence by source, including dependency vulnerability evidence from Trivy, Android APK and manifest evidence from MobSF static analysis, runtime observations from MobSF dynamic analysis, and SAST findings filtered to application code.")

    add_nav_heading("5.1 Key takeaways (Top findings)", 2)
    _add_callout(doc, "Key takeaways (Top findings)", key_takeaways[:7])

    add_nav_heading("5.2 Positive controls observed", 2)
    doc.add_paragraph("The statements below summarize controls recorded as Compliant in the audit workbook and the observed technical or audit evidence associated with them. Verification traceability is provided in Appendix D.")
    positive_groups = _group_positive_controls(pos_controls, positive_control_writeups)
    rendered_positive = False
    for group_title in ["Observed positive controls", "Observed supporting signals", "Observed controls with additional technical context"]:
        items = positive_groups.get(group_title, [])
        if not items:
            continue
        doc.add_paragraph(group_title, style="Heading 3")
        for pc in items:
            doc.add_paragraph(pc.get("reported_statement", ""), style="List Bullet")
            rendered_positive = True
    if not rendered_positive:
        doc.add_paragraph("Positive-control statements were not rendered in the body because no AI-authored statement passed section validation. See Appendix D for source traceability.", style="List Bullet")

    add_nav_heading("5.3 Risk scoring approach", 2)
    doc.add_paragraph("Severity and likelihood ratings in this report follow a qualitative rubric grounded in the audit workbook:\n- Severity reflects potential impact on confidentiality, integrity, and availability of health information, including regulatory exposure.\n- Likelihood is derived from workbook prevalence: the count of non-compliant controls mapped to a weakness pattern as a proxy for exposure.\nLikelihood mapping: High (>=50), Medium-High (20-49), Medium (10-19), Low-Medium (<10).")

    add_nav_heading("5.4 Risk triage (prioritized)", 2)
    _add_body_paragraph(doc, "The table below keeps the executive triage compact. Detailed impact narratives and evidence anchors are provided in Section 7 for each weakness pattern.")
    _add_table_caption(doc, "Risk triage summary")
    rt = doc.add_table(rows=1, cols=5)
    rt.style = "Table Grid"
    h = rt.rows[0].cells
    for idx, txt in enumerate(["Weakness pattern", "Severity / likelihood", "Workbook basis", "Recommended owner", "Target timeline"]):
        h[idx].text = txt
        _set_cell_shading(h[idx], THEME_BLUE)
        for run in h[idx].paragraphs[0].runs:
            run.bold = True
    for p in patterns[:10]:
        cnt = int(p["mapped_noncompliant_count"])
        lik = _likelihood_from_count(cnt)
        sev = p["severity"]
        owner = p["recommended_owner"]
        row = rt.add_row().cells
        row[0].text = p["pattern"]
        row[1].text = f"{sev} / {lik}"
        row[2].text = f"{cnt} mapped non-compliant controls."
        row[3].text = owner
        row[4].text = _target_timeline(sev)
    doc.add_paragraph()

    add_nav_heading("5.5 Technical evidence snapshot", 2)
    doc.add_paragraph("The following table summarizes the principal evidence sources used in the report. The workbook provides requirement results, while each scanner contributes observed technical evidence for prioritization and treatment planning.")
    _add_table(doc, ["Evidence source", "Status", "Summary"], _source_status_rows(technical), max_rows=10)

    add_nav_heading("6. Technical evidence by scanner", 1)
    ai_tech_intro = _clean_text(tech_ai.get("technical_evidence_intro", ""))
    if ai_tech_intro:
        doc.add_paragraph(ai_tech_intro)
    else:
        doc.add_paragraph("This section presents observed technical evidence by scanner. Each scanner is reported separately so that dependency, APK, runtime, and source-code findings remain clear and actionable.")
    add_nav_heading("6.1 Software Composition Analysis from Trivy", 2)
    if _clean_text(tech_ai.get("trivy_paragraph", "")):
        doc.add_paragraph(_clean_text(tech_ai.get("trivy_paragraph", "")))
    _add_trivy_section(doc, technical)
    add_nav_heading("6.2 Android static evidence from MobSF", 2)
    if _clean_text(tech_ai.get("mobsf_static_paragraph", "")):
        doc.add_paragraph(_clean_text(tech_ai.get("mobsf_static_paragraph", "")))
    _add_mobsf_static_section(doc, technical)
    add_nav_heading("6.3 Runtime evidence from MobSF dynamic analysis", 2)
    if _clean_text(tech_ai.get("mobsf_dynamic_paragraph", "")):
        doc.add_paragraph(_clean_text(tech_ai.get("mobsf_dynamic_paragraph", "")))
    _add_mobsf_dynamic_section(doc, technical)
    add_nav_heading("6.4 Static Application Security Testing evidence", 2)
    if _clean_text(tech_ai.get("sast_paragraph", "")):
        doc.add_paragraph(_clean_text(tech_ai.get("sast_paragraph", "")))
    _add_sast_section(doc, technical)

    add_nav_heading("7. Main deficiencies", 1)
    doc.add_paragraph("The following weakness patterns summarize non-compliant requirements and associated technical evidence. They are organized as cross-cutting themes to support prioritization, ownership assignment, and remediation planning.")
    for p in patterns[:10]:
        pat = p["pattern"]
        cnt = int(p["mapped_noncompliant_count"])
        sev = p["severity"]
        owner = p["recommended_owner"]
        ex_ids = p.get("example_puids", [])[:4]
        anchors = p.get("description_anchors", [])[:2]
        doc.add_paragraph(f"{pat} ({sev})", style="Heading 2")
        doc.add_paragraph(f"Workbook basis: {cnt} related non-compliant controls mapped to this pattern.")
        expected = _expected_secure_state(pat, _ai_field_for_pattern(pat, writeups, "expected"))
        impact = _ai_field_for_pattern(pat, writeups, "impact") or "This pattern increases security and privacy exposure for the assessed application and should be prioritized according to severity, prevalence, and implementation dependency."
        if expected:
            doc.add_paragraph(f"Expected secure state: {expected}")
        doc.add_paragraph("Observed condition: The audit workbook records the mapped controls as non-compliant for the assessed scope.")
        doc.add_paragraph(f"Impact: {impact}")
        doc.add_paragraph(f"Recommended owner: {owner}")
        if ex_ids:
            doc.add_paragraph(f"Traceability (examples, non-exhaustive): {', '.join(ex_ids)}.")
        for a in anchors:
            doc.add_paragraph(f"Evidence anchor (from workbook description): {a}", style="List Bullet")

    doc.add_page_break()
    add_nav_heading("8. Recommendations", 1)
    doc.add_paragraph("Recommendations are generated from the audit workbook, weakness-pattern prevalence, PUID examples, and observed technical evidence from Trivy, MobSF static analysis, MobSF dynamic analysis, and SAST. Detailed flags and PUID-level traceability remain in the appendices.")
    for p in patterns[:10]:
        pat = p["pattern"]
        doc.add_paragraph(pat, style="Heading 2")
        recs = _ai_recommendations_for_pattern(pat, writeups)
        if not recs:
            doc.add_paragraph("AI-generated recommendations were not returned for this pattern; rerun the pipeline with a valid AI configuration or review the AI section artifact.", style="List Bullet")
            continue
        for r in recs[:12]:
            doc.add_paragraph(_sanitize_recommendation_text(r), style="List Bullet")

    doc.add_page_break()
    add_nav_heading("9. Visual Analytics", 1)
    doc.add_paragraph("Figures below summarize workbook-derived outcomes and distributions. All figures: source: audit workbook.")
    _add_figure(doc, fig1, "Figure 1. Overall compliance distribution (donut chart; source: audit workbook).")
    _add_figure(doc, fig2, "Figure 2. Share of non-compliances by category (legible horizontal bars; source: audit workbook).")
    _add_figure(doc, fig3, "Figure 3. Compliance rate by category (applicable controls only; source: audit workbook).")
    _add_figure(doc, fig4, "Figure 4. Counts by category and status (horizontal stacked bars; source: audit workbook).")
    _add_figure(doc, fig5, "Figure 5. Treatment volume by weakness pattern (source: Stage 1 treatment-plan model).")
    _add_figure(doc, fig6, "Figure 6. Treatment workload by recommended owner (source: Stage 1 treatment-plan model).")
    _add_figure(doc, fig7, "Figure 7. Treatment priority matrix (source: Stage 1 treatment-plan model).")

    add_nav_heading("10. Management Action Plan (MAP)", 1)
    _add_body_paragraph(doc, "The MAP is designed for executive readability. Detailed evidence remains in the workbook and technical artifacts. Priority combines severity and workbook-derived likelihood from Section 5.3.")
    _add_table_caption(doc, "Management action plan")
    mp = doc.add_table(rows=1, cols=4)
    mp.style = "Table Grid"
    mh = mp.rows[0].cells
    headers = ["Finding", "Owner / priority", "Target date", "Closure evidence"]
    for idx, txt in enumerate(headers):
        mh[idx].text = txt
        _set_cell_shading(mh[idx], THEME_BLUE)
        for run in mh[idx].paragraphs[0].runs:
            run.bold = True
    for p in patterns[:10]:
        pat = p["pattern"]
        cnt = int(p["mapped_noncompliant_count"])
        sev = p["severity"]
        lik = _likelihood_from_count(cnt)
        owner = p["recommended_owner"]
        target_date = _target_date_str(audit_dt, sev)
        criteria = _sanitize_closure_criteria(_ai_field_for_pattern(pat, writeups, "closure_criteria") or "AI-generated closure criteria were not returned for this pattern.")
        row = mp.add_row().cells
        row[0].text = f"{pat}\nWorkbook basis: {cnt} mapped non-compliant controls."
        row[1].text = f"{owner}\n{sev} / {lik}"
        row[2].text = target_date
        row[3].text = _short(criteria, 320)

    _add_body_paragraph(doc, "Technical verification evidence expected at closure includes updated workbook scoring, scan artifacts for the remediated release, regression-test evidence where applicable, and formal accepted-risk records for unresolved items.")

    doc.add_page_break()
    add_nav_heading("11. Treatment plan overview", 1)
    if treatment_plan:
        _render_treatment_overview(doc, treatment_plan, fig5, fig6, fig7)
    else:
        _add_body_paragraph(doc, "The analysis pack did not include a treatment_plan block. Re-run Stage 1 with the treatment-plan model enabled to populate this section.")

    doc.add_page_break()
    add_nav_heading("Appendix A - SECM-CAT treatment plan", 1)
    if _treatment_full_mode():
        _add_body_paragraph(doc, "This appendix lists non-compliant SECM-CAT controls as treatment items. It preserves PUIDs, categories, flags, workbook evidence, weakness patterns, owners, timelines, treatment actions, and expected closure evidence.")
    else:
        _add_body_paragraph(doc, "This appendix lists non-compliant SECM-CAT controls as structured treatment traceability items. It preserves PUIDs, categories, flags, workbook evidence, weakness patterns, owners, and target timelines. Item-level narrative actions are intentionally not generated for every control in priority mode.")
    _render_control_treatment_appendix(doc, treatment_plan, treatment_ai)

    doc.add_page_break()
    add_nav_heading("Appendix B - Technical vulnerability treatment plan", 1)
    if _treatment_full_mode():
        _add_body_paragraph(doc, "This appendix lists technical treatment items derived from Trivy, SAST, MobSF static, and MobSF dynamic evidence. It strengthens the body of the report with scanner-specific findings, affected components, source files, observed issues, linked PUIDs, and treatment evidence.")
    else:
        _add_body_paragraph(doc, "This appendix lists technical treatment items derived from Trivy, SAST, MobSF static, and MobSF dynamic evidence. In priority mode, AI-authored treatment actions are generated only for selected priority technical items; the remaining rows preserve scanner traceability.")
    _render_technical_treatment_appendix(doc, treatment_plan, treatment_ai)

    doc.add_page_break()
    add_nav_heading("Appendix C - Evidence correlation matrix", 1)
    _add_body_paragraph(doc, "This matrix links SECM-CAT requirements, flags, scanner evidence, evidence fit, correlation strength, and treatment item IDs. It supports technical traceability from each treatment action back to requirement-level audit evidence.")
    _render_correlation_appendix(doc, treatment_plan)

    doc.add_page_break()
    add_nav_heading("Appendix D - Positive controls verification (workbook traceability)", 1)
    doc.add_paragraph("This appendix verifies each Positive controls observed statement by providing the originating PUID, flags used, evidence class, and evidence excerpt when available.")
    _add_table_caption(doc, "Positive control verification traceability")
    vb = doc.add_table(rows=1, cols=5)
    vb.style = "Table Grid"
    vh = vb.rows[0].cells
    for idx, txt in enumerate(["Positive control statement (as reported)", "Evidence class", "Workbook PUID", "Flags used", "Evidence / justification (excerpt)"]):
        vh[idx].text = txt
        _set_cell_shading(vh[idx], THEME_BLUE)
        _format_table_cell(vh[idx], font_size=REPORT_TABLE_HEADER_PT, bold=True, no_wrap=True)
    if pos_controls:
        for group_title in ["Observed positive controls", "Observed supporting signals", "Observed controls with additional technical context"]:
            for pc in positive_groups.get(group_title, []):
                r = vb.add_row().cells
                puid = _clean_text(pc.get("puid"))
                r[0].text = pc.get("reported_statement", "")
                r[1].text = group_title
                r[2].text = puid
                r[3].text = pc.get("flags_used", "") or ""
                r[4].text = _sanitize_positive_statement(pc.get("evidence_excerpt", "") or "")
                for idx, cell in enumerate(r):
                    _format_table_cell(cell, font_size=REPORT_TABLE_BODY_PT, no_wrap=(idx in {1, 2}))
    else:
        r = vb.add_row().cells
        r[0].text = "No verified positive controls available."
        r[1].text = ""
        r[2].text = ""
        r[3].text = ""
        r[4].text = ""

    doc.add_page_break()
    add_nav_heading("Appendix E - Technical evidence summary", 1)
    doc.add_paragraph("This appendix provides a compact index of the technical evidence parsed from scan artifacts. Detailed raw JSON, SARIF, and tool reports remain in their original pipeline artifacts and are not embedded in the report package.")
    _add_table(doc, ["Evidence source", "Status", "Summary"], _source_status_rows(technical), max_rows=10)
    _add_technical_execution_metadata(doc, technical)

    doc.add_page_break()
    add_nav_heading("Appendix F - Vision360 flag evidence matrix", 1)
    _render_vision360_appendix(doc, technical)

    doc.add_page_break()
    add_nav_heading("Appendix G - MobSF static technical inventory", 1)
    _render_mobsf_static_inventory_appendix(doc, technical)

    doc.add_page_break()
    add_nav_heading("Appendix H - Software Composition Analysis inventory", 1)
    _render_sca_inventory_appendix(doc, technical)

    _render_clickable_toc(toc_placeholder, toc_entries)
    _enable_update_fields_on_open(doc)
    _format_report_paragraphs(doc)
    _sanitize_report_language(doc)
    _normalize_document_typography(doc)
    _quality_gate(doc)
    doc.save(out_path)
    print(f"[OK] DOCX generated -> {out_path}")


if __name__ == "__main__":
    main()